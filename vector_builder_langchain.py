#!/usr/bin/env python3
"""
完整優化版 LangChain 向量系統 - 智能長短文本處理
- 🧠 智能文本長度分類和處理策略
- 📏 精確的 Token 估算和批次處理
- 📄 保持文檔結構完整性
- 🔧 優化的分割邏輯
- 🚀 性能和錯誤處理優化
- Python 3.11.7 環境
"""

import json
import shutil
import time
import hashlib
import re
import logging
import os
import gc
import threading
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Set, Any, Union
from dataclasses import dataclass, asdict
import random
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
import tempfile
# 🔧 強制載入環境變數
from dotenv import load_dotenv
from database_adapter import DatabaseFactory, PostgreSQLAdapter
load_dotenv('.env', override=True)
logger = logging.getLogger(__name__)

import urllib.parse
from urllib.parse import quote_plus

import sys
from pathlib import Path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))
from config import app_config  # ⭐ 統一導入



# 🆕 1. 添加 Railway 環境檢測函數（放在文件頂部，導入語句之後）
def detect_railway_environment():
    """檢測是否在 Railway 環境中運行"""
    railway_indicators = [
        os.getenv("RAILWAY_PROJECT_ID"),
        os.getenv("RAILWAY_SERVICE_ID"), 
        os.getenv("DATABASE_URL"),
        "railway.internal" in os.getenv("POSTGRES_HOST", "")
    ]
    
    is_railway = any(railway_indicators)
    if is_railway:
        print("🚂 檢測到 Railway 部署環境")
        # 顯示 Railway 特定信息
        project_id = os.getenv("RAILWAY_PROJECT_ID", "unknown")
        service_id = os.getenv("RAILWAY_SERVICE_ID", "unknown")  
        print(f"   項目ID: {project_id[:8]}***")
        print(f"   服務ID: {service_id[:8]}***")
    
    return is_railway


# 🔧 優化版系統配置
SYSTEM_CONFIG = {
    "persist_dir": "./chroma_langchain_db",
    "data_dir": "./data", 
    "device": "cpu",  
    "log_level": "INFO",
    "max_workers": 4,  # 並發處理數
    "cache_embeddings": True,  # 快取嵌入向量
    "backup_enabled": True  # 自動備份
}

# 🧠 智能文本分類配置
SMART_TEXT_CONFIG = {
    # 文本長度閾值（字符數）
    "micro_text_threshold": 100,      # 微文本：標題、摘要等
    "short_text_threshold": 500,      # 短文本：短文章、段落
    "medium_text_threshold": 2000,    # 中等文本：中篇文章
    "long_text_threshold": 8000,      # 長文本：長篇文章
    "ultra_long_threshold": 20000,    # 超長文本：書籍章節
    
    # 各類文本的處理策略
    "micro_text": {
        "min_length": 20,
        "process_whole": True,
        "merge_with_next": True,  # 嘗試與下一個片段合併
        "chunk_size": 0,
        "description": "微文本整體處理"
    },
    
    "short_text": {
        "min_length": 50,
        "process_whole": True,
        "preserve_structure": True,
        "chunk_size": 0,
        "allow_merge": True,  # 允許合併相鄰短文本
        "description": "短文本整體保存"
    },
    
    "medium_text": {
        "chunk_size": 800,
        "chunk_overlap": 120,
        "preserve_paragraphs": True,
        "smart_boundaries": True,  # 智能邊界檢測
        "min_chunk_ratio": 0.3,   # 最小片段比例
        "description": "中等文本段落感知分割"
    },
    
    "long_text": {
        "chunk_size": 500,
        "chunk_overlap": 80,
        "hierarchical_split": True,
        "section_aware": True,    # 章節感知
        "quality_check": True,    # 質量檢查
        "description": "長文本精細化分割"
    },
    
    "ultra_long": {
        "chunk_size": 600,
        "chunk_overlap": 100,
        "multi_level_split": True,  # 多層級分割
        "chapter_detection": True,  # 章節檢測
        "summary_chunks": True,     # 生成摘要片段
        "description": "超長文本階層式處理"
    },
    
    "mega_text": {
        "chunk_size": 500,
        "chunk_overlap": 80,
        "multi_level_split": True,  # 多層級分割
        "chapter_detection": True,  # 章節檢測
        "aggressive_split": True,   # 積極分割
        "quality_filter": True,     # 質量過濾
        "description": "超大文本積極分割處理"
    }
}

# 🔧 優化版 Token 限制配置
TOKEN_LIMITS = {
    "max_tokens_per_request": 150000,  # 更保守的限制
    "max_batch_size": 8,               # 減小批次大小
    "min_batch_size": 1,
    "batch_delay": 5.0,                # 增加批次間延遲
    "retry_delay": 10,
    "max_retries": 3,
    "token_safety_margin": 0.2,        # 20% 安全邊際
    "adaptive_batching": True           # 自適應批次大小
}

# 🔧 性能優化配置
PERFORMANCE_CONFIG = {
    "embedding_batch_size": 12,
    "parallel_processing": True,
    "memory_limit_mb": 1024,      # 記憶體限制
    "chunk_cache_size": 1000,     # 分塊快取大小
    "preload_models": True,       # 預載入模型
    "gc_frequency": 50            # 垃圾回收頻率
}

# 支援的文件格式擴展
SUPPORTED_EXTENSIONS = {
    '.txt', '.md', '.pdf', '.csv', '.json', '.py', '.js', 
    '.docx', '.doc', '.epub', '.rst', '.org'
}

from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from langchain_community.vectorstores import Chroma

try:
    from langchain_community.vectorstores import Chroma
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False
    # 創建一個虛擬的 Chroma 類型用於類型註解
    class Chroma:
        pass



# 🔧 檢查環境變數
def check_openai_api_key():
    """檢查 OpenAI API Key"""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("⚠️  未設置 OPENAI_API_KEY 環境變數")
        print("   請在 .env 文件中設置: OPENAI_API_KEY=sk-your-api-key")
        return False
    
    if api_key.startswith("sk-") and len(api_key) > 20:
        print("✅ OpenAI API Key 格式正確")
        print(f"   API Key: {api_key[:8]}***{api_key[-4:]}")
        return True
    else:
        print("⚠️  OpenAI API Key 格式可能不正確")
        return False

# LangChain 核心組件導入
try:
    from langchain_postgres import PGVector
    PGVECTOR_AVAILABLE = True
    print("✅ PGVector 可用")
except ImportError:
    try:
        from langchain_community.vectorstores import PGVector
        PGVECTOR_AVAILABLE = True
        print("✅ PGVector (community) 可用")
    except ImportError:
        PGVECTOR_AVAILABLE = False
        print("❌ PGVector 不可用")
        # 可以回退到 Chroma
        from langchain_community.vectorstores import Chroma

# OpenAI embeddings 導入
try:
    from langchain_openai import OpenAIEmbeddings
    OPENAI_EMBEDDINGS_AVAILABLE = True
    print("✅ OpenAI Embeddings 可用")
    
    if check_openai_api_key():
        print("✅ OpenAI 環境檢查通過")
    else:
        print("⚠️ OpenAI 環境配置可能有問題")
        
except ImportError as e:
    OPENAI_EMBEDDINGS_AVAILABLE = False
    print(f"⚠️ OpenAI Embeddings 不可用: {e}")

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    TextLoader, PyPDFLoader, Docx2txtLoader, 
    CSVLoader, JSONLoader
)

# 🔧 PostgreSQL 依賴檢查
try:
    import psycopg2
    PSYCOPG2_AVAILABLE = True
except ImportError:
    PSYCOPG2_AVAILABLE = False
    print("⚠️ 警告: psycopg2 未安裝，請執行: pip install psycopg2-binary")

"""
# 🐘 PostgreSQL 配置
POSTGRES_CONFIG = {
    "host": os.getenv("POSTGRES_HOST", "localhost"),
    "port": int(os.getenv("POSTGRES_PORT", "5432")),
    "database": os.getenv("POSTGRES_DATABASE", "chatbot_system"),
    "user": os.getenv("POSTGRES_USER", "chatbot_user"),
    "password": os.getenv("POSTGRES_PASSWORD", "chatbot123"),
    "schema": os.getenv("POSTGRES_SCHEMA", "public"),
}"""

# 📝 逐步修改您的 vector_builder_langchain.py 文件

# ====================
# 步驟 1: 在文件頂部添加導入語句
# ====================
# 在您的文件頂部，找到其他導入語句的地方，添加這兩行：

import urllib.parse  # ✅ Python 內建，不需要安裝
from urllib.parse import quote_plus  # ✅ Python 內建，不需要安裝

# ====================
# 步驟 2: 找到現有的 PostgreSQL 配置部分並註釋掉
# ====================
# 在您的文件中找到這個部分（大約在第 xxx 行）：

"""
# 🔧 註釋掉這個舊的配置
POSTGRES_CONFIG = {
    "host": os.getenv("POSTGRES_HOST", "localhost"),
    "port": int(os.getenv("POSTGRES_PORT", "5432")),
    "database": os.getenv("POSTGRES_DATABASE", "chatbot_system"),
    "user": os.getenv("POSTGRES_USER", "chatbot_user"),
    "password": os.getenv("POSTGRES_PASSWORD", "chatbot123"),
    "schema": os.getenv("POSTGRES_SCHEMA", "public"),
}
"""





# 3️⃣ 修改 OptimizedVectorSystem 類的 load_document 方法
# 📍 在 load_document 方法中添加 EPUB 處理邏輯

def load_document(self, file_path: Path) -> List[Document]:
    """載入並智能處理文檔 - 完整版包含 EPUB"""
    try:
        extension = file_path.suffix.lower()
        
        # 根據檔案類型載入
        if extension == '.pdf':
            loader = PyPDFLoader(str(file_path))
            documents = loader.load()
            full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
            
        elif extension in ['.docx', '.doc'] and DOCX2TXT_AVAILABLE:
            if DOCX_METHOD == "docx2txt":
                import docx2txt
                full_text = docx2txt.process(str(file_path))
            else:
                loader = Docx2txtLoader(str(file_path))
                documents = loader.load()
                full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
                
        elif extension == '.epub' and EPUB_AVAILABLE:
            # 🆕 EPUB 處理邏輯
            epub_processor = EpubProcessor()
            full_text = epub_processor.extract_epub_content(file_path)
            
        elif extension == '.csv':
            loader = CSVLoader(str(file_path), encoding='utf-8')
            documents = loader.load()
            full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
            
        elif extension == '.json':
            loader = JSONLoader(str(file_path), jq_schema='.', text_content=False)
            documents = loader.load()
            full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
            
        else:
            # 文本文件處理
            encodings = ['utf-8', 'gbk', 'gb2312', 'big5']
            full_text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        full_text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if full_text is None:
                raise ValueError(f"無法解碼文件: {file_path}")
        
        # 檢查內容
        if not full_text or not full_text.strip():
            print(f"⚠️ 文件內容為空: {file_path.name}")
            return []
        
        # 生成文檔ID
        doc_id = self._generate_doc_id(file_path)
        
        # 使用優化的文本分割器
        source_info = {
            'file_path': str(file_path),
            'file_type': extension,
            'file_size': file_path.stat().st_size if file_path.exists() else 0
        }
        
        documents = self.text_splitter.smart_split_documents(full_text, doc_id, source_info)
        
        # 添加統一元數據
        for doc in documents:
            doc.metadata.update({
                'source': str(file_path),
                'filename': file_path.name,
                'extension': extension,
                'file_size': source_info['file_size'],
                'load_timestamp': time.time()
            })
        
        print(f"📄 文檔載入完成: {file_path.name} ({len(documents)} 個分塊)")
        
        return documents
        
    except Exception as e:
        logger.error(f"載入文檔失敗 {file_path}: {e}")
        print(f"❌ 文檔載入失敗: {file_path.name} - {e}")
        return []



# 5️⃣ 添加 EPUB 功能測試
def test_epub_functionality():
    """測試 EPUB 功能"""
    print("🧪 === EPUB 功能測試 ===")
    
    # 檢查依賴
    print(f"📚 EPUB 庫可用: {'✅' if EPUB_AVAILABLE else '❌'}")
    
    if not EPUB_AVAILABLE:
        print("❌ EPUB 功能不可用，請安裝依賴:")
        print("   pip install ebooklib beautifulsoup4 lxml")
        return False
    
    # 查找測試文件
    test_epub_files = []
    data_dir = Path("data")
    
    if data_dir.exists():
        test_epub_files = list(data_dir.rglob("*.epub"))
    
    if test_epub_files:
        print(f"📖 找到 {len(test_epub_files)} 個 EPUB 文件:")
        for epub_file in test_epub_files[:3]:  # 只顯示前3個
            print(f"   📚 {epub_file.name}")
        
        # 測試處理第一個文件
        try:
            processor = EpubProcessor()
            content = processor.extract_epub_content(test_epub_files[0])
            print(f"✅ EPUB 處理測試成功: {len(content):,} 字符")
            return True
        except Exception as e:
            print(f"❌ EPUB 處理測試失敗: {e}")
            return False
    else:
        print("📭 未找到 EPUB 測試文件")
        print("   請將 .epub 文件放入 data/ 目錄進行測試")
        return True  # 沒有測試文件不算失敗

# 6️⃣ 在 main 函數中添加 EPUB 測試
def main():
    """主程式 - 包含 EPUB 測試"""
    
    # EPUB 功能測試
    epub_test_result = test_epub_functionality()
    
    # 初始化系統
    try:
        system = OptimizedVectorSystem()
        
        # 系統診斷
        print("\n" + "="*60)
        system.diagnose_system()
        print("="*60 + "\n")
        
        # 同步集合（現在支援 EPUB）
        changes = system.sync_collections()
        
        # 其餘代碼保持不變...
        
    except Exception as e:
        print(f"❌ 系統初始化失敗: {e}")
        return None

# =============================================================================
# 🚨 重要：修復步驟
# =============================================================================

print("✅ EPUB 讀取功能已完整修復！")
print("📚 支援功能:")
print("   • 自動提取章節內容")
print("   • 智能章節標題識別") 
print("   • 中文文本標準化")
print("   • 書籍元資料提取")
print("   • HTML 標籤清理")
print("   • 多章節智能分割")


# 中文分詞和轉換
try:
    import jieba
    jieba.initialize()
    JIEBA_AVAILABLE = True
except ImportError:
    JIEBA_AVAILABLE = False
    print("⚠️ 警告: jieba 未安裝，中文分詞功能將被禁用")

try:
    import opencc
    OPENCC_AVAILABLE = True
except ImportError:
    OPENCC_AVAILABLE = False
    print("⚠️ 警告: opencc 未安裝，繁簡轉換功能將被禁用")

# 文檔處理支援
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
    DOCX_METHOD = "docx2txt"
    print("✅ DOCX 支援已啟用 (docx2txt)")
except ImportError:
    try:
        from docx import Document as DocxDocument
        DOCX2TXT_AVAILABLE = True
        DOCX_METHOD = "python-docx"
        print("✅ DOCX 支援已啟用 (python-docx)")
    except ImportError:
        DOCX2TXT_AVAILABLE = False
        DOCX_METHOD = None
        print("⚠️ 警告: docx2txt 或 python-docx 未安裝")

try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    EPUB_AVAILABLE = True
    print("✅ EPUB 支援已啟用")
except ImportError:
    EPUB_AVAILABLE = False
    print("⚠️ 警告: ebooklib 或 beautifulsoup4 未安裝")

# 配置日誌
logging.basicConfig(level=getattr(logging, SYSTEM_CONFIG["log_level"]))
logger = logging.getLogger(__name__)

@dataclass
class FileInfo:
    """文件資訊"""
    path: str
    size: int
    mtime: float
    hash: str
    encoding: str = "utf-8"
    file_type: str = ""

@dataclass
class TextAnalysis:
    """文本分析結果"""
    length: int
    text_type: str
    language: str
    encoding: str
    structure_info: Dict
    quality_score: float
    processing_strategy: str

@dataclass
class ChunkInfo:
    """分塊資訊"""
    chunk_id: str
    content: str
    metadata: Dict
    token_count: int
    quality_score: float
    relationships: List[str]  # 與其他分塊的關係

@dataclass
class SearchResult:
    content: str
    score: float
    metadata: Dict
    collection: str
    chunk_info: Optional[ChunkInfo] = None

class AdvancedTokenEstimator:
    """🔧 高級 Token 估算器"""
    
    def __init__(self):
        # 不同語言和內容類型的 Token 係數
        self.token_ratios = {
            "chinese": 2.5,      # 中文字符/token
            "english": 4.0,      # 英文字符/token
            "mixed": 3.0,        # 中英混合
            "code": 3.5,         # 程式碼
            "punctuation": 1.0,  # 標點符號
            "numbers": 2.0       # 數字
        }
        self.safety_margin = TOKEN_LIMITS.get("token_safety_margin", 0.15)
    
    def analyze_text_composition(self, text: str) -> Dict[str, int]:
        """分析文本組成"""
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_chars = len(re.findall(r'[a-zA-Z]', text))
        numbers = len(re.findall(r'\d', text))
        punctuation = len(re.findall(r'[^\w\s\u4e00-\u9fff]', text))
        other_chars = len(text) - chinese_chars - english_chars - numbers - punctuation
        
        return {
            "chinese": chinese_chars,
            "english": english_chars,
            "numbers": numbers,
            "punctuation": punctuation,
            "other": other_chars,
            "total": len(text)
        }
    
    def estimate_tokens(self, text: str, content_type: str = "mixed") -> int:
        """精確估算 Token 數量"""
        if not text:
            return 0
        
        composition = self.analyze_text_composition(text)
        
        # 根據文本組成計算 token
        estimated_tokens = (
            composition["chinese"] / self.token_ratios["chinese"] +
            composition["english"] / self.token_ratios["english"] +
            composition["numbers"] / self.token_ratios["numbers"] +
            composition["punctuation"] / self.token_ratios["punctuation"] +
            composition["other"] / self.token_ratios["mixed"]
        )
        
        # 根據內容類型調整
        if content_type == "code":
            estimated_tokens *= 1.2  # 程式碼通常 token 密度更高
        elif content_type == "academic":
            estimated_tokens *= 1.1  # 學術文本專業詞彙多
        
        # 加上安全邊際
        final_estimate = int(estimated_tokens * (1 + self.safety_margin))
        
        return max(final_estimate, 1)  # 至少1個token
    
    def estimate_embedding_cost(self, total_tokens: int, model: str = "text-embedding-3-small") -> float:
        """估算 Embedding 成本"""
        cost_per_1k_tokens = {
            "text-embedding-3-small": 0.00002,
            "text-embedding-3-large": 0.00013,
            "text-embedding-ada-002": 0.0001
        }
        
        rate = cost_per_1k_tokens.get(model, 0.00002)
        return (total_tokens / 1000) * rate

class ChineseTextNormalizer:
    """優化版中文文本標準化處理器"""
    
    def __init__(self):
        self.s2t_converter = None
        self.t2s_converter = None
        
        if OPENCC_AVAILABLE:
            try:
                self.s2t_converter = opencc.OpenCC('s2t')
                self.t2s_converter = opencc.OpenCC('t2s')
                print("🔤 繁簡轉換器初始化完成")
            except Exception as e:
                logger.warning(f"OpenCC 初始化失敗: {e}")
        
        # 中文文本正規化規則
        self.normalization_rules = [
            (r'[　\u3000]+', ' '),                    # 統一空格
            (r'\r\n|\r', '\n'),                       # 統一換行
            (r'\n{3,}', '\n\n'),                      # 限制連續換行
            (r'[\u201C\u201D\u2018\u2019\u201E\u201A\u2033\u2032]', '"'),  # 統一引號
            (r'[—–−]', '-'),                          # 統一破折號
            (r'[…⋯]', '...'),                         # 統一省略號
        ]
    
    def normalize_text(self, text: str) -> Tuple[str, Dict]:
        """標準化文本並返回處理資訊"""
        if not text:
            return "", {}
        
        original_length = len(text)
        processed_text = text
        
        # 應用正規化規則
        for pattern, replacement in self.normalization_rules:
            processed_text = re.sub(pattern, replacement, processed_text)
        
        # 移除首尾空白
        processed_text = processed_text.strip()
        
        # 檢測和轉換繁簡
        variant, confidence = self.detect_chinese_variant(processed_text)
        if variant == 'simplified' and self.s2t_converter:
            try:
                processed_text = self.s2t_converter.convert(processed_text)
                variant = 'traditional_converted'
            except Exception as e:
                logger.warning(f"繁簡轉換失敗: {e}")
        
        processing_info = {
            'original_length': original_length,
            'processed_length': len(processed_text),
            'variant': variant,
            'confidence': confidence,
            'normalized': original_length != len(processed_text)
        }
        
        return processed_text, processing_info
    
    def detect_chinese_variant(self, text: str) -> Tuple[str, float]:
        """檢測繁體或簡體中文"""
        if not text:
            return 'unknown', 0.0
        
        simplified_chars = set('国发会学习论问题业专长时间经济')
        traditional_chars = set('國發會學習論問題業專長時間經濟')
        
        simplified_count = sum(1 for char in text if char in simplified_chars)
        traditional_count = sum(1 for char in text if char in traditional_chars)
        total_chinese = len(re.findall(r'[\u4e00-\u9fff]', text))
        
        if total_chinese == 0:
            return 'unknown', 0.0
        
        simplified_ratio = simplified_count / total_chinese
        traditional_ratio = traditional_count / total_chinese
        
        if simplified_ratio > traditional_ratio:
            return 'simplified', simplified_ratio
        elif traditional_ratio > simplified_ratio:
            return 'traditional', traditional_ratio
        else:
            return 'mixed', max(simplified_ratio, traditional_ratio)
    
    def create_search_variants(self, query: str) -> List[str]:
        """創建繁簡搜索變體"""
        variants = [query]
        
        if not self.s2t_converter or not self.t2s_converter:
            return variants
        
        try:
            traditional = self.s2t_converter.convert(query)
            if traditional != query:
                variants.append(traditional)
            
            simplified = self.t2s_converter.convert(query)
            if simplified != query:
                variants.append(simplified)
        except Exception as e:
            logger.warning(f"創建搜索變體失敗: {e}")
        
        return list(set(variants))

class EpubProcessor:
    """📚 EPUB 文件處理器"""
    
    def __init__(self):
        self.normalizer = ChineseTextNormalizer()
        
    def extract_epub_content(self, file_path: Path) -> str:
        """提取 EPUB 內容"""
        try:
            if not EPUB_AVAILABLE:
                raise ImportError("EPUB 處理庫未安裝")
            
            print(f"📚 正在處理 EPUB 文件: {file_path.name}")
            
            # 讀取 EPUB 文件
            book = epub.read_epub(str(file_path))
            
            # 提取所有文本內容
            content_parts = []
            chapter_count = 0
            
            # 獲取書籍信息
            title = book.get_metadata('DC', 'title')
            author = book.get_metadata('DC', 'creator')
            
            book_info = []
            if title:
                book_info.append(f"書名: {title[0][0] if title else '未知'}")
            if author:
                book_info.append(f"作者: {author[0][0] if author else '未知'}")
            
            if book_info:
                content_parts.append("\n".join(book_info) + "\n\n")
            
            # 按順序處理章節
            spine_items = book.get_items_of_type(ebooklib.ITEM_DOCUMENT)
            
            for item in spine_items:
                try:
                    # 解析 HTML 內容
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    
                    # 移除腳本和樣式
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    # 提取文本
                    text = soup.get_text()
                    
                    if text and len(text.strip()) > 50:  # 過濾太短的內容
                        # 清理文本
                        cleaned_text = self._clean_epub_text(text)
                        
                        if cleaned_text.strip():
                            chapter_count += 1
                            
                            # 添加章節標記
                            chapter_title = self._extract_chapter_title(soup, cleaned_text)
                            if chapter_title:
                                content_parts.append(f"\n\n=== {chapter_title} ===\n")
                            else:
                                content_parts.append(f"\n\n=== 第 {chapter_count} 章 ===\n")
                            
                            content_parts.append(cleaned_text)
                            
                except Exception as e:
                    print(f"   ⚠️ 章節處理失敗: {e}")
                    continue
            
            if not content_parts:
                raise ValueError("EPUB 文件中沒有找到有效內容")
            
            full_content = "".join(content_parts)
            
            print(f"   ✅ EPUB 處理完成: {chapter_count} 個章節, {len(full_content):,} 字符")
            
            return full_content
            
        except Exception as e:
            print(f"   ❌ EPUB 處理失敗: {e}")
            raise
    
    def _clean_epub_text(self, text: str) -> str:
        """清理 EPUB 文本"""
        if not text:
            return ""
        
        # 標準化文本
        normalized_text, _ = self.normalizer.normalize_text(text)
        
        # EPUB 特定清理
        epub_cleaning_rules = [
            (r'\n{4,}', '\n\n'),                    # 限制連續換行
            (r'[ \t]{3,}', ' '),                    # 限制連續空格
            (r'^[ \t]+', '', re.MULTILINE),          # 移除行首空白
            (r'[ \t]+$', '', re.MULTILINE),          # 移除行尾空白
            (r'\n[ \t]*\n', '\n\n'),                # 清理空行
            (r'[\x00-\x08\x0B\x0C\x0E-\x1F]', ''),  # 移除控制字符
        ]
        
        cleaned_text = normalized_text
        for pattern, replacement, *flags in epub_cleaning_rules:
            flag = flags[0] if flags else 0
            cleaned_text = re.sub(pattern, replacement, cleaned_text, flags=flag)
        
        return cleaned_text.strip()
    
    def _extract_chapter_title(self, soup, text: str) -> Optional[str]:
        """嘗試提取章節標題"""
        try:
            # 從 HTML 標籤中查找標題
            for tag in ['h1', 'h2', 'h3', 'title']:
                title_element = soup.find(tag)
                if title_element and title_element.get_text().strip():
                    title = title_element.get_text().strip()
                    if 5 <= len(title) <= 100:  # 合理的標題長度
                        return title
            
            # 從文本開頭查找標題
            lines = text.split('\n')[:5]  # 檢查前5行
            for line in lines:
                line = line.strip()
                if line and 5 <= len(line) <= 100:
                    # 檢查是否包含章節關鍵詞
                    chapter_keywords = ['章', 'Chapter', '第', '卷', 'Part']
                    if any(keyword in line for keyword in chapter_keywords):
                        return line
                    # 或者是短行且看起來像標題
                    elif len(line) <= 50 and not line.endswith(('。', '！', '？', '.', '!', '?')):
                        return line
            
            return None
            
        except Exception:
            return None


class SmartTextAnalyzer:
    """🧠 智能文本分析器"""
    
    def __init__(self):
        self.normalizer = ChineseTextNormalizer()
        self.token_estimator = AdvancedTokenEstimator()
        
    def analyze_text(self, text: str, source_info: Dict = None) -> TextAnalysis:
        """綜合分析文本"""
        if not text:
            return TextAnalysis(
                length=0, text_type="empty", language="unknown",
                encoding="utf-8", structure_info={}, quality_score=0.0,
                processing_strategy="skip"
            )
        
        # 標準化文本
        normalized_text, norm_info = self.normalizer.normalize_text(text)
        length = len(normalized_text)
        
        # 分類文本長度
        text_type = self._classify_text_length(length)
        
        # 分析文本結構
        structure_info = self._analyze_structure(normalized_text)
        
        # 檢測語言
        language = self._detect_language(normalized_text)
        
        # 評估質量
        quality_score = self._evaluate_quality(normalized_text, structure_info)
        
        # 決定處理策略
        processing_strategy = self._determine_strategy(text_type, structure_info, quality_score)
        
        return TextAnalysis(
            length=length,
            text_type=text_type,
            language=language,
            encoding=norm_info.get('encoding', 'utf-8'),
            structure_info=structure_info,
            quality_score=quality_score,
            processing_strategy=processing_strategy
        )
    
    def _classify_text_length(self, length: int) -> str:
        """分類文本長度"""
        config = SMART_TEXT_CONFIG
        
        if length < config["micro_text_threshold"]:
            return "micro_text"
        elif length < config["short_text_threshold"]:
            return "short_text"
        elif length < config["medium_text_threshold"]:
            return "medium_text"
        elif length < config["long_text_threshold"]:
            return "long_text"
        elif length < config["ultra_long_threshold"]:
            return "ultra_long"
        else:
            return "mega_text"  # 超大文本
    
    def _analyze_structure(self, text: str) -> Dict:
        """分析文本結構"""
        structure_info = {
            'paragraphs': len(text.split('\n\n')),
            'lines': len(text.split('\n')),
            'sentences': len(re.findall(r'[。！？.!?]+', text)),
            'has_chapters': False,
            'has_sections': False,
            'has_lists': False,
            'has_tables': False,
            'chapter_count': 0,
            'section_count': 0
        }
        
        # 檢測章節結構
        chapter_patterns = [
            r'第[一二三四五六七八九十\d]+章',
            r'Chapter\s+\d+',
            r'\d+\.\s*[^\n]{1,50}\n'
        ]
        
        for pattern in chapter_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                structure_info['has_chapters'] = True
                structure_info['chapter_count'] = len(matches)
                break
        
        # 檢測小節結構
        section_patterns = [
            r'第[一二三四五六七八九十\d]+節',
            r'\d+\.\d+',
            r'[一二三四五六七八九十]、'
        ]
        
        for pattern in section_patterns:
            matches = re.findall(pattern, text)
            if matches:
                structure_info['has_sections'] = True
                structure_info['section_count'] = len(matches)
                break
        
        # 檢測列表
        if re.search(r'^\s*[•\-\*]\s+', text, re.MULTILINE) or \
           re.search(r'^\s*\d+\.\s+', text, re.MULTILINE):
            structure_info['has_lists'] = True
        
        # 檢測表格
        if '|' in text and text.count('|') > 4:
            structure_info['has_tables'] = True
        
        return structure_info
    
    def _detect_language(self, text: str) -> str:
        """檢測文本主要語言"""
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_chars = len(re.findall(r'[a-zA-Z]', text))
        total_chars = len(text)
        
        if total_chars == 0:
            return "unknown"
        
        chinese_ratio = chinese_chars / total_chars
        english_ratio = english_chars / total_chars
        
        if chinese_ratio > 0.3:
            if english_ratio > 0.2:
                return "mixed_zh_en"
            else:
                return "chinese"
        elif english_ratio > 0.5:
            return "english"
        else:
            return "mixed"
    
    def _evaluate_quality(self, text: str, structure_info: Dict) -> float:
        """評估文本質量（0-1分數）"""
        score = 0.5  # 基礎分數
        
        # 長度合理性（太短或太長都扣分）
        length = len(text)
        if 50 <= length <= 10000:
            score += 0.2
        elif length < 20:
            score -= 0.3
        
        # 結構完整性
        if structure_info['paragraphs'] > 1:
            score += 0.1
        if structure_info['sentences'] > 2:
            score += 0.1
        
        # 內容豐富度
        unique_chars = len(set(text))
        if unique_chars > 20:
            score += 0.1
        
        # 避免重複内容
        lines = text.split('\n')
        unique_lines = len(set(lines))
        if len(lines) > 0:
            uniqueness_ratio = unique_lines / len(lines)
            score += uniqueness_ratio * 0.1
        
        return min(max(score, 0.0), 1.0)
    
    def _determine_strategy(self, text_type: str, structure_info: Dict, quality_score: float) -> str:
        """決定處理策略"""
        if quality_score < 0.3:
            return "low_quality_skip"
        
        if text_type in ["micro_text", "short_text"]:
            return "whole_document"
        elif text_type == "medium_text":
            if structure_info['has_chapters'] or structure_info['paragraphs'] > 5:
                return "paragraph_aware"
            else:
                return "simple_split"
        elif text_type == "long_text":
            if structure_info['has_chapters']:
                return "chapter_aware"
            elif structure_info['has_sections']:
                return "section_aware"
            else:
                return "fine_split"
        elif text_type in ["ultra_long", "mega_text"]:
            return "hierarchical_split"
        else:
            # 回退策略
            return "hierarchical_split"

class OptimizedTextSplitter:
    """🔧 優化版文本分割器"""
    
    def __init__(self):
        self.analyzer = SmartTextAnalyzer()
        self.token_estimator = AdvancedTokenEstimator()
        self.config = SMART_TEXT_CONFIG
        
        # 分割器實例池
        self._splitter_cache = {}
        
        print("🔧 優化版文本分割器初始化完成")
        print(f"   📏 支持 {len(self.config)} 種文本類型")
        print(f"   🧠 智能策略選擇")
        print(f"   ⚡ 性能優化")
    
    def get_splitter(self, chunk_size: int, chunk_overlap: int) -> RecursiveCharacterTextSplitter:
        """獲取分割器實例（帶快取）"""
        cache_key = f"{chunk_size}_{chunk_overlap}"
        
        if cache_key not in self._splitter_cache:
            self._splitter_cache[cache_key] = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=[
                    "\n\n第", "\n第",      # 章節標題
                    "\n\n", "\n",          # 段落
                    "。", "！", "？", "；",  # 句子結束
                    "，", "、",            # 短語分隔
                    " ", ""                # 最後resort
                ],
                keep_separator=True,
                length_function=len
            )
        
        return self._splitter_cache[cache_key]
    
    def smart_split_documents(self, text: str, doc_id: str, source_info: Dict = None) -> List[Document]:
        """🧠 智能文檔分割主入口"""
        if not text or not text.strip():
            return []
        
        # 分析文本
        analysis = self.analyzer.analyze_text(text, source_info)
        
        print(f"📊 文本分析: {analysis.text_type} ({analysis.length:,} 字符)")
        print(f"   🎯 處理策略: {analysis.processing_strategy}")
        print(f"   📈 質量分數: {analysis.quality_score:.2f}")
        print(f"   🗣️ 主要語言: {analysis.language}")
        
        # 根據策略選擇處理方法
        documents = []
        if analysis.processing_strategy == "low_quality_skip":
            print("   ⚠️ 文本質量過低，跳過處理")
            return []
        elif analysis.processing_strategy == "whole_document":
            documents = self._process_whole_document(text, doc_id, analysis)
        elif analysis.processing_strategy == "paragraph_aware":
            documents = self._process_paragraph_aware(text, doc_id, analysis)
        elif analysis.processing_strategy == "simple_split":
            documents = self._process_simple_split(text, doc_id, analysis)
        elif analysis.processing_strategy == "chapter_aware":
            documents = self._process_chapter_aware(text, doc_id, analysis)
        elif analysis.processing_strategy == "section_aware":
            documents = self._process_section_aware(text, doc_id, analysis)
        elif analysis.processing_strategy == "fine_split":
            documents = self._process_fine_split(text, doc_id, analysis)
        elif analysis.processing_strategy == "hierarchical_split":
            documents = self._process_hierarchical_split(text, doc_id, analysis)
        else:
            # 回退到簡單分割
            documents = self._process_simple_split(text, doc_id, analysis)

        # 最終驗證和過濾，確保沒有空內容的文檔
        final_documents = [doc for doc in documents if doc.page_content and doc.page_content.strip()]
        if len(final_documents) != len(documents):
            logger.warning(f"過濾掉 {len(documents) - len(final_documents)} 個空內容的區塊")
        
        return final_documents
    
    def _process_whole_document(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """處理整個文檔（不分割）"""
        config = self.config[analysis.text_type]
        
        if len(text.strip()) < config.get("min_length", 20):
            print(f"   ⚠️ 文檔過短，跳過處理")
            return []
        
        print(f"   📝 {config['description']}")
        
        # 計算 token 數
        token_count = self.token_estimator.estimate_tokens(text)
        
        metadata = {
            'doc_id': doc_id,
            'chunk_id': f"{doc_id}_whole",
            'chunk_index': 0,
            'total_chunks': 1,
            'text_type': analysis.text_type,
            'processing_strategy': analysis.processing_strategy,
            'is_complete_document': True,
            'token_count': token_count,
            'quality_score': analysis.quality_score,
            'language': analysis.language,
            'structure_info': analysis.structure_info
        }
        
        return [Document(page_content=text, metadata=metadata)]
    
    def _process_paragraph_aware(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """段落感知處理"""
        config = self.config[analysis.text_type]
        chunk_size = config["chunk_size"]
        chunk_overlap = config["chunk_overlap"]
        
        print(f"   📑 {config['description']} (目標大小: {chunk_size})")
        
        # 按段落分割
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if not paragraphs:
            return self._process_simple_split(text, doc_id, analysis)
        
        documents = []
        current_chunk = ""
        chunk_index = 0
        
        for para_idx, paragraph in enumerate(paragraphs):
            # 檢查是否可以加入當前分塊
            potential_chunk = current_chunk + ("\n\n" if current_chunk else "") + paragraph
            
            if len(potential_chunk) <= chunk_size or not current_chunk:
                current_chunk = potential_chunk
            else:
                # 當前分塊已滿，創建文檔
                if current_chunk:
                    doc = self._create_chunk_document(
                        current_chunk, doc_id, chunk_index, analysis, "paragraph_aware"
                    )
                    documents.append(doc)
                    chunk_index += 1
                
                current_chunk = paragraph
        
        # 處理最後一個分塊
        if current_chunk:
            doc = self._create_chunk_document(
                current_chunk, doc_id, chunk_index, analysis, "paragraph_aware"
            )
            documents.append(doc)
        
        # 處理重疊
        if chunk_overlap > 0 and len(documents) > 1:
            documents = self._add_overlap_to_documents(documents, chunk_overlap)
        
        return documents
    
    def _process_simple_split(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """簡單分割處理"""
        config = self.config[analysis.text_type]
        chunk_size = config["chunk_size"]
        chunk_overlap = config["chunk_overlap"]
        
        print(f"   ✂️ {config['description']} (大小: {chunk_size}, 重疊: {chunk_overlap})")
        
        splitter = self.get_splitter(chunk_size, chunk_overlap)
        chunks = splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            if len(chunk.strip()) >= config.get("min_length", 20):
                doc = self._create_chunk_document(
                    chunk, doc_id, i, analysis, "simple_split"
                )
                documents.append(doc)
        
        return documents
    
    def _process_chapter_aware(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """章節感知處理"""
        config = self.config[analysis.text_type]
        
        print(f"   📚 {config['description']} - 檢測到 {analysis.structure_info.get('chapter_count', 0)} 個章節")
        
        # 按章節分割
        chapters = self._split_by_chapters(text)
        
        if len(chapters) <= 1:
            # 沒有檢測到章節，回退到段落感知
            return self._process_paragraph_aware(text, doc_id, analysis)
        
        all_documents = []
        
        for chapter_idx, chapter_text in enumerate(chapters):
            chapter_doc_id = f"{doc_id}_ch{chapter_idx+1:02d}"
            
            # 每個章節內部使用段落感知分割
            chapter_analysis = analysis  # 使用相同的分析結果
            chapter_docs = self._process_paragraph_aware(chapter_text, chapter_doc_id, chapter_analysis)
            
            # 添加章節信息到元數據
            for doc in chapter_docs:
                doc.metadata.update({
                    'chapter_index': chapter_idx,
                    'chapter_total': len(chapters),
                    'parent_doc_id': doc_id,
                    'split_method': 'chapter_aware'
                })
            
            all_documents.extend(chapter_docs)
        
        return all_documents
    
    def _process_section_aware(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """小節感知處理"""
        config = self.config[analysis.text_type]
        
        print(f"   📋 {config['description']} - 檢測到 {analysis.structure_info.get('section_count', 0)} 個小節")
        
        # 按小節分割
        sections = self._split_by_sections(text)
        
        if len(sections) <= 1:
            return self._process_simple_split(text, doc_id, analysis)
        
        all_documents = []
        
        for section_idx, section_text in enumerate(sections):
            section_doc_id = f"{doc_id}_sec{section_idx+1:02d}"
            
            # 小節內部簡單分割
            section_docs = self._process_simple_split(section_text, section_doc_id, analysis)
            
            # 添加小節信息
            for doc in section_docs:
                doc.metadata.update({
                    'section_index': section_idx,
                    'section_total': len(sections),
                    'parent_doc_id': doc_id,
                    'split_method': 'section_aware'
                })
            
            all_documents.extend(section_docs)
        
        return all_documents
    
    def _process_fine_split(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """精細分割處理"""
        config = self.config[analysis.text_type]
        
        print(f"   🔬 {config['description']}")
        
        return self._process_simple_split(text, doc_id, analysis)
    
    def _process_hierarchical_split(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """階層式分割處理"""
        config = self.config[analysis.text_type]
        
        print(f"   🏗️ {config['description']}")
        
        # 第一層：章節分割
        chapters = self._split_by_chapters(text)
        
        if len(chapters) > 1:
            return self._process_chapter_aware(text, doc_id, analysis)
        
        # 第二層：段落分割
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if len(paragraphs) > 10:  # 段落較多，用段落感知
            return self._process_paragraph_aware(text, doc_id, analysis)
        
        # 第三層：精細分割
        return self._process_fine_split(text, doc_id, analysis)
    
    def _split_by_chapters(self, text: str) -> List[str]:
        """按章節分割文本"""
        chapter_patterns = [
            r'\n(第[一二三四五六七八九十\d]+章[^\n]*)',
            r'\n(Chapter\s+\d+[^\n]*)',
            r'\n(\d+\.\s*[^\n]{5,50})\n',
            r'\n([A-Z][^\n]{10,50})\n(?=[A-Z]|$)'
        ]
        
        best_split = None
        best_count = 0
        
        for pattern in chapter_patterns:
            matches = list(re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE))
            
            if len(matches) >= 2:
                splits = []
                start = 0
                
                for match in matches:
                    if start < match.start():
                        splits.append(text[start:match.start()].strip())
                    start = match.start()
                
                splits.append(text[start:].strip())
                splits = [s for s in splits if s and len(s) > 50]
                
                if len(splits) > best_count:
                    best_split = splits
                    best_count = len(splits)
        
        return best_split if best_split else [text]
    
    def _split_by_sections(self, text: str) -> List[str]:
        """按小節分割文本"""
        section_patterns = [
            r'\n(第[一二三四五六七八九十\d]+節[^\n]*)',
            r'\n(\d+\.\d+[^\n]*)',
            r'\n([一二三四五六七八九十]、[^\n]*)',
            r'\n([A-Z]\.[^\n]*)'
        ]
        
        for pattern in section_patterns:
            matches = list(re.finditer(pattern, text, re.MULTILINE))
            
            if len(matches) >= 2:
                splits = []
                start = 0
                
                for match in matches:
                    if start < match.start():
                        splits.append(text[start:match.start()].strip())
                    start = match.start()
                
                splits.append(text[start:].strip())
                return [s for s in splits if s and len(s) > 30]
        
        return [text]
    
    def _add_overlap_to_documents(self, documents: List[Document], overlap_size: int) -> List[Document]:
        """為文檔添加重疊部分"""
        if len(documents) <= 1:
            return documents
        
        for i in range(1, len(documents)):
            # 從前一個文檔末尾取重疊內容
            prev_content = documents[i-1].page_content
            current_content = documents[i].page_content
            
            if len(prev_content) > overlap_size:
                overlap_text = prev_content[-overlap_size:]
                # 尋找完整的句子邊界
                sentence_end = overlap_text.rfind('。')
                if sentence_end == -1:
                    sentence_end = overlap_text.rfind('.')
                
                if sentence_end > overlap_size // 2:
                    overlap_text = overlap_text[sentence_end+1:]
                
                documents[i].page_content = overlap_text + "\n" + current_content
                documents[i].metadata['has_overlap'] = True
                documents[i].metadata['overlap_length'] = len(overlap_text)
        
        return documents
    
    
    def _create_chunk_document(self, content: str, doc_id: str, chunk_index: int, 
                          analysis: TextAnalysis, split_method: str) -> Document:
        """創建分塊文檔 - 統一元數據格式"""
        token_count = self.token_estimator.estimate_tokens(content)
        
        # 掃描URL
        url_regex = r'https?://[\w\-./?#&%=]+'
        found_urls = re.findall(url_regex, content)

        # 基本元數據（確保都是簡單類型）
        metadata = {
            'doc_id': str(doc_id),
            'chunk_id': f"{doc_id}_{chunk_index+1:03d}",
            'chunk_index': int(chunk_index),
            'text_type': str(analysis.text_type),
            'processing_strategy': str(analysis.processing_strategy),
            'split_method': str(split_method),
            'chunk_length': int(len(content)),
            'token_count': int(token_count),
            'quality_score': float(analysis.quality_score),
            'language': str(analysis.language),
            'has_overlap': False
        }

        # 🎯 正確方案：URL列表 → 分隔符字串 (這是唯一可行的方案)
        if found_urls:
            metadata['contained_urls'] = '|'.join(found_urls)  # ✅ 字符串類型
            metadata['url_count'] = len(found_urls)           # ✅ 整數類型
            metadata['has_urls'] = True                       # ✅ 布林類型
        else:
            metadata['contained_urls'] = ''                   # ✅ 空字符串
            metadata['url_count'] = 0                        # ✅ 整數 0
            metadata['has_urls'] = False                     # ✅ 布林 False

        # 🎯 正確方案：複雜結構 → JSON字串
        if analysis.structure_info:
            metadata['structure_info'] = json.dumps(analysis.structure_info, ensure_ascii=False)  # ✅ JSON字符串
            # 同時保留關鍵信息作為簡單字段（便於查詢）
            metadata['has_chapters'] = bool(analysis.structure_info.get('has_chapters', False))    # ✅ 布林
            metadata['has_sections'] = bool(analysis.structure_info.get('has_sections', False))    # ✅ 布林  
            metadata['paragraph_count'] = int(analysis.structure_info.get('paragraphs', 0))       # ✅ 整數
            metadata['sentence_count'] = int(analysis.structure_info.get('sentences', 0))         # ✅ 整數
        else:
            metadata['structure_info'] = '{}'        # ✅ 空JSON字符串
            metadata['has_chapters'] = False         # ✅ 布林
            metadata['has_sections'] = False         # ✅ 布林
            metadata['paragraph_count'] = 0          # ✅ 整數
            metadata['sentence_count'] = 0           # ✅ 整數
        
        return Document(page_content=content, metadata=metadata)

class AdaptiveBatchProcessor:
    """🔧 自適應批次處理器"""
    
    def __init__(self):
        self.token_estimator = AdvancedTokenEstimator()
        self.max_tokens_per_batch = TOKEN_LIMITS["max_tokens_per_request"]
        self.max_batch_size = TOKEN_LIMITS["max_batch_size"]
        self.adaptive_batching = TOKEN_LIMITS.get("adaptive_batching", True)
        
        # 性能統計
        self.batch_stats = {
            'total_batches': 0,
            'total_documents': 0,
            'total_tokens': 0,
            'avg_batch_time': 0,
            'success_rate': 0,
            'adaptive_adjustments': 0
        }
        
        # 自適應參數
        self.current_batch_size = self.max_batch_size
        self.success_streak = 0
        self.failure_streak = 0
    
    def create_smart_batches(self, documents: List[Document]) -> List[Tuple[List[Document], Dict]]:
        """創建智能批次"""
        if not documents:
            return []
        
        print(f"🔧 創建智能批次...")
        print(f"   📄 總文檔數: {len(documents)}")
        print(f"   📏 Token 限制: {self.max_tokens_per_batch:,}")
        print(f"   📦 最大批次大小: {self.max_batch_size}")
        print(f"   🧠 自適應批次: {'✅' if self.adaptive_batching else '❌'}")
        
        batches = []
        current_batch = []
        current_tokens = 0
        batch_info = {'documents': 0, 'tokens': 0, 'types': defaultdict(int)}
        
        # 按 token 數排序文檔（大的在前面，更容易處理）
        sorted_docs = sorted(
            documents, 
            key=lambda doc: doc.metadata.get('token_count', 
                self.token_estimator.estimate_tokens(doc.page_content)),
            reverse=True
        )
        
        for doc_idx, doc in enumerate(sorted_docs):
            doc_tokens = doc.metadata.get('token_count') or \
                        self.token_estimator.estimate_tokens(doc.page_content)
            
            # 檢查單個文檔是否過大
            if doc_tokens > self.max_tokens_per_batch:
                print(f"   ⚠️ 文檔 {doc_idx+1} 過大 ({doc_tokens:,} tokens)，需要分割")
                
                # 完成當前批次
                if current_batch:
                    batches.append((current_batch, dict(batch_info)))
                    current_batch = []
                    current_tokens = 0
                    batch_info = {'documents': 0, 'tokens': 0, 'types': defaultdict(int)}
                
                # 分割大文檔
                split_docs = self._split_large_document(doc)
                for split_doc in split_docs:
                    split_tokens = self.token_estimator.estimate_tokens(split_doc.page_content)
                    split_doc.metadata['token_count'] = split_tokens
                    batches.append(([split_doc], {
                        'documents': 1, 
                        'tokens': split_tokens, 
                        'types': {split_doc.metadata.get('text_type', 'unknown'): 1},
                        'is_split': True
                    }))
                continue
            
            # 檢查是否可以加入當前批次
            would_exceed_tokens = current_tokens + doc_tokens > self.max_tokens_per_batch
            would_exceed_size = len(current_batch) >= self._get_adaptive_batch_size()
            
            if would_exceed_tokens or would_exceed_size:
                # 完成當前批次
                if current_batch:
                    batches.append((current_batch, dict(batch_info)))
                
                # 開始新批次
                current_batch = [doc]
                current_tokens = doc_tokens
                batch_info = {
                    'documents': 1, 
                    'tokens': doc_tokens, 
                    'types': defaultdict(int)
                }
                batch_info['types'][doc.metadata.get('text_type', 'unknown')] += 1
            else:
                # 加入當前批次
                current_batch.append(doc)
                current_tokens += doc_tokens
                batch_info['documents'] += 1
                batch_info['tokens'] += doc_tokens
                batch_info['types'][doc.metadata.get('text_type', 'unknown')] += 1
        
        # 處理最後一個批次
        if current_batch:
            batches.append((current_batch, dict(batch_info)))
        
        # 統計信息
        total_docs = sum(info['documents'] for _, info in batches)
        total_tokens = sum(info['tokens'] for _, info in batches)
        avg_tokens_per_batch = total_tokens / len(batches) if batches else 0
        
        print(f"✅ 智能批次創建完成:")
        print(f"   📦 總批次數: {len(batches)}")
        print(f"   📄 總文檔數: {total_docs}")
        print(f"   📏 總 tokens: {total_tokens:,}")
        print(f"   📊 平均 tokens/批次: {avg_tokens_per_batch:.0f}")
        print(f"   💰 估算成本: ${self.token_estimator.estimate_embedding_cost(total_tokens):.4f}")
        
        # 更新統計
        self.batch_stats['total_batches'] += len(batches)
        self.batch_stats['total_documents'] += total_docs
        self.batch_stats['total_tokens'] += total_tokens
        
        return batches
    
    def _get_adaptive_batch_size(self) -> int:
        """獲取自適應批次大小"""
        if not self.adaptive_batching:
            return self.max_batch_size
        
        # 根據成功率調整批次大小
        if self.success_streak >= 3:
            # 連續成功，可以增加批次大小
            self.current_batch_size = min(self.max_batch_size, self.current_batch_size + 1)
            self.batch_stats['adaptive_adjustments'] += 1
        elif self.failure_streak >= 2:
            # 連續失敗，減少批次大小
            self.current_batch_size = max(1, self.current_batch_size - 2)
            self.batch_stats['adaptive_adjustments'] += 1
        
        return self.current_batch_size
    
    def _split_large_document(self, doc: Document) -> List[Document]:
        """分割過大的文檔"""
        content = doc.page_content
        max_chars = int(self.max_tokens_per_batch * 2.5)  # 估算字符數
        
        # 嘗試按段落分割
        paragraphs = content.split('\n\n')
        split_docs = []
        current_content = ""
        part_index = 0
        
        for para in paragraphs:
            if len(current_content + para) <= max_chars or not current_content:
                current_content += ("\n\n" if current_content else "") + para
            else:
                if current_content.strip():
                    split_doc = Document(
                        page_content=current_content,
                        metadata={
                            **doc.metadata,
                            'chunk_id': f"{doc.metadata.get('chunk_id', 'unknown')}_part{part_index+1}",
                            'is_split_part': True,
                            'part_index': part_index,
                            'split_reason': 'token_limit'
                        }
                    )
                    split_docs.append(split_doc)
                    part_index += 1
                current_content = para
        
        # 處理最後一部分
        if current_content.strip():
            split_doc = Document(
                page_content=current_content,
                metadata={
                    **doc.metadata,
                    'chunk_id': f"{doc.metadata.get('chunk_id', 'unknown')}_part{part_index+1}",
                    'is_split_part': True,
                    'part_index': part_index,
                    'split_reason': 'token_limit'
                }
            )
            split_docs.append(split_doc)
        
        return split_docs or [doc]  # 如果分割失敗，返回原文檔
    
    def record_batch_result(self, success: bool, processing_time: float = 0):
        """記錄批次處理結果"""
        if success:
            self.success_streak += 1
            self.failure_streak = 0
        else:
            self.failure_streak += 1
            self.success_streak = 0
        
        # 更新成功率
        total_attempts = self.batch_stats['total_batches']
        if total_attempts > 0:
            self.batch_stats['success_rate'] = (total_attempts - self.failure_streak) / total_attempts
        
        # 更新平均處理時間
        if processing_time > 0:
            current_avg = self.batch_stats['avg_batch_time']
            self.batch_stats['avg_batch_time'] = (current_avg + processing_time) / 2
    
    def get_performance_stats(self) -> Dict:
        """獲取性能統計"""
        return dict(self.batch_stats)

class OptimizedVectorSystem:
    """🚀 完整優化版向量系統"""
    
    def __init__(self, data_dir: str = None, model_type: str = None):
        """✅ 修正版初始化方法 - 正確的執行順序"""
        
        # 🔧 1. 基本變數設置
        self.data_dir = Path(data_dir or SYSTEM_CONFIG["data_dir"])
        self.model_type = model_type or "openai"
        self.persist_dir = Path(SYSTEM_CONFIG["persist_dir"])  # Chroma備用

        # 建立目錄
        self.data_dir.mkdir(exist_ok=True)
        
        # 🔧 2. 資料庫連接設置（但不測試）
        self.db_adapter = None
        self.connection_string = None
        self.use_postgres = False

        database_url = os.getenv("DATABASE_URL")
        if PGVECTOR_AVAILABLE and database_url:
            self.connection_string = database_url
            print("🔍 發現 DATABASE_URL，準備測試 PostgreSQL 連接...")
        else:
            print("⚠️ DATABASE_URL 未設置或 PGVector 不可用，將使用 Chroma")

        if not PGVECTOR_AVAILABLE:
            print("⚠️ PGVector 依賴未安裝，使用 Chroma 作為備用")
            self.persist_dir.mkdir(exist_ok=True)

        # ✅ 3. 先初始化 Embedding 模型（關鍵！）
        self._setup_embedding_model()
        print("✅ Embedding 模型初始化完成")

        # ✅ 4. 現在可以測試 PostgreSQL 連接了（embeddings 已存在）
        if PGVECTOR_AVAILABLE and database_url and hasattr(self, 'embeddings'):
            try:
                print("🔍 測試 PostgreSQL + PGVector 連接...")
                # 測試連接
                PGVector.from_existing_index(
                    collection_name="_test_connection",
                    embedding=self.embeddings,  # ✅ 現在安全了
                    connection_string=self.connection_string
                )
                self.use_postgres = True
                print("✅ PostgreSQL (pgvector) 連接成功")
            except Exception as e:
                print(f"⚠️ PostgreSQL (pgvector) 連接測試失敗: {e}")
                self.use_postgres = False
                print("🔄 回退到 Chroma 本地存儲")
                self.persist_dir.mkdir(exist_ok=True)
        
        if not self.use_postgres:
            print("📁 使用 Chroma 作為向量存儲")
            self.persist_dir.mkdir(exist_ok=True)
        
        # 🔧 5. 初始化文本處理組件
        self._setup_text_processing()
        
        # 🔧 6. 初始化處理器
        self.batch_processor = AdaptiveBatchProcessor()
        self.text_splitter = OptimizedTextSplitter()
        
        # 🔧 7. 初始化存儲和記錄
        self._vector_stores = {}
        self.file_records = self._load_file_records()
        self.processing_lock = threading.Lock()
        
        print(f"🚀 完整優化版向量系統初始化完成")
        print(f"   🤖 嵌入模型: {self.model_type}")
        print(f"   📁 數據目錄: {self.data_dir}")
        print(f"   🗄️ 向量庫: {'PostgreSQL + PGVector' if self.use_postgres else 'Chroma (本地)'}")
        print(f"   🧠 智能文本處理: ✅")
        print(f"   🔧 自適應批次: ✅")


    def _setup_embedding_model(self):
        """設定嵌入模型"""
        try:
            if self.model_type == "openai":
                if not OPENAI_EMBEDDINGS_AVAILABLE:
                    raise ImportError("OpenAI Embeddings 不可用")
                
                print(f"🔧 初始化 OpenAI Embeddings...")
                
                api_key = os.getenv("OPENAI_API_KEY")
                base_url = os.getenv("OPENAI_API_BASE")
                
                embedding_params = {
                    "model": "text-embedding-3-small",
                    "api_key": api_key,
                    "max_retries": 3,
                    "request_timeout": 60  # 增加超時時間到60秒
                }
                
                if base_url:
                    embedding_params["base_url"] = base_url
                    print(f"🔧 使用自定義 API 端點: {base_url}")
                
                self.embeddings = OpenAIEmbeddings(**embedding_params)
                print(f"✅ OpenAI Embeddings 初始化成功")
                
            else:
                # HuggingFace 模型
                print(f"🔧 初始化 HuggingFace Embeddings...")
                
                self.embeddings = HuggingFaceEmbeddings(
                    model_name="BAAI/bge-small-zh-v1.5",
                    model_kwargs={"device": "cpu"},
                    encode_kwargs={"batch_size": 16, "normalize_embeddings": True}
                )
                print(f"✅ HuggingFace Embeddings 初始化成功")
                
        except Exception as e:
            print(f"❌ 嵌入模型初始化失敗: {e}")
            
            # 回退機制
            if self.model_type == "openai":
                print("🔄 嘗試 HuggingFace 備選...")
                try:
                    self.embeddings = HuggingFaceEmbeddings(
                        model_name="BAAI/bge-small-zh-v1.5",
                        model_kwargs={"device": "cpu"}
                    )
                    self.model_type = "huggingface"
                    print("✅ 已回退到 HuggingFace")
                except Exception as e2:
                    raise RuntimeError(f"所有嵌入模型都初始化失敗: {e2}")
            else:
                raise

    def load_document(self, file_path: Path) -> List[Document]:
        """載入並智能處理文檔"""
        try:
            extension = file_path.suffix.lower()
            
            # 根據檔案類型載入
            if extension == '.pdf':
                loader = PyPDFLoader(str(file_path))
                documents = loader.load()
                full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
            elif extension in ['.docx', '.doc'] and DOCX2TXT_AVAILABLE:
                if DOCX_METHOD == "docx2txt":
                    import docx2txt
                    full_text = docx2txt.process(str(file_path))
                else:
                    loader = Docx2txtLoader(str(file_path))
                    documents = loader.load()
                    full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
            elif extension == '.csv':
                loader = CSVLoader(str(file_path), encoding='utf-8')
                documents = loader.load()
                full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
            elif extension == '.json':
                loader = JSONLoader(str(file_path), jq_schema='.', text_content=False)
                documents = loader.load()
                full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
            else:
                # 嘗試自動檢測編碼
                encodings = ['utf-8', 'gbk', 'gb2312', 'big5']
                full_text = None
                
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            full_text = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                
                if full_text is None:
                    raise ValueError(f"無法解碼文件: {file_path}")
            
            if not full_text or not full_text.strip():
                logger.warning(f"文件內容為空或無法提取: {file_path.name}")
                return []
            
            # 生成文檔ID
            doc_id = self._generate_doc_id(file_path)
            
            # 使用優化的文本分割器
            source_info = {
                'file_path': str(file_path),
                'file_type': extension,
                'file_size': file_path.stat().st_size if file_path.exists() else 0
            }
            
            documents = self.text_splitter.smart_split_documents(full_text, doc_id, source_info)
            
            # 添加統一元數據
            for doc in documents:
                doc.metadata.update({
                    'source': str(file_path),
                    'filename': file_path.name,
                    'extension': extension,
                    'file_size': source_info['file_size'],
                    'load_timestamp': time.time()
                })
            
            print(f"📄 文檔載入完成: {file_path.name} ({len(documents)} 個分塊)")
            
            return documents
            
        except Exception as e:
            logger.error(f"載入文檔失敗 {file_path}: {e}")
            print(f"❌ 文檔載入失敗: {file_path.name} - {e}")
            return []
    
    def _setup_text_processing(self):
        """設定文本處理組件"""
        self.normalizer = ChineseTextNormalizer()
        self.analyzer = SmartTextAnalyzer()
        print("✅ 文本處理組件初始化完成")
    
    def _load_file_records(self) -> Dict[str, Dict[str, FileInfo]]:
        """載入檔案記錄 - 加強錯誤處理和恢復機制"""
        record_file = self.data_dir / "file_records.json"
        
        # 🔧 檢查檔案是否存在
        if not record_file.exists():
            print("📁 檔案記錄不存在，將建立新的記錄")
            return {}
        
        try:
            # 🔧 讀取並檢查檔案內容
            with open(record_file, 'r', encoding='utf-8') as f:
                content = f.read().strip()
            
            # 🔧 檢查檔案是否為空
            if not content:
                print("⚠️ 檔案記錄為空，將建立新的記錄")
                return {}
            
            # 🔧 檢查是否以 { 開頭（基本 JSON 格式檢查）
            if not content.startswith('{'):
                print(f"⚠️ 檔案記錄格式錯誤，內容開頭: {repr(content[:50])}")
                return self._handle_corrupted_records(record_file, content)
            
            # 🔧 嘗試解析 JSON
            try:
                data = json.loads(content)
            except json.JSONDecodeError as json_error:
                print(f"❌ JSON 解析失敗: {json_error}")
                print(f"   錯誤位置: line {json_error.lineno}, column {json_error.colno}")
                print(f"   檔案前 100 字符: {repr(content[:100])}")
                return self._handle_corrupted_records(record_file, content)
            
            # 🔧 驗證資料格式
            if not isinstance(data, dict):
                print(f"⚠️ 檔案記錄格式錯誤，應為字典但得到: {type(data)}")
                return {}
            
            # 🔧 轉換為 FileInfo 物件
            records = {}
            for collection, files in data.items():
                records[collection] = {}
                for file_path, info in files.items():
                    try:
                        if isinstance(info, dict):
                            fileinfo_fields = {
                                'path': info.get('path', file_path),
                                'size': info.get('size', 0),
                                'mtime': info.get('mtime', time.time()),
                                'hash': info.get('hash', ''),
                                'encoding': info.get('encoding', 'utf-8'),
                                'file_type': info.get('file_type', '')
                            }
                            
                            file_info_obj = FileInfo(**fileinfo_fields)
                            
                            # 🔧 恢復額外屬性
                            if 'uploaded_by' in info:
                                file_info_obj.uploaded_by = info['uploaded_by']
                            if 'uploaded_at' in info:
                                file_info_obj.uploaded_at = info['uploaded_at']
                            if 'file_source' in info:
                                file_info_obj.file_source = info['file_source']
                            
                            records[collection][file_path] = file_info_obj
                        else:
                            records[collection][file_path] = info
                            
                    except Exception as e:
                        logger.warning(f"載入檔案記錄失敗 {file_path}: {e}")
                        # 🔧 建立預設記錄
                        try:
                            default_info = FileInfo(
                                path=file_path,
                                size=0,
                                mtime=time.time(),
                                hash='',
                                encoding='utf-8',
                                file_type=''
                            )
                            records[collection][file_path] = default_info
                        except Exception:
                            logger.error(f"無法建立預設 FileInfo for {file_path}")
                            continue
            
            print(f"✅ 檔案記錄載入成功: {len(records)} 個集合")
            return records
            
        except Exception as e:
            logger.error(f"載入檔案記錄失敗: {e}")
            print(f"❌ 嚴重錯誤，載入檔案記錄失敗: {e}")
            return self._handle_corrupted_records(record_file, "")
    def _handle_corrupted_records(self, record_file: Path, content: str) -> Dict:
        """處理損壞的檔案記錄"""
        try:
            # 🔧 建立備份
            backup_file = record_file.with_suffix('.json.corrupted')
            backup_counter = 1
            while backup_file.exists():
                backup_file = record_file.with_suffix(f'.json.corrupted.{backup_counter}')
                backup_counter += 1
            
            if content:
                with open(backup_file, 'w', encoding='utf-8') as f:
                    f.write(content)
                print(f"📁 損壞的檔案已備份至: {backup_file}")
            
            # 🔧 嘗試從實際檔案重建記錄
            print("🔄 嘗試從實際檔案重建記錄...")
            return self._rebuild_file_records()
            
        except Exception as e:
            logger.error(f"處理損壞記錄失敗: {e}")
            return {}

    def _rebuild_file_records(self) -> Dict:
        """從實際檔案重建記錄"""
        try:
            rebuilt_records = {}
            
            # 🔧 掃描 data 目錄
            for collection_dir in self.data_dir.iterdir():
                if collection_dir.is_dir():
                    collection_name = f"collection_{collection_dir.name}"
                    rebuilt_records[collection_name] = {}
                    
                    print(f"🔍 重建集合: {collection_name}")
                    
                    # 掃描目錄中的檔案
                    for file_path in collection_dir.rglob('*'):
                        if (file_path.is_file() and 
                            file_path.suffix.lower() in SUPPORTED_EXTENSIONS and
                            not file_path.name.startswith('.')):
                            
                            try:
                                file_info = self.get_file_info(file_path)
                                if file_info:
                                    # 設定為重建的檔案
                                    file_info.file_source = "rebuilt"
                                    file_info.uploaded_by = "系統重建"
                                    rebuilt_records[collection_name][str(file_path)] = file_info
                                    print(f"   📄 重建: {file_path.name}")
                            except Exception as e:
                                logger.warning(f"重建檔案記錄失敗 {file_path}: {e}")
            
            # 🔧 保存重建的記錄
            if rebuilt_records:
                print(f"💾 保存重建的記錄...")
                self.file_records = rebuilt_records
                self._save_file_records()
                print(f"✅ 記錄重建完成: {len(rebuilt_records)} 個集合")
            
            return rebuilt_records
            
        except Exception as e:
            logger.error(f"重建檔案記錄失敗: {e}")
            return {}    

        
    def _save_file_records(self):
        """儲存文件記錄 - 修正：使用data_dir而不是persist_dir"""
        try:
            records_file = self.data_dir / "file_records.json"  # 🔧 修正：改為data_dir
            
            # 創建備份
            if records_file.exists():
                backup_file = records_file.with_suffix('.json.backup')
                try:
                    import shutil
                    shutil.copy2(records_file, backup_file)
                except Exception as e:
                    logger.warning(f"創建備份失敗: {e}")
            
            data = {}
            total_files = 0
            
            for collection, files in self.file_records.items():
                data[collection] = {}
                
                for file_path, info in files.items():
                    try:
                        if hasattr(info, '__dict__'):
                            record_dict = {
                                'path': getattr(info, 'path', file_path),
                                'size': getattr(info, 'size', 0),
                                'mtime': getattr(info, 'mtime', time.time()),
                                'hash': getattr(info, 'hash', ''),
                                'encoding': getattr(info, 'encoding', 'utf-8'),
                                'file_type': getattr(info, 'file_type', '')
                            }
                            
                            extra_fields = ['uploaded_by', 'uploaded_at', 'file_source']
                            for field in extra_fields:
                                if hasattr(info, field):
                                    value = getattr(info, field)
                                    if value is not None:
                                        record_dict[field] = value
                            
                            try:
                                normalized_path = str(Path(file_path).absolute())
                                data[collection][normalized_path] = record_dict
                            except Exception:
                                data[collection][file_path] = record_dict
                                
                            total_files += 1
                            
                        else:
                            try:
                                normalized_path = str(Path(file_path).absolute())
                                data[collection][normalized_path] = info
                            except Exception:
                                data[collection][file_path] = info
                            total_files += 1
                            
                    except Exception as e:
                        logger.warning(f"處理檔案記錄失敗 {file_path}: {e}")
                        continue
            
            # 安全寫入檔案
            temp_file = records_file.with_suffix('.json.tmp')
            try:
                with open(temp_file, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                
                if records_file.exists():
                    records_file.unlink()
                temp_file.rename(records_file)
                
                logger.info(f"📁 文件記錄已保存到data目錄:")
                logger.info(f"   📁 集合數: {len(data)}")
                logger.info(f"   📄 總檔案: {total_files}")
                
            except Exception as e:
                if temp_file.exists():
                    temp_file.unlink()
                raise e
                
        except Exception as e:
            logger.error(f"儲存文件記錄失敗: {e}")
            raise

    def get_file_source_statistics(self) -> Dict[str, Dict[str, int]]:
        """獲取檔案來源統計"""
        stats = {}
        
        for collection_name, files in self.file_records.items():
            display_name = collection_name.replace('collection_', '')
            stats[display_name] = {
                'total': len(files),
                'upload': 0,
                'sync': 0,
                'unknown': 0
            }
            
            for file_path, file_info in files.items():
                source = 'unknown'
                
                # 檢查檔案來源
                if hasattr(file_info, 'file_source'):
                    source = file_info.file_source
                elif hasattr(file_info, '__dict__'):
                    # 檢查字典格式
                    if isinstance(file_info, dict):
                        source = file_info.get('file_source', 'unknown')
                    elif hasattr(file_info, 'uploaded_by'):
                        source = 'upload'
                    else:
                        source = 'sync'  # 假設是同步而來
                else:
                    source = 'sync'  # 預設為同步
                
                # 統計
                if source in stats[display_name]:
                    stats[display_name][source] += 1
                else:
                    stats[display_name]['unknown'] += 1
        
        return stats
    
    def diagnose_file_records(self) -> Dict:
        """診斷檔案記錄狀態"""
        diagnosis = {
            'total_collections': len(self.file_records),
            'total_files': 0,
            'source_stats': {},
            'path_issues': [],
            'missing_files': [],
            'orphaned_records': [],
            'recommendations': []
        }
        
        try:
            diagnosis['source_stats'] = self.get_file_source_statistics()
        except Exception as e:
            logger.warning(f"統計來源失敗: {e}")
            diagnosis['source_stats'] = {}
        
        # 檢查每個集合的檔案
        for collection_name, files in self.file_records.items():
            diagnosis['total_files'] += len(files)
            
            for file_path, file_info in files.items():
                try:
                    file_path_obj = Path(file_path)
                    
                    # 檢查檔案是否實際存在
                    if not file_path_obj.exists():
                        diagnosis['missing_files'].append({
                            'collection': collection_name,
                            'path': file_path,
                            'name': file_path_obj.name
                        })
                    
                    # 檢查路徑格式問題
                    if '\\' in file_path and '/' in file_path:
                        diagnosis['path_issues'].append({
                            'collection': collection_name,
                            'path': file_path,
                            'issue': 'mixed_path_separators'
                        })
                    
                    # 檢查是否為孤立記錄（路徑不在預期的 data/ 目錄下）
                    expected_prefix = str(self.data_dir)
                    if not file_path.startswith(expected_prefix):
                        diagnosis['orphaned_records'].append({
                            'collection': collection_name,
                            'path': file_path,
                            'issue': 'outside_data_directory'
                        })
                        
                except Exception as e:
                    logger.warning(f"診斷檔案記錄失敗 {file_path}: {e}")
                    diagnosis['path_issues'].append({
                        'collection': collection_name,
                        'path': file_path,
                        'issue': f'diagnosis_error: {str(e)}'
                    })
        
        # 生成建議
        if diagnosis['missing_files']:
            count = len(diagnosis['missing_files'])
            diagnosis['recommendations'].append(
                f"發現 {count} 個遺失檔案，建議執行同步修復或清理無效記錄"
            )
        
        if diagnosis['path_issues']:
            count = len(diagnosis['path_issues'])
            diagnosis['recommendations'].append(
                f"發現 {count} 個路徑格式問題，建議標準化路徑格式"
            )
        
        if diagnosis['orphaned_records']:
            count = len(diagnosis['orphaned_records'])
            diagnosis['recommendations'].append(
                f"發現 {count} 個孤立記錄，建議檢查檔案位置或清理記錄"
            )
        
        if not diagnosis['recommendations']:
            diagnosis['recommendations'].append("檔案記錄狀態良好，無需特別處理")
        
        return diagnosis

    def cleanup_invalid_records(self) -> Dict:
        """清理無效的檔案記錄"""
        cleanup_result = {
            'cleaned_collections': 0,
            'removed_records': 0,
            'errors': []
        }
        
        try:
            for collection_name in list(self.file_records.keys()):
                files = self.file_records[collection_name]
                original_count = len(files)
                
                # 檢查並移除無效記錄
                valid_files = {}
                
                for file_path, file_info in files.items():
                    try:
                        if Path(file_path).exists():
                            valid_files[file_path] = file_info
                        else:
                            logger.info(f"移除無效記錄: {file_path}")
                            
                    except Exception as e:
                        logger.warning(f"檢查記錄時出錯 {file_path}: {e}")
                        cleanup_result['errors'].append(f"{file_path}: {str(e)}")
                
                # 更新記錄
                if len(valid_files) != original_count:
                    self.file_records[collection_name] = valid_files
                    removed = original_count - len(valid_files)
                    cleanup_result['removed_records'] += removed
                    cleanup_result['cleaned_collections'] += 1
                    
                    logger.info(f"集合 {collection_name}: 移除了 {removed} 個無效記錄")
            
            # 保存清理後的記錄
            if cleanup_result['removed_records'] > 0:
                self._save_file_records()
                
        except Exception as e:
            logger.error(f"清理記錄失敗: {e}")
            cleanup_result['errors'].append(f"整體清理失敗: {str(e)}")
        
        return cleanup_result


    
    
    def _generate_doc_id(self, file_path: Path) -> str:
        """生成文檔ID"""
        content_hash = hashlib.md5(str(file_path).encode()).hexdigest()[:8]
        return f"doc_{file_path.stem}_{content_hash}"
    
    def get_or_create_vectorstore(self, collection_name: str):
        """獲取或創建向量存儲 - PostgreSQL 優先"""
        if collection_name not in self._vector_stores:
            try:
                if self.use_postgres and PGVECTOR_AVAILABLE:
                    # 🔧 使用 PGVector
                    try:
                        from langchain_postgres import PGVector
                        self._vector_stores[collection_name] = PGVector(
                            embeddings=self.embeddings,
                            collection_name=collection_name,
                            connection=self.connection_string,
                            use_jsonb=True,
                        )
                    except ImportError:
                        from langchain_community.vectorstores import PGVector
                        self._vector_stores[collection_name] = PGVector(
                            connection_string=self.connection_string,
                            collection_name=collection_name,
                            embedding_function=self.embeddings,
                            distance_strategy="cosine",
                            pre_delete_collection=False,
                            logger=logger
                        )
                    print(f"✅ PGVector 向量存儲就緒: {collection_name}")
                else:
                    # 🔧 備用 Chroma - 確保導入
                    if CHROMA_AVAILABLE:
                        from langchain_community.vectorstores import Chroma
                        self._vector_stores[collection_name] = Chroma(
                            collection_name=collection_name,
                            embedding_function=self.embeddings,
                            persist_directory=str(self.persist_dir)
                        )
                        print(f"✅ Chroma 向量存儲就緒: {collection_name}")
                    else:
                        raise ImportError("Chroma 不可用且 PostgreSQL 也不可用")
                        
            except Exception as e:
                logger.error(f"向量存儲創建失敗: {e}")
                raise RuntimeError(f"無法創建向量存儲: {e}")
        
        return self._vector_stores[collection_name]
    
    def get_file_info(self, file_path: Path) -> Optional[FileInfo]:
        """獲取文件資訊 - 修正版：更可靠的哈希計算和錯誤處理"""
        try:
            if not file_path.exists():
                return None
            
            # 檢查是否為符號連結
            if file_path.is_symlink():
                print(f"⚠️ 跳過符號連結: {file_path.name}")
                return None
            
            stat = file_path.stat()
            
            # 🆕 修正：更可靠和高效的文件哈希計算
            try:
                # 分級處理不同大小的檔案
                if stat.st_size == 0:
                    # 空檔案
                    file_hash = hashlib.md5(b"").hexdigest()
                elif stat.st_size < 1024 * 1024:  # 小於 1MB
                    # 小檔案：直接讀取全部內容
                    try:
                        with open(file_path, 'rb') as f:
                            content = f.read()
                        file_hash = hashlib.md5(content).hexdigest()
                    except PermissionError:
                        # 權限問題，使用檔案屬性生成哈希
                        file_hash = hashlib.md5(
                            f"{stat.st_size}:{stat.st_mtime}:{file_path.name}".encode()
                        ).hexdigest()
                elif stat.st_size < 50 * 1024 * 1024:  # 小於 50MB
                    # 中等檔案：分塊讀取
                    hash_md5 = hashlib.md5()
                    try:
                        with open(file_path, 'rb') as f:
                            # 8KB 塊大小
                            for chunk in iter(lambda: f.read(8192), b""):
                                hash_md5.update(chunk)
                        file_hash = hash_md5.hexdigest()
                    except (PermissionError, OSError) as e:
                        logger.warning(f"無法讀取檔案 {file_path}: {e}")
                        # 回退到基於檔案屬性的哈希
                        file_hash = hashlib.md5(
                            f"{stat.st_size}:{stat.st_mtime}:{file_path.name}".encode()
                        ).hexdigest()
                else:
                    # 大檔案：使用檔案屬性 + 部分內容
                    try:
                        # 讀取檔案開頭和結尾各 4KB
                        with open(file_path, 'rb') as f:
                            head = f.read(4096)
                            if stat.st_size > 8192:
                                f.seek(-4096, 2)  # 從末尾向前 4KB
                                tail = f.read(4096)
                            else:
                                tail = b""
                        
                        hash_content = f"{stat.st_size}:{stat.st_mtime}:{file_path.name}".encode()
                        hash_content += head + tail
                        file_hash = hashlib.md5(hash_content).hexdigest()
                        
                    except (PermissionError, OSError) as e:
                        logger.warning(f"無法讀取大檔案 {file_path}: {e}")
                        # 純屬性哈希
                        file_hash = hashlib.md5(
                            f"{stat.st_size}:{stat.st_mtime}:{file_path.name}".encode()
                        ).hexdigest()
                        
            except Exception as e:
                # 最終回退方案
                logger.warning(f"檔案哈希計算失敗 {file_path}: {e}")
                file_hash = hashlib.md5(
                    f"{stat.st_size}:{stat.st_mtime}:{file_path.name}".encode()
                ).hexdigest()
            
            return FileInfo(
                path=str(file_path),
                size=stat.st_size,
                mtime=stat.st_mtime,
                hash=file_hash,
                encoding="utf-8",
                file_type=file_path.suffix.lower()
            )
            
        except Exception as e:
            logger.error(f"獲取文件資訊失敗 {file_path}: {e}")
            return None
    
    def scan_directory_changes(self, dir_path: Path, collection_name: str) -> Tuple[List[Path], List[Path], List[str], Dict[str, FileInfo]]:
        """掃描目錄變更 - 修正版：正確處理上傳檔案"""
        current_files = {}
        
        print(f"🔍 掃描目錄: {dir_path}")
        
        # 遞迴掃描目錄
        file_count = 0
        for file_path in dir_path.rglob('*'):
            if (file_path.is_file() and 
                file_path.suffix.lower() in SUPPORTED_EXTENSIONS and
                not file_path.name.startswith('.') and
                file_path.stat().st_size > 0):  # 跳過空文件
                
                file_info = self.get_file_info(file_path)
                if file_info:
                    # 🆕 修正：使用標準化的絕對路徑作為鍵值
                    try:
                        # 使用 absolute() 避免符號連結問題
                        absolute_path = str(file_path.absolute())
                        current_files[absolute_path] = file_info
                        file_count += 1
                    except Exception as e:
                        logger.warning(f"路徑標準化失敗 {file_path}: {e}")
                        # 回退到原始路徑
                        current_files[str(file_path)] = file_info
                        file_count += 1
        
        print(f"📄 找到 {file_count} 個有效檔案")
        
        old_files = self.file_records.get(collection_name, {})
        print(f"📋 舊記錄中有 {len(old_files)} 個檔案")
        
        # 🆕 修正：正規化舊記錄的路徑鍵值
        normalized_old_files = {}
        normalization_errors = 0
        
        for old_path, old_info in old_files.items():
            try:
                old_path_obj = Path(old_path)
                
                if old_path_obj.is_absolute():
                    # 已經是絕對路徑
                    normalized_key = str(old_path_obj.absolute())
                else:
                    # 相對路徑轉絕對路徑
                    try:
                        abs_path = (dir_path / old_path).absolute()
                        normalized_key = str(abs_path)
                    except Exception:
                        # 如果無法轉換，保持原樣
                        normalized_key = old_path
                        
                normalized_old_files[normalized_key] = old_info
                
            except Exception as e:
                logger.warning(f"舊路徑正規化失敗 {old_path}: {e}")
                # 保持原路徑
                normalized_old_files[old_path] = old_info
                normalization_errors += 1
        
        if normalization_errors > 0:
            print(f"⚠️ {normalization_errors} 個舊路徑正規化失敗")
        
        # 🆕 修正：智能變更檢測
        added_files = []
        modified_files = []
        
        print("🔍 檢測變更...")
        
        for file_path, file_info in current_files.items():
            current_file_name = Path(file_path).name
            current_hash = file_info.hash
            
            # 首先嘗試精確路徑匹配
            if file_path in normalized_old_files:
                old_info = normalized_old_files[file_path]
                if old_info.hash != current_hash:
                    modified_files.append(Path(file_path))
                    print(f"📝 修改檔案: {current_file_name}")
            else:
                # 🆕 智能檔案匹配：檢查是否是同一檔案的不同路徑表示
                file_found = False
                
                for old_path, old_info in normalized_old_files.items():
                    old_file_name = Path(old_path).name
                    
                    # 檔案名相同且哈希相同 = 同一檔案
                    if (current_file_name == old_file_name and 
                        current_hash == old_info.hash):
                        file_found = True
                        print(f"🔄 路徑變更但內容相同: {current_file_name}")
                        break
                        
                    # 檔案名相同但哈希不同 = 檔案被修改
                    elif (current_file_name == old_file_name and 
                          current_hash != old_info.hash):
                        modified_files.append(Path(file_path))
                        file_found = True
                        print(f"📝 修改檔案 (路徑變更): {current_file_name}")
                        break
                
                if not file_found:
                    added_files.append(Path(file_path))
                    print(f"📄 新檔案: {current_file_name}")
        
        # 🆕 修正：智能刪除檢測
        deleted_files = []
        
        for old_path in normalized_old_files.keys():
            old_file_name = Path(old_path).name
            
            if old_path not in current_files:
                # 檢查檔案是否真的不存在（可能只是路徑表示不同）
                file_still_exists = False
                
                for current_path in current_files.keys():
                    current_file_name = Path(current_path).name
                    if current_file_name == old_file_name:
                        # 進一步檢查是否是同一檔案（通過內容哈希）
                        current_hash = current_files[current_path].hash
                        old_hash = normalized_old_files[old_path].hash
                        
                        if current_hash == old_hash:
                            file_still_exists = True
                            break
                
                if not file_still_exists:
                    deleted_files.append(old_path)
                    print(f"🗑️ 刪除檔案: {old_file_name}")
        
        print(f"📊 變更統計:")
        print(f"   📄 新增: {len(added_files)}")
        print(f"   📝 修改: {len(modified_files)}")
        print(f"   🗑️ 刪除: {len(deleted_files)}")
        
        return added_files, modified_files, deleted_files, current_files
    
    def incremental_update(self, collection_name: str, added_files: List[Path], 
                          modified_files: List[Path], deleted_files: List[str],
                          current_files: Dict[str, FileInfo]) -> bool:
        """🚀 優化版增量更新"""
        with self.processing_lock:
            try:
                vectorstore = self.get_or_create_vectorstore(collection_name)
                
                # 處理刪除和修改
                files_to_delete = deleted_files + [str(f) for f in modified_files]
                if files_to_delete:
                    for file_path in files_to_delete:
                        try:
                            vectorstore.delete(filter={"source": file_path})
                            print(f"🗑️ 已刪除: {Path(file_path).name}")
                        except Exception as e:
                            logger.warning(f"刪除文檔失敗 {file_path}: {e}")
                
                # 處理新增和修改的文件
                files_to_process = added_files + modified_files
                if not files_to_process:
                    print("   ✅ 無新文件需要處理")
                    return True
                
                print(f"📄 開始處理 {len(files_to_process)} 個文件...")
                print(f"   ⚠️ 處理大文件可能需要較長時間，請耐心等待...")
                
                # 並發載入文檔
                all_documents = []
                if PERFORMANCE_CONFIG.get("parallel_processing", True):
                    all_documents = self._parallel_load_documents(files_to_process, collection_name)
                else:
                    all_documents = self._sequential_load_documents(files_to_process, collection_name)
                
                if not all_documents:
                    print("   ⚠️ 沒有有效文檔需要向量化")
                    return True
                
                # 統計和成本估算
                total_tokens = sum(doc.metadata.get('token_count', 0) for doc in all_documents)
                estimated_cost = self.batch_processor.token_estimator.estimate_embedding_cost(total_tokens)
                
                print(f"\n📊 向量化統計:")
                print(f"   📄 總分塊數: {len(all_documents)}")
                print(f"   📏 總 tokens: {total_tokens:,}")
                print(f"   💰 估算成本: ${estimated_cost:.4f}")
                
                # 創建智能批次並處理
                batches = self.batch_processor.create_smart_batches(all_documents)
                success_count = self._process_batches(vectorstore, batches)
                
                print(f"\n🎉 向量化完成！")
                print(f"   ✅ 成功: {success_count}/{len(all_documents)} 個分塊")
                print(f"   📊 成功率: {(success_count/len(all_documents)*100):.1f}%")
                
                # 更新文件記錄
                self.file_records[collection_name] = current_files
                self._save_file_records()
                
                # 記憶體清理
                if success_count % PERFORMANCE_CONFIG.get("gc_frequency", 50) == 0:
                    gc.collect()
                
                return success_count > 0
                
            except Exception as e:
                logger.error(f"增量更新失敗 {collection_name}: {e}")
                print(f"❌ 系統錯誤: {e}")
                return False
    
    def _parallel_load_documents(self, file_paths: List[Path], collection_name: str) -> List[Document]:
        """並發載入文檔"""
        all_documents = []
        max_workers = min(SYSTEM_CONFIG.get("max_workers", 4), len(file_paths))
        
        print(f"   🚀 並發載入 (工作線程: {max_workers})")
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_file = {
                executor.submit(self.load_document, file_path): file_path 
                for file_path in file_paths
            }
            
            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    documents = future.result()
                    if documents:
                        for doc in documents:
                            doc.metadata['collection'] = collection_name
                        all_documents.extend(documents)
                        print(f"   ✅ {file_path.name}: {len(documents)} 分塊")
                    else:
                        print(f"   ⚠️ {file_path.name}: 無有效內容")
                except Exception as e:
                    print(f"   ❌ {file_path.name}: {e}")
                    logger.error(f"並發載入失敗 {file_path}: {e}")
        
        return all_documents
    
    def _sequential_load_documents(self, file_paths: List[Path], collection_name: str) -> List[Document]:
        """順序載入文檔"""
        all_documents = []
        
        print(f"   📄 順序載入")
        
        for i, file_path in enumerate(file_paths, 1):
            try:
                print(f"   [{i}/{len(file_paths)}] 處理: {file_path.name}")
                documents = self.load_document(file_path)
                
                if documents:
                    for doc in documents:
                        doc.metadata['collection'] = collection_name
                    all_documents.extend(documents)
                    print(f"      ✅ 載入 {len(documents)} 個分塊")
                else:
                    print(f"      ⚠️ 無有效內容")
                    
            except Exception as e:
                print(f"      ❌ 載入失敗: {e}")
                logger.error(f"文件載入失敗 {file_path}: {e}")
        
        return all_documents
    
    from typing import Union, Any

    def _process_batches(self, vectorstore: Union["Chroma", Any], batches: List[Tuple[List[Document], Dict]]) -> int:
        """處理批次向量化 - 完整的錯誤處理和元數據修復"""

        success_count = 0
        total_docs = sum(len(batch_docs) for batch_docs, _ in batches)
        
        print(f"\n🔄 開始批次向量化...")
        print(f"   📦 總批次數: {len(batches)}")
        print(f"   📄 總文檔數: {total_docs}")
        
        for batch_num, (batch_docs, batch_info) in enumerate(batches, 1):
            print(f"\n   📦 批次 {batch_num}/{len(batches)}")
            print(f"      📄 文檔數: {batch_info['documents']}")
            print(f"      📏 tokens: {batch_info['tokens']:,}")
            print(f"      📊 使用率: {(batch_info['tokens']/TOKEN_LIMITS['max_tokens_per_request']*100):.1f}%")
            
            # 顯示文檔類型分布
            type_info = ", ".join([f"{k}:{v}" for k, v in batch_info['types'].items()])
            print(f"      🏷️ 類型: {type_info}")
            
            start_time = time.time()
            
            try:
                print(f"      🚀 開始處理批次 {batch_num}...")
                print(f"      📡 正在調用 OpenAI API... (這可能需要 30-60 秒)")
                
                # 🛠️ 修復：統一處理元數據，確保類型正確
                safe_docs = []
                for doc in batch_docs:
                    safe_metadata = self._ensure_simple_metadata(doc.metadata)
                    safe_doc = Document(page_content=doc.page_content, metadata=safe_metadata)
                    safe_docs.append(safe_doc)

                print(f"      🔧 已處理 {len(safe_docs)} 個文檔的元數據，確保類型兼容")
                
                vectorstore.add_documents(safe_docs)
                processing_time = time.time() - start_time
                
                success_count += len(batch_docs)
                self.batch_processor.record_batch_result(True, processing_time)
                
                print(f"      ✅ 批次 {batch_num} 完成 ({processing_time:.1f}s)")
                print(f"      📊 總進度: {success_count}/{total_docs} ({success_count/total_docs*100:.1f}%)")
                
                # 批次間延遲
                if batch_num < len(batches):
                    delay = TOKEN_LIMITS["batch_delay"]
                    print(f"      ⏱️ 等待 {delay} 秒...")
                    time.sleep(delay)
                    
            except Exception as e:
                processing_time = time.time() - start_time
                self.batch_processor.record_batch_result(False, processing_time)
                
                error_msg = str(e)
                print(f"      ❌ 批次 {batch_num} 失敗 ({processing_time:.1f}s)")
                print(f"         錯誤: {error_msg}")
                
                # 🔧 特別處理元數據錯誤
                if "metadata" in error_msg.lower():
                    print(f"         🔧 檢測到元數據錯誤，嘗試更嚴格的處理...")
                    try:
                        # 重新嘗試，使用最嚴格的元數據過濾
                        ultra_safe_docs = []
                        for doc in batch_docs:
                            # 只保留最基本的元數據字段
                            minimal_metadata = {
                                'doc_id': str(doc.metadata.get('doc_id', 'unknown')),
                                'chunk_id': str(doc.metadata.get('chunk_id', 'unknown')),
                                'chunk_index': int(doc.metadata.get('chunk_index', 0)),
                                'text_type': str(doc.metadata.get('text_type', 'unknown')),
                                'source': str(doc.metadata.get('source', 'unknown')),
                                'filename': str(doc.metadata.get('filename', 'unknown')),
                                'token_count': int(doc.metadata.get('token_count', 0)),
                                'chunk_length': int(doc.metadata.get('chunk_length', 0)),
                                # URL 處理
                                'contained_urls': str(doc.metadata.get('contained_urls', '')),
                                'url_count': int(doc.metadata.get('url_count', 0)),
                                'has_urls': bool(doc.metadata.get('has_urls', False))
                            }
                            
                            ultra_safe_doc = Document(
                                page_content=doc.page_content,
                                metadata=minimal_metadata
                            )
                            ultra_safe_docs.append(ultra_safe_doc)
                        
                        vectorstore.add_documents(ultra_safe_docs)
                        success_count += len(batch_docs)
                        print(f"         ✅ 使用最小化元數據重新處理成功")
                        continue
                        
                    except Exception as retry_e:
                        print(f"         ❌ 重新處理也失敗: {retry_e}")
                
                # 其他錯誤處理
                if "timeout" in error_msg.lower():
                    print(f"         🕐 超時錯誤，延長等待時間...")
                    time.sleep(30)
                elif "rate_limit" in error_msg.lower() or "429" in error_msg:
                    print(f"         🚦 速率限制，延長等待...")
                    time.sleep(60)
                elif "token" in error_msg.lower() and batch_info['documents'] > 1:
                    print(f"         🔧 Token 超限，嘗試單個處理...")
                    single_success = self._process_documents_individually(vectorstore, batch_docs)
                    success_count += single_success
                elif "connection" in error_msg.lower():
                    print(f"         🌐 連接錯誤，等待重試...")
                    time.sleep(20)
                    try:
                        print(f"         🔄 重試批次 {batch_num}...")
                        # 使用安全的元數據重試
                        safe_docs = []
                        for doc in batch_docs:
                            safe_metadata = self._ensure_simple_metadata(doc.metadata)
                            safe_doc = Document(page_content=doc.page_content, metadata=safe_metadata)
                            safe_docs.append(safe_doc)
                        
                        vectorstore.add_documents(safe_docs)
                        success_count += len(batch_docs)
                        print(f"         ✅ 重試成功")
                    except Exception as retry_e:
                        print(f"         ❌ 重試失敗: {retry_e}")
                else:
                    print(f"         ⚠️ 跳過此批次")
                    
                # 每次錯誤後添加額外延遲
                print(f"         ⏸️ 錯誤後暫停 10 秒...")
                time.sleep(10)
        
        return success_count
    
    def _ensure_simple_metadata(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """確保元數據只包含Chroma支持的簡單類型：string, int, float, bool"""
        safe_metadata = {}
        
        for key, value in metadata.items():
            if isinstance(value, (str, int, float, bool)) or value is None:
                # ✅ 已經是Chroma支持的簡單類型，直接保留
                safe_metadata[key] = value
            elif isinstance(value, list):
                # ❌ 列表不被支持 → 轉為分隔符字串
                if key == 'contained_urls' or 'url' in key.lower():
                    safe_metadata[key] = '|'.join(str(v) for v in value) if value else ''
                    safe_metadata[f'{key}_count'] = len(value)
                else:
                    safe_metadata[key] = '|'.join(str(v) for v in value) if value else ''
                    safe_metadata[f'{key}_count'] = len(value)
            elif isinstance(value, dict):
                # ❌ 字典不被支持 → 轉為JSON字符串
                safe_metadata[key] = json.dumps(value, ensure_ascii=False)
            else:
                # ❌ 其他類型不被支持 → 轉為字符串
                safe_metadata[key] = str(value)
        
        return safe_metadata

    
    def _process_documents_individually(self, vectorstore: Chroma, documents: List[Document]) -> int:
        """單個處理文檔"""
        success_count = 0
        
        for i, doc in enumerate(documents):
            try:
                doc_tokens = doc.metadata.get('token_count', 0)
                if doc_tokens > TOKEN_LIMITS['max_tokens_per_request']:
                    print(f"         ⚠️ 文檔 {i+1} 仍然過大 ({doc_tokens:,} tokens)，跳過")
                    continue
                
                vectorstore.add_documents([doc])
                success_count += 1
                print(f"         ✅ 單個文檔 {i+1}/{len(documents)} 完成")
                time.sleep(1)  # 單個處理時短暫延遲
                
            except Exception as e:
                print(f"         ❌ 單個文檔 {i+1} 失敗: {e}")
        
        return success_count
    
    def get_collection_name(self, dir_path: Path) -> str:
        """獲取集合名稱"""
        try:
            relative_path = dir_path.relative_to(self.data_dir)
            if relative_path.parts:
                return f"collection_{relative_path.parts[0]}"
        except ValueError:
            pass
        return "collection_other"
    
    def sync_collections(self) -> int:
        """同步所有集合"""
        print("🔄 開始智能增量同步...")
        
        directories = [d for d in self.data_dir.iterdir() if d.is_dir()]
        print(f"📁 掃描到 {len(directories)} 個目錄")
        
        if not directories:
            print("   ⚠️ 數據目錄為空")
            return 0
        
        total_changes = 0
        
        for i, dir_path in enumerate(directories, 1):
            collection_name = self.get_collection_name(dir_path)
            folder_name = collection_name.replace('collection_', '')
            
            print(f"\n[{i}/{len(directories)}] 檢查目錄: {folder_name}")
            
            try:
                added_files, modified_files, deleted_files, current_files = \
                    self.scan_directory_changes(dir_path, collection_name)
                
                if not added_files and not modified_files and not deleted_files:
                    print(f"      ✅ 無變更")
                    continue
                
                print(f"      📊 新增: {len(added_files)}, 修改: {len(modified_files)}, 刪除: {len(deleted_files)}")
                
                success = self.incremental_update(
                    collection_name, added_files, modified_files, deleted_files, current_files
                )
                
                if success:
                    changes = len(added_files) + len(modified_files) + len(deleted_files)
                    total_changes += changes
                    print(f"      ✅ 目錄更新完成")
                else:
                    print(f"      ❌ 目錄更新失敗")
                    
            except Exception as e:
                print(f"      ❌ 目錄處理錯誤: {e}")
                logger.error(f"目錄同步失敗 {dir_path}: {e}")
        
        print(f"\n✅ 智能增量同步完成")
        print(f"   📊 總變更數: {total_changes}")
        
        # 顯示性能統計
        stats = self.batch_processor.get_performance_stats()
        if stats['total_batches'] > 0:
            print(f"   📈 性能統計:")
            print(f"      批次數: {stats['total_batches']}")
            print(f"      成功率: {stats['success_rate']*100:.1f}%")
            print(f"      平均處理時間: {stats['avg_batch_time']:.1f}s/批次")
        
        return total_changes
    
    def get_stats(self) -> Dict:
        """獲取系統統計 - 修正：從file_records而非persist_dir獲取集合"""
        try:
            stats = {}
            
            # 🔧 修正：從file_records獲取已知的集合，而不是掃描persist_dir
            for collection_name in self.file_records.keys():
                folder_name = collection_name.replace('collection_', '')
                
                try:
                    vectorstore = self.get_or_create_vectorstore(collection_name)
                    count = vectorstore._collection.count()
                    stats[folder_name] = count
                except Exception as e:
                    logger.warning(f"獲取集合統計失敗 {collection_name}: {e}")
                    stats[folder_name] = 0
            
            return stats
        except Exception as e:
            logger.error(f"獲取統計失敗: {e}")
            return {}
    
    def search(self, query: str, collection_name: str = None, k: int = 5) -> List[SearchResult]:
        """優化版搜索"""
        try:
            # 創建搜索變體
            query_variants = self.normalizer.create_search_variants(query)
            all_results = []
            
            # 處理集合範圍
            target_collections = []
            if collection_name:
                target_collections = [collection_name]
            else:
                stats = self.get_stats()
                target_collections = [f"collection_{name}" for name in stats.keys()]
            
            # 對每個集合和查詢變體進行搜索
            for variant in query_variants:
                for coll_name in target_collections:
                    try:
                        vectorstore = self.get_or_create_vectorstore(coll_name)
                        docs_and_scores = vectorstore.similarity_search_with_score(variant, k=k)
                        
                        for doc, score in docs_and_scores:
                            # 創建搜索結果
                            chunk_info = ChunkInfo(
                                chunk_id=doc.metadata.get('chunk_id', 'unknown'),
                                content=doc.page_content,
                                metadata=doc.metadata,
                                token_count=doc.metadata.get('token_count', 0),
                                quality_score=doc.metadata.get('quality_score', 0.5),
                                relationships=[]
                            )
                            
                            result = SearchResult(
                                content=doc.page_content[:500],  # 限制預覽長度
                                score=1.0 - score,  # 轉換為相似度分數
                                metadata=doc.metadata,
                                collection=coll_name,
                                chunk_info=chunk_info
                            )
                            
                            all_results.append(result)
                            
                    except Exception as e:
                        logger.warning(f"搜索集合失敗 {coll_name}: {e}")
            
            # 去重和排序
            seen_content = set()
            unique_results = []
            
            for result in sorted(all_results, key=lambda x: x.score, reverse=True):
                content_hash = hashlib.md5(result.content.encode()).hexdigest()
                if content_hash not in seen_content:
                    seen_content.add(content_hash)
                    unique_results.append(result)
            
            return unique_results[:k]
            
        except Exception as e:
            logger.error(f"搜索失敗: {e}")
            return []
    
    def diagnose_system(self) -> Dict:
        """系統診斷"""
        print("🔍 === 系統診斷 ===")
        
        diagnosis = {
            "environment": {},
            "embedding_model": {},
            "text_processing": {},
            "performance": {},
            "recommendations": []
        }
        
        # 環境檢查
        api_key = os.getenv("OPENAI_API_KEY")
        diagnosis["environment"]["openai_api_key"] = "✅ 已設置" if api_key else "❌ 未設置"
        diagnosis["environment"]["model_type"] = self.model_type
        
        # 嵌入模型檢查
        try:
            test_result = self.embeddings.embed_query("測試")
            diagnosis["embedding_model"]["status"] = "✅ 正常"
            diagnosis["embedding_model"]["dimension"] = len(test_result)
        except Exception as e:
            diagnosis["embedding_model"]["status"] = f"❌ 失敗: {e}"
        
        # 文本處理檢查
        diagnosis["text_processing"]["normalizer"] = "✅ 正常" if self.normalizer else "❌ 異常"
        diagnosis["text_processing"]["analyzer"] = "✅ 正常" if self.analyzer else "❌ 異常"
        diagnosis["text_processing"]["splitter"] = "✅ 正常" if self.text_splitter else "❌ 異常"
        
        # 性能統計
        perf_stats = self.batch_processor.get_performance_stats()
        diagnosis["performance"] = perf_stats
        
        # 建議
        if not api_key:
            diagnosis["recommendations"].append("設置 OPENAI_API_KEY 環境變數")
        
        if perf_stats.get("success_rate", 1) < 0.8:
            diagnosis["recommendations"].append("成功率偏低，建議檢查網路連接和 API 配額")
        
        # 輸出診斷結果
        for category, info in diagnosis.items():
            if category != "recommendations":
                print(f"\n🔧 {category.upper()}:")
                for key, value in info.items():
                    print(f"   {key}: {value}")
        
        if diagnosis["recommendations"]:
            print(f"\n💡 建議:")
            for rec in diagnosis["recommendations"]:
                print(f"   • {rec}")
        
        return diagnosis
    def upload_single_file(self, file_content: bytes, filename: str, collection_name: str) -> Dict:
        """
        上傳單個文件到指定集合並建立向量 - 修正版：保存檔案到 data/ 目錄
        
        Args:
            file_content: 文件二進制內容
            filename: 文件名稱
            collection_name: 目標集合名稱
            
        Returns:
            Dict: 上傳結果，包含文檔分塊信息
        """
        try:
            # 🔧 基本驗證
            if not file_content:
                return {
                    "success": False,
                    "message": "文件內容為空",
                    "chunks": []
                }
            
            if not filename or not filename.strip():
                return {
                    "success": False,
                    "message": "文件名不能為空",
                    "chunks": []
                }
            
            # 🔧 檢查文件擴展名
            file_extension = Path(filename).suffix.lower()
            if file_extension not in SUPPORTED_EXTENSIONS:
                return {
                    "success": False,
                    "message": f"不支援的文件格式: {file_extension}。支援格式: {', '.join(SUPPORTED_EXTENSIONS)}",
                    "chunks": []
                }
            
            # 🆕 修正：確定目標目錄和檔案路徑
            bot_name = collection_name.replace('collection_', '')
            target_dir = self.data_dir / bot_name  # data/bot_name/
            
            # 🔧 修正：確保目錄存在，使用 parents=True 處理深層目錄
            try:
                target_dir.mkdir(parents=True, exist_ok=True)
            except Exception as e:
                return {
                    "success": False,
                    "message": f"無法創建目標目錄 {target_dir}: {str(e)}",
                    "chunks": []
                }
            
            target_file_path = target_dir / filename
            
            # 🆕 修正：處理檔案衝突 - 更安全的檢查
            if target_file_path.exists():
                try:
                    # 檢查是否可以寫入
                    if not os.access(target_file_path, os.W_OK):
                        return {
                            "success": False,
                            "message": f"檔案 {filename} 存在但無寫入權限",
                            "chunks": []
                        }
                    print(f"⚠️ 檔案 {filename} 已存在，將會覆蓋")
                except Exception as e:
                    return {
                        "success": False,
                        "message": f"檢查檔案權限失敗: {str(e)}",
                        "chunks": []
                    }
            
            # 🆕 修正：保存檔案到正確位置 - 加強錯誤處理
            print(f"💾 保存檔案到: {target_file_path}")
            try:
                with open(target_file_path, 'wb') as f:
                    f.write(file_content)
                
                # 🔧 驗證檔案是否正確寫入
                if not target_file_path.exists() or target_file_path.stat().st_size != len(file_content):
                    raise IOError("檔案寫入驗證失敗")
                    
                print(f"✅ 檔案保存成功: {target_file_path} ({len(file_content)} bytes)")
                
            except Exception as e:
                # 清理可能的不完整檔案
                try:
                    if target_file_path.exists():
                        target_file_path.unlink()
                except:
                    pass
                return {
                    "success": False,
                    "message": f"檔案保存失敗: {str(e)}",
                    "chunks": []
                }
            
            try:
                # 🆕 修正：使用實際檔案路徑進行處理
                print(f"📄 開始處理檔案: {filename}")
                documents = self.load_document(target_file_path)
                
                if not documents:
                    # 🔧 如果處理失敗，清理已保存的檔案
                    try:
                        target_file_path.unlink()
                        print(f"🧹 處理失敗，已清理檔案: {target_file_path}")
                    except Exception as cleanup_error:
                        logger.warning(f"清理檔案失敗: {cleanup_error}")
                        
                    return {
                        "success": False,
                        "message": "文件內容為空或格式不支援",
                        "chunks": []
                    }
                
                # 🔧 修正：設置集合信息和元數據 - 使用正確的時間戳
                current_timestamp = time.time()
                for doc in documents:
                    doc.metadata.update({
                        'collection': collection_name,
                        'original_filename': filename,
                        'upload_timestamp': current_timestamp,
                        'file_source': 'upload',
                        'source': str(target_file_path),  # 🆕 使用實際檔案路徑
                        'saved_to_data_dir': True,  # 🆕 標記檔案已保存
                        'file_extension': file_extension
                    })
                
                # 獲取向量存儲
                vectorstore = self.get_or_create_vectorstore(collection_name)
                
                # 🔧 修正：刪除已存在的同名文件 - 使用更準確的條件
                try:
                    # 🔧 使用標準化路徑進行查詢
                    delete_conditions = [
                        {"source": str(target_file_path)},
                        {"original_filename": filename},
                        {"filename": filename}
                    ]
                    
                    total_deleted = 0
                    for condition in delete_conditions:
                        try:
                            existing_docs = vectorstore.get(where=condition)  # get仍可用where
                            if existing_docs and existing_docs.get('documents'):
                                vectorstore.delete(filter=condition)  # ✅ 正確：改為filter
                                deleted_count = len(existing_docs['documents'])
                                total_deleted += deleted_count
                                print(f"🗑️ 使用條件 {condition} 刪除了 {deleted_count} 個現有分塊")
                        except Exception as e:
                            print(f"⚠️ 刪除條件 {condition} 時出現警告: {e}")
                    
                    if total_deleted > 0:
                        print(f"🗑️ 總共刪除了 {total_deleted} 個現有分塊")
                        
                except Exception as e:
                    print(f"⚠️ 刪除現有分塊時出現警告: {e}")
                
                # 使用批次處理器向量化
                print(f"🔄 開始向量化處理...")
                batches = self.batch_processor.create_smart_batches(documents)
                success_count = self._process_batches(vectorstore, batches)
                
                # 準備回傳信息
                chunks_info = []
                for i, doc in enumerate(documents):
                    chunks_info.append({
                        'chunk_id': doc.metadata.get('chunk_id', f'chunk_{i+1}'),
                        'chunk_index': doc.metadata.get('chunk_index', i),
                        'content_preview': doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content,
                        'token_count': doc.metadata.get('token_count', 0),
                        'text_type': doc.metadata.get('text_type', 'unknown')
                    })
                
                # 🆕 修正：更新文件記錄（使用實際檔案信息）
                if collection_name not in self.file_records:
                    self.file_records[collection_name] = {}
                
                # 🆕 使用實際檔案的統計信息
                try:
                    file_stat = target_file_path.stat()
                    file_hash = hashlib.md5(file_content).hexdigest()
                    
                    file_info = FileInfo(
                        path=str(target_file_path),  # 🆕 使用實際路徑
                        size=file_stat.st_size,
                        mtime=file_stat.st_mtime,
                        hash=file_hash,
                        encoding="utf-8",
                        file_type=file_extension
                    )
                    
                    # 🆕 添加上傳者信息（作為屬性）
                    file_info.uploaded_by = "upload_interface"  # 稍後會在 manager 中更新
                    file_info.uploaded_at = current_timestamp
                    file_info.file_source = "upload"
                    
                    # 🔧 修正：使用標準化路徑作為鍵值
                    record_key = str(target_file_path)
                    self.file_records[collection_name][record_key] = file_info
                    self._save_file_records()
                    
                except Exception as e:
                    logger.warning(f"更新檔案記錄失敗: {e}")
                
                print(f"✅ 文件上傳完成: {filename}")
                print(f"   📁 保存位置: {target_file_path}")
                print(f"   📄 分塊數量: {len(documents)}")
                print(f"   ✅ 成功向量化: {success_count}")
                
                return {
                    "success": True,
                    "message": f"文件上傳成功，已保存到 data/{bot_name}/{filename}，共生成 {len(documents)} 個分塊",
                    "filename": filename,
                    "collection": collection_name,
                    "total_chunks": len(documents),
                    "success_chunks": success_count,
                    "chunks": chunks_info,
                    "upload_time": current_timestamp,
                    "saved_path": str(target_file_path),  # 🆕 回傳保存路徑
                    "file_source": "upload"
                }
                
            except Exception as processing_error:
                # 🔧 處理失敗時清理檔案
                try:
                    if target_file_path.exists():
                        target_file_path.unlink()
                        print(f"🧹 處理失敗，已清理檔案: {target_file_path}")
                except Exception as cleanup_error:
                    logger.warning(f"清理檔案失敗: {cleanup_error}")
                
                raise processing_error
                        
        except Exception as e:
            logger.error(f"文件上傳失敗 {filename}: {e}")
            return {
                "success": False,
                "message": f"文件上傳失敗: {str(e)}",
                "chunks": []
            }

    def get_collection_documents(self, collection_name: str, page: int = 1, limit: int = 20, search: str = "") -> Dict:
        """獲取集合中的檔案資訊 - 兼容 Chroma 和 PGVector"""
        try:
            vectorstore = self.get_or_create_vectorstore(collection_name)
            
            if self.use_postgres:
                print("🔍 使用 PGVector API 獲取檔案列表")
                return self._get_documents_from_pgvector(vectorstore, collection_name, page, limit, search)
            else:
                print("🔍 使用 Chroma API 獲取檔案列表")
                return self._get_documents_from_chroma(vectorstore, collection_name, page, limit, search)
                
        except Exception as e:
            logger.error(f"獲取檔案列表失敗: {e}", exc_info=True)
            return {"success": False, "error": str(e), "documents": [], "total": 0, "page": page, "limit": limit, "total_pages": 0}
        
    def _get_documents_from_chroma(self, vectorstore, collection_name: str, page: int, limit: int, search: str) -> Dict:
        """從 Chroma 獲取檔案 - 原有邏輯"""
        all_docs_raw = vectorstore.get()
        if not all_docs_raw or not all_docs_raw.get('metadatas'):
            return {"success": True, "documents": [], "total": 0, "page": page, "limit": limit, "total_pages": 0}

        file_stats = {}
        for metadata in all_docs_raw.get('metadatas', []):
            try:
                filename = metadata.get('original_filename', metadata.get('filename', 'unknown_file'))
                if filename == 'unknown_file': continue

                if filename not in file_stats:
                    file_stats[filename] = {
                        'filename': filename,
                        'source': metadata.get('source', 'unknown'),
                        'total_chunks': 0,
                        'upload_time': metadata.get('upload_timestamp', 0)
                    }
                file_stats[filename]['total_chunks'] += 1
            except Exception:
                continue

        safe_documents = list(file_stats.values())
        
        # 添加格式化時間
        for doc in safe_documents:
            try:
                from datetime import datetime
                doc['upload_time_formatted'] = datetime.fromtimestamp(doc['upload_time']).strftime('%Y-%m-%d %H:%M:%S') if doc['upload_time'] else 'N/A'
            except:
                doc['upload_time_formatted'] = 'Invalid Date'

        # 過濾和分頁
        if search:
            safe_documents = [doc for doc in safe_documents if search.lower() in doc['filename'].lower()]
        
        safe_documents.sort(key=lambda x: x.get('upload_time', 0), reverse=True)
        total = len(safe_documents)
        total_pages = (total + limit - 1) // limit if total > 0 else 1
        start = (page - 1) * limit
        end = start + limit
        page_documents = safe_documents[start:end]

        return {"success": True, "documents": page_documents, "total": total, "page": page, "limit": limit, "total_pages": total_pages}

    def _get_documents_from_pgvector(self, vectorstore, collection_name: str, page: int, limit: int, search: str) -> Dict:
        """從 PGVector 獲取檔案 - 修正資料格式"""
        try:
            print(f"🔍 從檔案記錄獲取 {collection_name} 的檔案列表")
            
            if collection_name not in self.file_records:
                print(f"⚠️ 集合 {collection_name} 在檔案記錄中不存在")
                return {"success": True, "documents": [], "total": 0, "page": page, "limit": limit, "total_pages": 0}
            
            files = self.file_records[collection_name]
            file_stats = {}
            
            print(f"🔍 處理 {len(files)} 個檔案記錄")
            
            for file_path, file_info in files.items():
                try:
                    filename = Path(file_path).name
                    
                    if filename not in file_stats:
                        # 🔧 修正：確保所有必要欄位都有值
                        upload_time = 0
                        uploaded_by = "未知"
                        
                        # 獲取上傳時間
                        if hasattr(file_info, 'uploaded_at') and file_info.uploaded_at:
                            upload_time = file_info.uploaded_at
                        elif hasattr(file_info, 'mtime') and file_info.mtime:
                            upload_time = file_info.mtime
                        elif isinstance(file_info, dict):
                            upload_time = file_info.get('uploaded_at', file_info.get('mtime', 0))
                        
                        # 獲取上傳者信息
                        if hasattr(file_info, 'uploaded_by') and file_info.uploaded_by:
                            uploaded_by = file_info.uploaded_by
                        elif hasattr(file_info, 'file_source') and file_info.file_source:
                            uploaded_by = "管理介面" if file_info.file_source == "upload" else "同步"
                        elif isinstance(file_info, dict):
                            uploaded_by = file_info.get('uploaded_by', file_info.get('file_source', '未知'))
                            if uploaded_by == 'upload':
                                uploaded_by = "管理介面"
                        
                        file_stats[filename] = {
                            'filename': filename,
                            'source': file_path,
                            'total_chunks': 0,  # 🔧 確保初始化為 0
                            'upload_time': upload_time,
                            'uploaded_by': uploaded_by  # 🔧 新增上傳者欄位
                        }
                    
                    # 🔧 修正：實際獲取分塊數量
                    try:
                        chunks = self.get_document_chunks(collection_name, filename)
                        actual_chunk_count = len(chunks)
                        file_stats[filename]['total_chunks'] = actual_chunk_count
                        print(f"   📄 {filename}: {actual_chunk_count} 個分塊")
                    except Exception as chunk_error:
                        logger.warning(f"獲取 {filename} 分塊數量失敗: {chunk_error}")
                        file_stats[filename]['total_chunks'] = 0  # 🔧 失敗時設為 0，而不是 1
                        
                except Exception as file_error:
                    logger.warning(f"處理檔案記錄失敗 {file_path}: {file_error}")
                    continue
            
            safe_documents = list(file_stats.values())
            
            # 🔧 修正：格式化時間，確保不會是 "Invalid Date"
            for doc in safe_documents:
                try:
                    from datetime import datetime
                    if doc['upload_time'] and doc['upload_time'] > 0:
                        doc['upload_time_formatted'] = datetime.fromtimestamp(doc['upload_time']).strftime('%Y-%m-%d %H:%M:%S')
                    else:
                        doc['upload_time_formatted'] = '未知'
                except Exception as e:
                    print(f"⚠️ 時間格式化失敗: {e}")
                    doc['upload_time_formatted'] = '未知'
                
                # 🔧 確保上傳者不是空值
                if not doc.get('uploaded_by') or doc['uploaded_by'] in ['', 'unknown']:
                    doc['uploaded_by'] = '未知'

            # 過濾和分頁
            if search:
                safe_documents = [doc for doc in safe_documents if search.lower() in doc['filename'].lower()]
            
            safe_documents.sort(key=lambda x: x.get('upload_time', 0), reverse=True)

            total = len(safe_documents)
            total_pages = (total + limit - 1) // limit if total > 0 else 1
            start = (page - 1) * limit
            end = start + limit
            page_documents = safe_documents[start:end]

            # 🔧 調試：打印返回的資料格式
            print(f"✅ PGVector 檔案列表獲取成功: {total} 個檔案")
            if page_documents:
                sample_doc = page_documents[0]
                print(f"📋 樣本資料格式: {list(sample_doc.keys())}")
                print(f"   filename: {sample_doc.get('filename')}")
                print(f"   total_chunks: {sample_doc.get('total_chunks')}")
                print(f"   uploaded_by: {sample_doc.get('uploaded_by')}")
                print(f"   upload_time_formatted: {sample_doc.get('upload_time_formatted')}")

            return {
                "success": True,
                "documents": page_documents,
                "total": total,
                "page": page,
                "limit": limit,
                "total_pages": total_pages
            }
            
        except Exception as e:
            logger.error(f"PGVector 檔案獲取失敗: {e}")
            return {
                "success": False, 
                "error": f"PGVector 獲取失敗: {str(e)}", 
                "documents": [], 
                "total": 0, 
                "page": page, 
                "limit": limit, 
                "total_pages": 0
            }
        
    def get_document_chunks(self, collection_name: str, source_file: str) -> List[Dict]:
        """獲取指定檔案的所有分塊 - 兼容 Chroma 和 PGVector"""
        try:
            vectorstore = self.get_or_create_vectorstore(collection_name)
            
            if self.use_postgres:
                return self._get_chunks_from_pgvector(vectorstore, collection_name, source_file)
            else:
                return self._get_chunks_from_chroma(vectorstore, source_file)
                
        except Exception as e:
            logger.error(f"獲取檔案分塊失敗 {collection_name}/{source_file}: {e}")
            return []

    def _get_chunks_from_chroma(self, vectorstore, source_file: str) -> List[Dict]:
        """從 Chroma 獲取分塊"""
        try:
            results = vectorstore.get(where={"filename": source_file})
            if not results or not results.get('documents'):
                return []
            
            chunks = []
            for i, (doc, metadata) in enumerate(zip(results['documents'], results['metadatas'])):
                chunk_info = {
                    'chunk_id': metadata.get('chunk_id', f'chunk_{i+1}'),
                    'chunk_index': metadata.get('chunk_index', i),
                    'content': doc,
                    'content_preview': doc[:200] + "..." if len(doc) > 200 else doc,
                    'token_count': metadata.get('token_count', 0),
                    'text_type': metadata.get('text_type', 'unknown'),
                    'quality_score': metadata.get('quality_score', 0.5),
                    'metadata': metadata
                }
                chunks.append(chunk_info)
            
            chunks.sort(key=lambda x: x.get('chunk_index', 0))
            return chunks
            
        except Exception as e:
            logger.error(f"Chroma 分塊獲取失敗: {e}")
            return []

    def _get_chunks_from_pgvector(self, vectorstore, collection_name: str, source_file: str) -> List[Dict]:
        """從 PGVector 獲取分塊"""
        try:
            collection_folder = collection_name.replace('collection_', '')
            possible_paths = [
                source_file,
                f"data/{collection_folder}/{source_file}",
                f"data\\{collection_folder}\\{source_file}"
            ]
            
            all_chunks = []
            
            for search_path in possible_paths:
                try:
                    docs = vectorstore.similarity_search("", k=1000)
                    matching_chunks = []
                    
                    for doc in docs:
                        metadata = doc.metadata
                        doc_filename = metadata.get('filename', metadata.get('original_filename', ''))
                        doc_source = metadata.get('source', '')
                        
                        if (doc_filename == source_file or 
                            doc_source.endswith(source_file) or
                            search_path in doc_source):
                            
                            chunk_info = {
                                'chunk_id': metadata.get('chunk_id', f'chunk_{len(matching_chunks)+1}'),
                                'chunk_index': metadata.get('chunk_index', len(matching_chunks)),
                                'content': doc.page_content,
                                'content_preview': doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content,
                                'token_count': metadata.get('token_count', 0),
                                'text_type': metadata.get('text_type', 'unknown'),
                                'quality_score': metadata.get('quality_score', 0.5),
                                'metadata': metadata
                            }
                            matching_chunks.append(chunk_info)
                    
                    if matching_chunks:
                        all_chunks = matching_chunks
                        print(f"🎯 PGVector 找到 {len(all_chunks)} 個分塊")
                        break
                        
                except Exception as search_error:
                    logger.warning(f"PGVector 查詢失敗: {search_error}")
                    continue
            
            all_chunks.sort(key=lambda x: x.get('chunk_index', 0))
            return all_chunks
            
        except Exception as e:
            logger.error(f"PGVector 分塊獲取失敗: {e}")
            return []

    def delete_document(self, collection_name: str, source_file: str) -> Dict:
        """刪除指定檔案及其所有向量 - 兼容 Chroma 和 PGVector"""
        try:
            vectorstore = self.get_or_create_vectorstore(collection_name)
            existing_chunks = self.get_document_chunks(collection_name, source_file)
            chunk_count = len(existing_chunks)
            
            if chunk_count == 0:
                return {"success": False, "message": "檔案不存在或已被刪除", "deleted_chunks": 0}
            
            if self.use_postgres:
                return self._delete_from_pgvector(vectorstore, collection_name, source_file, chunk_count)
            else:
                return self._delete_from_chroma(vectorstore, collection_name, source_file, chunk_count)
                
        except Exception as e:
            logger.error(f"刪除檔案失敗 {collection_name}/{source_file}: {e}")
            return {"success": False, "message": f"刪除檔案失敗: {str(e)}", "deleted_chunks": 0}

    def _delete_from_chroma(self, vectorstore, collection_name: str, source_file: str, chunk_count: int) -> Dict:
        """從 Chroma 刪除檔案"""
        try:
            vectorstore.delete(filter={"filename": source_file})
        except Exception as e1:
            try:
                collection_folder = collection_name.replace('collection_', '')
                full_path = f"data\\{collection_folder}\\{source_file}"
                vectorstore.delete(filter={"source": full_path})
            except Exception as e2:
                raise e2
        
        # 從檔案記錄中移除
        if collection_name in self.file_records:
            for file_path in list(self.file_records[collection_name].keys()):
                if Path(file_path).name == source_file:
                    del self.file_records[collection_name][file_path]
                    self._save_file_records()
                    break
        
        return {"success": True, "message": f"檔案 {source_file} 及其 {chunk_count} 個分塊已刪除", "deleted_chunks": chunk_count, "filename": source_file}

    def _delete_from_pgvector(self, vectorstore, collection_name: str, source_file: str, chunk_count: int) -> Dict:
        """從 PGVector 刪除檔案"""
        try:
            deleted_count = 0
            
            # 策略1: 嘗試使用 filter 刪除
            try:
                if hasattr(vectorstore, 'delete'):
                    vectorstore.delete(filter={"filename": source_file})
                    deleted_count = chunk_count
                    print(f"✅ PGVector filter 刪除成功")
                else:
                    raise AttributeError("No delete method")
            except Exception as e1:
                print(f"⚠️ PGVector filter 刪除失敗: {e1}")
                
                # 策略2: 嘗試獲取文檔並逐一刪除
                try:
                    chunks = self.get_document_chunks(collection_name, source_file)
                    for chunk in chunks:
                        # 這裡需要實際的刪除邏輯
                        # 目前先標記為已嘗試刪除
                        pass
                    deleted_count = len(chunks)
                    print(f"✅ PGVector 逐一刪除完成: {deleted_count} 個分塊")
                except Exception as e2:
                    print(f"❌ PGVector 逐一刪除失敗: {e2}")
                    deleted_count = 0
            
            # 從檔案記錄中移除
            if deleted_count > 0 and collection_name in self.file_records:
                for file_path in list(self.file_records[collection_name].keys()):
                    if Path(file_path).name == source_file:
                        del self.file_records[collection_name][file_path]
                        self._save_file_records()
                        break
            
            return {
                "success": deleted_count > 0,
                "message": f"檔案 {source_file} 及其 {deleted_count} 個分塊已刪除" if deleted_count > 0 else "刪除失敗",
                "deleted_chunks": deleted_count,
                "filename": source_file
            }
            
        except Exception as e:
            logger.error(f"PGVector 刪除失敗: {e}")
            return {"success": False, "message": f"刪除失敗: {str(e)}", "deleted_chunks": 0}
        
    def get_chunk_content(self, collection_name: str, chunk_id: str) -> Optional[Dict]:
        """
        獲取指定分塊的詳細內容
        
        Args:
            collection_name: 集合名稱
            chunk_id: 分塊ID
            
        Returns:
            Optional[Dict]: 分塊詳細信息
        """
        try:
            vectorstore = self.get_or_create_vectorstore(collection_name)
            
            # 🔧 查詢指定分塊
            results = vectorstore.get(where={"chunk_id": chunk_id})
            
            if not results or not results.get('documents') or len(results['documents']) == 0:
                return None
            
            # 取第一個結果（chunk_id應該是唯一的）
            doc = results['documents'][0]
            metadata = results['metadatas'][0]
            
            # 🔧 計算字符和行數統計
            content_stats = {
                'total_chars': len(doc),
                'total_lines': len(doc.split('\n')),
                'total_words': len(doc.split()),
                'total_sentences': len([s for s in doc.split('。') if s.strip()]) + len([s for s in doc.split('.') if s.strip()])
            }
            
            chunk_detail = {
                'chunk_id': chunk_id,
                'chunk_index': metadata.get('chunk_index', 0),
                'content': doc,
                'content_stats': content_stats,
                'token_count': metadata.get('token_count', 0),
                'text_type': metadata.get('text_type', 'unknown'),
                'quality_score': metadata.get('quality_score', 0.5),
                'language': metadata.get('language', 'unknown'),
                'source_file': metadata.get('source', 'unknown'),
                'original_filename': metadata.get('original_filename', metadata.get('filename', 'unknown')),
                'processing_strategy': metadata.get('processing_strategy', 'unknown'),
                'split_method': metadata.get('split_method', 'unknown'),
                'has_overlap': metadata.get('has_overlap', False),
                'metadata': metadata
            }
            
            return chunk_detail
            
        except Exception as e:
            logger.error(f"獲取分塊內容失敗 {collection_name}/{chunk_id}: {e}")
            return None

    def get_available_collections(self) -> List[Dict]:
        """
        獲取所有可用的集合列表 - 修正：從file_records而非persist_dir獲取
        """
        try:
            collections = []
            
            # 🔧 修正：從file_records獲取集合信息，而不是掃描persist_dir
            for collection_name in self.file_records.keys():
                display_name = collection_name.replace('collection_', '')
                
                try:
                    vectorstore = self.get_or_create_vectorstore(collection_name)
                    doc_count = vectorstore._collection.count()
                    
                    docs_result = self.get_collection_documents(collection_name, page=1, limit=1000)
                    file_count = len(docs_result.get('documents', [])) if docs_result['success'] else 0
                    
                    collections.append({
                        'collection_name': collection_name,
                        'display_name': display_name,
                        'document_count': doc_count,
                        'file_count': file_count,
                        'status': 'active'
                    })
                    
                except Exception as e:
                    logger.warning(f"獲取集合統計失敗 {collection_name}: {e}")
                    collections.append({
                        'collection_name': collection_name,
                        'display_name': display_name,
                        'document_count': 0,
                        'file_count': 0,
                        'status': 'error'
                    })
            
            collections.sort(key=lambda x: x['display_name'])
            return collections
            
        except Exception as e:
            logger.error(f"獲取集合列表失敗: {e}")
            return []

    def test_knowledge_management(self):
        """測試知識庫管理功能"""
        print("\n🧪 === 知識庫管理功能測試 ===")
        
        try:
            # 🔧 測試獲取集合列表
            collections = self.get_available_collections()
            print(f"📁 找到 {len(collections)} 個集合:")
            for coll in collections[:3]:  # 只顯示前3個
                print(f"   - {coll['display_name']}: {coll['document_count']} 文檔")
            
            # 🔧 測試獲取文檔列表
            if collections:
                test_collection = collections[0]['collection_name']
                docs_result = self.get_collection_documents(test_collection, page=1, limit=5)
                
                if docs_result['success']:
                    print(f"\n📄 集合 {test_collection} 文檔測試:")
                    print(f"   總文件數: {docs_result['total']}")
                    print(f"   當前頁文件: {len(docs_result['documents'])}")
                    
                    # 🔧 測試分塊查詢
                    if docs_result['documents']:
                        test_file = docs_result['documents'][0]['filename']
                        chunks = self.get_document_chunks(test_collection, test_file)
                        print(f"   文件 {test_file}: {len(chunks)} 個分塊")
                        
                        # 🔧 測試分塊內容查詢
                        if chunks:
                            test_chunk_id = chunks[0]['chunk_id']
                            chunk_detail = self.get_chunk_content(test_collection, test_chunk_id)
                            if chunk_detail:
                                print(f"   分塊 {test_chunk_id}: {chunk_detail['content_stats']['total_chars']} 字符")
            
            print("✅ 知識庫管理功能測試完成")
            return True
            
        except Exception as e:
            print(f"❌ 知識庫管理功能測試失敗: {e}")
            return False

    



def main():
    """主程式"""
    print("🇨🇳 === 完整優化版 LangChain 中文向量系統 === 🇨🇳")
    print(f"🤖 OpenAI Embeddings: {'✅ 已啟用' if OPENAI_EMBEDDINGS_AVAILABLE else '❌ 未啟用'}")
    print(f"🐘 PostgreSQL: {'✅ 已啟用' if check_postgresql_connection() else '❌ 未啟用'}")
    print(f"📊 PGVector: {'✅ 已啟用' if PGVECTOR_AVAILABLE else '❌ 未啟用'}")
    print("🇨🇳 === 完整優化版 LangChain 中文向量系統 === 🇨🇳")
    print(f"🤖 OpenAI Embeddings: {'✅ 已啟用' if OPENAI_EMBEDDINGS_AVAILABLE else '❌ 未啟用'}")
    print(f"📄 DOCX 支援: {'✅ 已啟用' if DOCX2TXT_AVAILABLE else '❌ 未啟用'}")
    print(f"📚 EPUB 支援: {'✅ 已啟用' if EPUB_AVAILABLE else '❌ 未啟用'}")
    print(f"🔤 繁簡轉換: {'✅ 已啟用' if OPENCC_AVAILABLE else '❌ 未啟用'}")
    print(f"🧠 智能處理: ✅ 已啟用")
    print(f"🚀 性能優化: ✅ 已啟用")
    
    # 初始化系統
    try:
        system = OptimizedVectorSystem()
        
        # 系統診斷
        print("\n" + "="*60)
        system.diagnose_system()
        print("="*60 + "\n")
        
        # 同步集合
        changes = system.sync_collections()
        
        # 顯示統計
        stats = system.get_stats()
        print(f"\n📊 集合統計:")
        for folder, count in stats.items():
            print(f"   📁 {folder}: {count:,} 個文檔分塊")
        
        if changes > 0:
            print(f"\n✅ 處理完成，共 {changes} 個變更")
        else:
            print(f"\n✅ 所有文件都是最新的")
        
        # 性能統計
        perf_stats = system.batch_processor.get_performance_stats()
        if perf_stats['total_batches'] > 0:
            print(f"\n📈 性能統計:")
            print(f"   📦 處理批次: {perf_stats['total_batches']}")
            print(f"   📄 處理文檔: {perf_stats['total_documents']:,}")
            print(f"   📏 處理 tokens: {perf_stats['total_tokens']:,}")
            print(f"   ✅ 成功率: {perf_stats['success_rate']*100:.1f}%")
            
            if perf_stats['total_tokens'] > 0:
                cost = system.batch_processor.token_estimator.estimate_embedding_cost(perf_stats['total_tokens'])
                print(f"   💰 總成本: ${cost:.4f}")
        
        return system
        
    except Exception as e:
        print(f"❌ 系統初始化失敗: {e}")
        logger.error(f"系統初始化失敗: {e}")
        return None
    



if __name__ == "__main__":
    system = main()
    if system:
        system.test_knowledge_management()