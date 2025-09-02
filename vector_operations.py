#!/usr/bin/env python3
"""
向量操作核心層 - VectorOperationsCore
職責：底層數據操作、向量處理、文檔載入、集合管理
包含26個底層核心方法，從原 OptimizedVectorSystem 精確移動而來
"""

import time
import json
import hashlib
import re
import logging
import os
import gc
import threading
import tempfile
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Set, Any, Union
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, asdict

# LangChain核心
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    TextLoader, PyPDFLoader, Docx2txtLoader, 
    CSVLoader, JSONLoader, PandasExcelLoader
)

# --- Start of content from core_config.py ---

# 🆕 1. 添加 Railway 環境檢測函數（放在文件頂部，導入語句之後）
def detect_railway_environment():
    """檢測是否在 Railway 環境中運行"""
    railway_indicators = [
        os.getenv("RAILWAY_PROJECT_ID"),
        os.getenv("RAILWAY_SERVICE_ID"), 
        os.getenv("DATABASE_URL"),
        "railway.internal" in os.getenv("POSTGRES_HOST", "")
    ]
    
    is_railway = any(railway_indicators)
    if is_railway:
        print("🚂 檢測到 Railway 部署環境")
        # 顯示 Railway 特定信息
        project_id = os.getenv("RAILWAY_PROJECT_ID", "unknown")
        service_id = os.getenv("RAILWAY_SERVICE_ID", "unknown")  
        print(f"   項目ID: {project_id[:8]}***")
        print(f"   服務ID: {service_id[:8]}***")
    
    return is_railway

# 🔧 優化版系統配置
SYSTEM_CONFIG = {
    "persist_dir": "./chroma_langchain_db",
    "data_dir": "./data", 
    "device": "cpu",  
    "log_level": "INFO",
    "max_workers": 4,  # 並發處理數
    "cache_embeddings": True,  # 快取嵌入向量
    "backup_enabled": True  # 自動備份
}

# 🧠 智能文本分類配置
SMART_TEXT_CONFIG = {
    # 文本長度閾值（字符數）
    "micro_text_threshold": 100,      # 微文本：標題、摘要等
    "short_text_threshold": 500,      # 短文本：短文章、段落
    "medium_text_threshold": 2000,    # 中等文本：中篇文章
    "long_text_threshold": 8000,    # 長文本：長篇文章
    "ultra_long_threshold": 20000,    # 超長文本：書籍章節
    
    # 各類文本的處理策略
    "micro_text": {
        "min_length": 20,
        "process_whole": True,
        "merge_with_next": True,  # 嘗試與下一個片段合併
        "chunk_size": 0,
        "description": "微文本整體處理"
    },
    
    "short_text": {
        "min_length": 50,
        "process_whole": True,
        "preserve_structure": True,
        "chunk_size": 0,
        "allow_merge": True,  # 允許合併相鄰短文本
        "description": "短文本整體保存"
    },
    
    "medium_text": {
        "chunk_size": 800,
        "chunk_overlap": 120,
        "preserve_paragraphs": True,
        "smart_boundaries": True,  # 智能邊界檢測
        "min_chunk_ratio": 0.3,   # 最小片段比例
        "description": "中等文本段落感知分割"
    },
    
    "long_text": {
        "chunk_size": 500,
        "chunk_overlap": 80,
        "hierarchical_split": True,
        "section_aware": True,    # 章節感知
        "quality_check": True,    # 質量檢查
        "description": "長文本精細化分割"
    },
    
    "ultra_long": {
        "chunk_size": 600,
        "chunk_overlap": 100,
        "multi_level_split": True,  # 多層級分割
        "chapter_detection": True,  # 章節檢測
        "summary_chunks": True,     # 生成摘要片段
        "description": "超長文本階層式處理"
    },
    
    "mega_text": {
        "chunk_size": 500,
        "chunk_overlap": 80,
        "multi_level_split": True,  # 多層級分割
        "chapter_detection": True,  # 章節檢測
        "aggressive_split": True,   # 積極分割
        "quality_filter": True,     # 質量過濾
        "description": "超大文本積極分割處理"
    }
}

# 🔧 優化版 Token 限制配置
TOKEN_LIMITS = {
    "max_tokens_per_request": 150000,  # 更保守的限制
    "max_batch_size": 8,               # 減小批次大小
    "min_batch_size": 1,
    "batch_delay": 5.0,                # 增加批次間延遲
    "retry_delay": 10,
    "max_retries": 3,
    "token_safety_margin": 0.2,        # 20% 安全邊際
    "adaptive_batching": True           # 自適應批次大小
}

# 🔧 性能優化配置
PERFORMANCE_CONFIG = {
    "embedding_batch_size": 12,
    "parallel_processing": True,
    "memory_limit_mb": 1024,      # 記憶體限制
    "chunk_cache_size": 1000,     # 分塊快取大小
    "preload_models": True,       # 預載入模型
    "gc_frequency": 50            # 垃圾回收頻率
}

# 支持的文件格式擴展
SUPPORTED_EXTENSIONS = {
    '.txt', '.md', '.pdf', '.csv', '.json', '.py', '.js', 
    '.docx', '.doc', '.epub', '.rst', '.org', '.xlsx', '.xls'
}

from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from langchain_community.vectorstores import Chroma

try:
    from langchain_community.vectorstores import Chroma
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False
    # 創建一個虛擬的 Chroma 類型用於類型註解
    class Chroma:
        pass

# 🔧 檢查環境變數
def check_openai_api_key():
    """檢查 OpenAI API Key"""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("⚠️ 未設置 OPENAI_API_KEY 環境變數")
        print("   請在 .env 文件中設置: OPENAI_API_KEY=sk-your-api-key")
        return False
    
    if api_key.startswith("sk-") and len(api_key) > 20:
        print("✅ OpenAI API Key 格式正確")
        print(f"   API Key: {api_key[:8]}***{api_key[-4:]}")
        return True
    else:
        print("⚠️ OpenAI API Key 格式可能不正確")
        return False

# LangChain 核心組件導入
try:
    from langchain_postgres import PGVector
    PGVECTOR_AVAILABLE = True
    print("✅ PGVector 可用")
except ImportError:
    try:
        from langchain_community.vectorstores import PGVector
        PGVECTOR_AVAILABLE = True
        print("✅ PGVector (community) 可用")
    except ImportError:
        PGVECTOR_AVAILABLE = False
        print("❌ PGVector 不可用")
        # 可以回退到 Chroma
        from langchain_community.vectorstores import Chroma

# OpenAI embeddings 導入
try:
    from langchain_openai import OpenAIEmbeddings
    OPENAI_EMBEDDINGS_AVAILABLE = True
    print("✅ OpenAI Embeddings 可用")
    
    if check_openai_api_key():
        print("✅ OpenAI 環境檢查通過")
    else:
        print("⚠️ OpenAI 環境配置可能有問題")
        
except ImportError as e:
    OPENAI_EMBEDDINGS_AVAILABLE = False
    print(f"⚠️ OpenAI Embeddings 不可用: {e}")

# 🔧 PostgreSQL 依賴檢查
try:
    import psycopg2
    PSYCOPG2_AVAILABLE = True
except ImportError:
    PSYCOPG2_AVAILABLE = False
    print("⚠️ 警告: psycopg2 未安裝，請執行: pip install psycopg2-binary")

# 中文分詞和轉換
try:
    import jieba
    jieba.initialize()
    JIEBA_AVAILABLE = True
except ImportError:
    JIEBA_AVAILABLE = False
    print("⚠️ 警告: jieba 未安裝，中文分詞功能將被禁用")

try:
    import opencc
    OPENCC_AVAILABLE = True
except ImportError:
    OPENCC_AVAILABLE = False
    print("⚠️ 警告: opencc 未安裝，繁簡轉換功能將被禁用")

# 文檔處理支持
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
    DOCX_METHOD = "docx2txt"
    print("✅ DOCX 支持已啟用 (docx2txt)")
except ImportError:
    try:
        from docx import Document as DocxDocument
        DOCX2TXT_AVAILABLE = True
        DOCX_METHOD = "python-docx"
        print("✅ DOCX 支持已啟用 (python-docx)")
    except ImportError:
        DOCX2TXT_AVAILABLE = False
        DOCX_METHOD = None
        print("⚠️ 警告: docx2txt 或 python-docx 未安裝")

try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    EPUB_AVAILABLE = True
    print("✅ EPUB 支持已啟用")
except ImportError:
    EPUB_AVAILABLE = False
    print("⚠️ 警告: ebooklib 或 beautifulsoup4 未安裝")

@dataclass
class FileInfo:
    """文件資訊"""
    path: str
    size: int
    mtime: float
    hash: str
    encoding: str = "utf-8"
    file_type: str = ""

@dataclass
class TextAnalysis:
    """文本分析結果"""
    length: int
    text_type: str
    language: str
    encoding: str
    structure_info: Dict
    quality_score: float
    processing_strategy: str

@dataclass
class ChunkInfo:
    """分塊資訊"""
    chunk_id: str
    content: str
    metadata: Dict
    token_count: int
    quality_score: float
    relationships: List[str]  # 與其他分塊的關係

@dataclass
class SearchResult:
    content: str
    score: float
    metadata: Dict
    collection: str
    chunk_info: Optional[ChunkInfo] = None

# --- End of content from core_config.py ---


# --- Start of content from text_processing.py ---

logger = logging.getLogger(__name__)

class AdvancedTokenEstimator:
    """🔧 高級 Token 估算器"""
    
    def __init__(self):
        # 不同語言和內容類型的 Token 係數
        self.token_ratios = {
            "chinese": 2.5,      # 中文字符/token
            "english": 4.0,      # 英文字符/token
            "mixed": 3.0,        # 中英混合
            "code": 3.5,         # 程式碼
            "punctuation": 1.0,  # 標點符號
            "numbers": 2.0       # 數字
        }
        self.safety_margin = TOKEN_LIMITS.get("token_safety_margin", 0.15)
    
    def analyze_text_composition(self, text: str) -> Dict[str, int]:
        """分析文本組成"""
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_chars = len(re.findall(r'[a-zA-Z]', text))
        numbers = len(re.findall(r'\d', text))
        punctuation = len(re.findall(r'[^\w\s\u4e00-\u9fff]', text))
        other_chars = len(text) - chinese_chars - english_chars - numbers - punctuation
        
        return {
            "chinese": chinese_chars,
            "english": english_chars,
            "numbers": numbers,
            "punctuation": punctuation,
            "other": other_chars,
            "total": len(text)
        }
    
    def estimate_tokens(self, text: str, content_type: str = "mixed") -> int:
        """精確估算 Token 數量"""
        if not text:
            return 0
        
        composition = self.analyze_text_composition(text)
        
        # 根據文本組成計算 token
        estimated_tokens = (
            composition["chinese"] / self.token_ratios["chinese"] +
            composition["english"] / self.token_ratios["english"] +
            composition["numbers"] / self.token_ratios["numbers"] +
            composition["punctuation"] / self.token_ratios["punctuation"] +
            composition["other"] / self.token_ratios["mixed"]
        )
        
        # 根據內容類型調整
        if content_type == "code":
            estimated_tokens *= 1.2  # 程式碼通常 token 密度更高
        elif content_type == "academic":
            estimated_tokens *= 1.1  # 學術文本專業詞彙多
        
        # 加上安全邊際
        final_estimate = int(estimated_tokens * (1 + self.safety_margin))
        
        return max(final_estimate, 1)  # 至少1個token
    
    def estimate_embedding_cost(self, total_tokens: int, model: str = "text-embedding-3-small") -> float:
        """估算 Embedding 成本"""
        cost_per_1k_tokens = {
            "text-embedding-3-small": 0.00002,
            "text-embedding-3-large": 0.00013,
            "text-embedding-ada-002": 0.0001
        }
        
        rate = cost_per_1k_tokens.get(model, 0.00002)
        return (total_tokens / 1000) * rate

class ChineseTextNormalizer:
    """優化版中文文本標準化處理器"""
    
    def __init__(self):
        self.s2t_converter = None
        self.t2s_converter = None
        
        if OPENCC_AVAILABLE:
            try:
                self.s2t_converter = opencc.OpenCC('s2t')
                self.t2s_converter = opencc.OpenCC('t2s')
                print("🔤 繁簡轉換器初始化完成")
            except Exception as e:
                logger.warning(f"OpenCC 初始化失敗: {e}")
        
        # 中文文本正規化規則
        self.normalization_rules = [
            (r'[\u3000]+', ' '),
            (r'\r\n|\r', '\n'),                       # 統一換行
            (r'\n{3,}', '\n\n'),                      # 限制連續換行
            (r'[\u201C\u201D\u2018\u2019\u201E\u201A\u2033\u2032]', '"'),  # 統一引號
            (r'[——–∶]', '-'),                          # 統一破折號
            (r'[…⋯]', '...'),                         # 統一省略號
        ]
    
    def normalize_text(self, text: str) -> Tuple[str, Dict]:
        """標準化文本並返回處理資訊"""
        if not text:
            return "", {}
        
        original_length = len(text)
        processed_text = text
        
        # 應用正規化規則
        for pattern, replacement in self.normalization_rules:
            processed_text = re.sub(pattern, replacement, processed_text)
        
        # 移除首尾空白
        processed_text = processed_text.strip()
        
        # 檢測和轉換繁簡
        variant, confidence = self.detect_chinese_variant(processed_text)
        if variant == 'simplified' and self.s2t_converter:
            try:
                processed_text = self.s2t_converter.convert(processed_text)
                variant = 'traditional_converted'
            except Exception as e:
                logger.warning(f"繁簡轉換失敗: {e}")
        
        processing_info = {
            'original_length': original_length,
            'processed_length': len(processed_text),
            'variant': variant,
            'confidence': confidence,
            'normalized': original_length != len(processed_text)
        }
        
        return processed_text, processing_info
    
    def detect_chinese_variant(self, text: str) -> Tuple[str, float]:
        """檢測繁體或簡體中文"""
        if not text:
            return 'unknown', 0.0
        
        simplified_chars = set('国发会学习论问题业专长时间经济')
        traditional_chars = set('國發會學習論問題業專長時間經濟')
        
        simplified_count = sum(1 for char in text if char in simplified_chars)
        traditional_count = sum(1 for char in text if char in traditional_chars)
        total_chinese = len(re.findall(r'[\u4e00-\u9fff]', text))
        
        if total_chinese == 0:
            return 'unknown', 0.0
        
        simplified_ratio = simplified_count / total_chinese
        traditional_ratio = traditional_count / total_chinese
        
        if simplified_ratio > traditional_ratio:
            return 'simplified', simplified_ratio
        elif traditional_ratio > simplified_ratio:
            return 'traditional', traditional_ratio
        else:
            return 'mixed', max(simplified_ratio, traditional_ratio)
    
    def create_search_variants(self, query: str) -> List[str]:
        """創建繁簡搜索變體"""
        variants = [query]
        
        if not self.s2t_converter or not self.t2s_converter:
            return variants
        
        try:
            traditional = self.s2t_converter.convert(query)
            if traditional != query:
                variants.append(traditional)
            
            simplified = self.t2s_converter.convert(query)
            if simplified != query:
                variants.append(simplified)
        except Exception as e:
            logger.warning(f"創建搜索變體失敗: {e}")
        
        return list(set(variants))

class EpubProcessor:
    """📚 EPUB 文件處理器"""
    
    def __init__(self):
        self.normalizer = ChineseTextNormalizer()
        
    def extract_epub_content(self, file_path: Path) -> str:
        """提取 EPUB 內容"""
        try:
            if not EPUB_AVAILABLE:
                raise ImportError("EPUB 處理庫未安裝")
            
            print(f"📚 正在處理 EPUB 文件: {file_path.name}")
            
            # 讀取 EPUB 文件
            book = epub.read_epub(str(file_path))
            
            # 提取所有文本內容
            content_parts = []
            chapter_count = 0
            
            # 獲取書籍信息
            title = book.get_metadata('DC', 'title')
            author = book.get_metadata('DC', 'creator')
            
            book_info = []
            if title:
                book_info.append(f"書名: {title[0][0] if title else '未知'}")
            if author:
                book_info.append(f"作者: {author[0][0] if author else '未知'}")
            
            if book_info:
                content_parts.append("\n".join(book_info) + "\n\n")
            
            # 按順序處理章節
            spine_items = book.get_items_of_type(ebooklib.ITEM_DOCUMENT)
            
            for item in spine_items:
                try:
                    # 解析 HTML 內容
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    
                    # 移除腳本和樣式
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    # 提取文本
                    text = soup.get_text()
                    
                    if text and len(text.strip()) > 50:  # 過濾太短的內容
                        # 清理文本
                        cleaned_text = self._clean_epub_text(text)
                        
                        if cleaned_text.strip():
                            chapter_count += 1
                            
                            # 添加章節標記
                            chapter_title = self._extract_chapter_title(soup, cleaned_text)
                            if chapter_title:
                                content_parts.append(f"\n\n=== {chapter_title} ===\n")
                            else:
                                content_parts.append(f"\n\n=== 第 {chapter_count} 章 ===\n")
                            
                            content_parts.append(cleaned_text)
                            
                except Exception as e:
                    print(f"   ⚠️ 章節處理失敗: {e}")
                    continue
            
            if not content_parts:
                raise ValueError("EPUB 文件中沒有找到有效內容")
            
            full_content = "".join(content_parts)
            
            print(f"   ✅ EPUB 處理完成: {chapter_count} 個章節, {len(full_content):,} 字符")
            
            return full_content
            
        except Exception as e:
            print(f"   ❌ EPUB 處理失敗: {e}")
            raise
    
    def _clean_epub_text(self, text: str) -> str:
        """清理 EPUB 文本"""
        if not text:
            return ""
        
        # 標準化文本
        normalized_text, _ = self.normalizer.normalize_text(text)
        
        # EPUB 特定清理
        epub_cleaning_rules = [
            (r'\n{4,}', '\n\n'),                    # 限制連續換行
            (r'[ \t]{3,}', ' ' ),                    # 限制連續空格
            (r'^[ \t]+', '', re.MULTILINE),          # 移除行首空白
            (r'[ \t]+$', '', re.MULTILINE),          # 移除行尾空白
            (r'\n[ \t]*\n', '\n\n'),                # 清理空行
            (r'[\x00-\x08\x0B\x0C\x0E-\x1F]', ''),  # 移除控制字符
        ]
        
        cleaned_text = normalized_text
        for pattern, replacement, *flags in epub_cleaning_rules:
            flag = flags[0] if flags else 0
            cleaned_text = re.sub(pattern, replacement, cleaned_text, flags=flag)
        
        return cleaned_text.strip()
    
    def _extract_chapter_title(self, soup, text: str) -> Optional[str]:
        """嘗試提取章節標題"""
        try:
            # 從 HTML 標籤中查找標題
            for tag in ['h1', 'h2', 'h3', 'title']:
                title_element = soup.find(tag)
                if title_element and title_element.get_text().strip():
                    title = title_element.get_text().strip()
                    if 5 <= len(title) <= 100:  # 合理的標題長度
                        return title
            
            # 從文本開頭查找標題
            lines = text.split('\n')[:5]  # 檢查前5行
            for line in lines:
                line = line.strip()
                if line and 5 <= len(line) <= 100:
                    # 檢查是否包含章節關鍵詞
                    chapter_keywords = ['章', 'Chapter', '第', '卷', 'Part']
                    if any(keyword in line for keyword in chapter_keywords):
                        return line
                    # 或者是短行且看起來像標題
                    elif len(line) <= 50 and not line.endswith(('。', '！', '？', '.', '!', '?')):
                        return line
            
            return None
            
        except Exception:
            return None

class SmartTextAnalyzer:
    """🧠 智能文本分析器"""
    
    def __init__(self):
        self.normalizer = ChineseTextNormalizer()
        self.token_estimator = AdvancedTokenEstimator()
        
    def analyze_text(self, text: str, source_info: Dict = None) -> TextAnalysis:
        """綜合分析文本"""
        if not text:
            return TextAnalysis(
                length=0, text_type="empty", language="unknown",
                encoding="utf-8", structure_info={}, quality_score=0.0,
                processing_strategy="skip"
            )
        
        # 標準化文本
        normalized_text, norm_info = self.normalizer.normalize_text(text)
        length = len(normalized_text)
        
        # 分類文本長度
        text_type = self._classify_text_length(length)
        
        # 分析文本結構
        structure_info = self._analyze_structure(normalized_text)
        
        # 檢測語言
        language = self._detect_language(normalized_text)
        
        # 評估質量
        quality_score = self._evaluate_quality(normalized_text, structure_info)
        
        # 決定處理策略
        processing_strategy = self._determine_strategy(text_type, structure_info, quality_score)
        
        return TextAnalysis(
            length=length,
            text_type=text_type,
            language=language,
            encoding=norm_info.get('encoding', 'utf-8'),
            structure_info=structure_info,
            quality_score=quality_score,
            processing_strategy=processing_strategy
        )
    
    def _classify_text_length(self, length: int) -> str:
        """分類文本長度"""
        config = SMART_TEXT_CONFIG
        
        if length < config["micro_text_threshold"]:
            return "micro_text"
        elif length < config["short_text_threshold"]:
            return "short_text"
        elif length < config["medium_text_threshold"]:
            return "medium_text"
        elif length < config["long_text_threshold"]:
            return "long_text"
        elif length < config["ultra_long_threshold"]:
            return "ultra_long"
        else:
            return "mega_text"  # 超大文本
    
    def _analyze_structure(self, text: str) -> Dict:
        """分析文本結構"""
        structure_info = {
            'paragraphs': len(text.split('\n\n')),
            'lines': len(text.split('\n')),
            'sentences': len(re.findall(r'[。！？.!?]+', text)),
            'has_chapters': False,
            'has_sections': False,
            'has_lists': False,
            'has_tables': False,
            'chapter_count': 0,
            'section_count': 0
        }
        
        # 檢測章節結構
        chapter_patterns = [
            r'第[一二三四五六七八九十\d]+章',
            r'Chapter\s+\d+',
            r'\d+\.\s*[^\n]{1,50}\n'
        ]
        
        for pattern in chapter_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                structure_info['has_chapters'] = True
                structure_info['chapter_count'] = len(matches)
                break
        
        # 檢測小節結構
        section_patterns = [
            r'第[一二三四五六七八九十\d]+節',
            r'\d+\.\d+',
            r'[一二三四五六七八九十]、'
        ]
        
        for pattern in section_patterns:
            matches = re.findall(pattern, text)
            if matches:
                structure_info['has_sections'] = True
                structure_info['section_count'] = len(matches)
                break
        
        # 檢測列表
        if re.search(r'^\s*[•\-*]\s+', text, re.MULTILINE) or \
           re.search(r'^\s*\d+\.\s+', text, re.MULTILINE):
            structure_info['has_lists'] = True
        
        # 檢測表格
        if '|' in text and text.count('|') > 4:
            structure_info['has_tables'] = True
        
        return structure_info
    
    def _detect_language(self, text: str) -> str:
        """檢測文本主要語言"""
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_chars = len(re.findall(r'[a-zA-Z]', text))
        total_chars = len(text)
        
        if total_chars == 0:
            return "unknown"
        
        chinese_ratio = chinese_chars / total_chars
        english_ratio = english_chars / total_chars
        
        if chinese_ratio > 0.3:
            if english_ratio > 0.2:
                return "mixed_zh_en"
            else:
                return "chinese"
        elif english_ratio > 0.5:
            return "english"
        else:
            return "mixed"
    
    def _evaluate_quality(self, text: str, structure_info: Dict) -> float:
        """評估文本質量（0-1分數）"""
        score = 0.5  # 基礎分數
        
        # 長度合理性（太短或太長都扣分）
        length = len(text)
        if 50 <= length <= 10000:
            score += 0.2
        elif length < 20:
            score -= 0.3
        
        # 結構完整性
        if structure_info['paragraphs'] > 1:
            score += 0.1
        if structure_info['sentences'] > 2:
            score += 0.1
        
        # 內容豐富度
        unique_chars = len(set(text))
        if unique_chars > 20:
            score += 0.1
        
        # 避免重複內容
        lines = text.split('\n')
        unique_lines = len(set(lines))
        if len(lines) > 0:
            uniqueness_ratio = unique_lines / len(lines)
            score += uniqueness_ratio * 0.1
        
        return min(max(score, 0.0), 1.0)
    
    def _determine_strategy(self, text_type: str, structure_info: Dict, quality_score: float) -> str:
        """決定處理策略"""
        if quality_score < 0.3:
            return "low_quality_skip"
        
        if text_type in ["micro_text", "short_text"]:
            return "whole_document"
        elif text_type == "medium_text":
            if structure_info['has_chapters'] or structure_info['paragraphs'] > 5:
                return "paragraph_aware"
            else:
                return "simple_split"
        elif text_type == "long_text":
            if structure_info['has_chapters']:
                return "chapter_aware"
            elif structure_info['has_sections']:
                return "section_aware"
            else:
                return "fine_split"
        elif text_type in ["ultra_long", "mega_text"]:
            return "hierarchical_split"
        else:
            # 回退策略
            return "hierarchical_split"

class OptimizedTextSplitter:
    """🔧 優化版文本分割器"""
    
    def __init__(self):
        self.analyzer = SmartTextAnalyzer()
        self.token_estimator = AdvancedTokenEstimator()
        self.config = SMART_TEXT_CONFIG
        
        # 分割器實例池
        self._splitter_cache = {}
        
        print("🔧 優化版文本分割器初始化完成")
        print(f"   支持 {len(self.config)} 種文本類型")
        print(f"   智能策略選擇")
        print(f"   性能優化")
    
    def get_splitter(self, chunk_size: int, chunk_overlap: int) -> RecursiveCharacterTextSplitter:
        """獲取分割器實例（帶快取）"""
        cache_key = f"{chunk_size}_{chunk_overlap}"
        
        if cache_key not in self._splitter_cache:
            self._splitter_cache[cache_key] = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=[
                    "\n\n第", "\n第",      # 章節標題
                    "\n\n", "\n",          # 段落
                    "。", "！", "？", "；",  # 句子結束
                    "，", "、",            # 短語分隔
                    " ", ""
                ],
                keep_separator=True,
                length_function=len
            )
        
        return self._splitter_cache[cache_key]
    
    def smart_split_documents(self, text: str, doc_id: str, source_info: Dict = None) -> List[Document]:
        """🧠 智能文檔分割主入口"""
        if not text or not text.strip():
            return []
        
        # 分析文本
        analysis = self.analyzer.analyze_text(text, source_info)
        
        print(f"📊 文本分析: {analysis.text_type} ({analysis.length:,} 字符)")
        print(f"   處理策略: {analysis.processing_strategy}")
        print(f"   質量分數: {analysis.quality_score:.2f}")
        print(f"   主要語言: {analysis.language}")
        
        # 根據策略選擇處理方法
        documents = []
        if analysis.processing_strategy == "low_quality_skip":
            print("   跳過處理")
            return []
        elif analysis.processing_strategy == "whole_document":
            documents = self._process_whole_document(text, doc_id, analysis)
        elif analysis.processing_strategy == "paragraph_aware":
            documents = self._process_paragraph_aware(text, doc_id, analysis)
        elif analysis.processing_strategy == "simple_split":
            documents = self._process_simple_split(text, doc_id, analysis)
        elif analysis.processing_strategy == "chapter_aware":
            documents = self._process_chapter_aware(text, doc_id, analysis)
        elif analysis.processing_strategy == "section_aware":
            documents = self._process_section_aware(text, doc_id, analysis)
        elif analysis.processing_strategy == "fine_split":
            documents = self._process_fine_split(text, doc_id, analysis)
        elif analysis.processing_strategy == "hierarchical_split":
            documents = self._process_hierarchical_split(text, doc_id, analysis)
        else:
            # 回退到簡單分割
            documents = self._process_simple_split(text, doc_id, analysis)

        # 最終驗證和過濾，確保沒有空內容的文檔
        final_documents = [doc for doc in documents if doc.page_content and doc.page_content.strip()]
        if len(final_documents) != len(documents):
            logger.warning(f"過濾掉 {len(documents) - len(final_documents)} 個空內容的倫塊")
        
        return final_documents
    
    def _process_whole_document(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """處理整個文檔（不分割）"""
        config = self.config[analysis.text_type]
        
        if len(text.strip()) < config.get("min_length", 20):
            print("   文檔過短，跳過處理")
            return []
        
        print(f"   {config['description']}")
        
        # 計算 token 數
        token_count = self.token_estimator.estimate_tokens(text)
        
        metadata = {
            'doc_id': doc_id,
            'chunk_id': f"{doc_id}_whole",
            'chunk_index': 0,
            'total_chunks': 1,
            'text_type': analysis.text_type,
            'processing_strategy': analysis.processing_strategy,
            'is_complete_document': True,
            'token_count': token_count,
            'quality_score': analysis.quality_score,
            'language': analysis.language,
            'structure_info': analysis.structure_info
        }
        
        return [Document(page_content=text, metadata=metadata)]
    
    def _process_paragraph_aware(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """段落感知處理"""
        config = self.config[analysis.text_type]
        chunk_size = config["chunk_size"]
        chunk_overlap = config["chunk_overlap"]
        
        print(f"   {config['description']} (目標大小: {chunk_size})")
        
        # 按段落分割
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if not paragraphs:
            return self._process_simple_split(text, doc_id, analysis)
        
        documents = []
        current_chunk = ""
        chunk_index = 0
        
        for para_idx, paragraph in enumerate(paragraphs):
            # 檢查是否可以加入當前分塊
            potential_chunk = current_chunk + ("\n\n" if current_chunk else "") + paragraph
            
            if len(potential_chunk) <= chunk_size or not current_chunk:
                current_chunk = potential_chunk
            else:
                # 當前分塊已滿，創建文檔
                if current_chunk:
                    doc = self._create_chunk_document(
                        current_chunk, doc_id, chunk_index, analysis, "paragraph_aware"
                    )
                    documents.append(doc)
                    chunk_index += 1
                
                current_chunk = paragraph
        
        # 處理最後一個分塊
        if current_chunk:
            doc = self._create_chunk_document(
                current_chunk, doc_id, chunk_index, analysis, "paragraph_aware"
            )
            documents.append(doc)
        
        # 處理重疊
        if chunk_overlap > 0 and len(documents) > 1:
            documents = self._add_overlap_to_documents(documents, chunk_overlap)
        
        return documents
    
    def _process_simple_split(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """簡單分割處理"""
        config = self.config[analysis.text_type]
        chunk_size = config["chunk_size"]
        chunk_overlap = config["chunk_overlap"]
        
        print(f"   {config['description']} (大小: {chunk_size}, 重疊: {chunk_overlap})")
        
        splitter = self.get_splitter(chunk_size, chunk_overlap)
        chunks = splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            if len(chunk.strip()) >= config.get("min_length", 20):
                doc = self._create_chunk_document(
                    chunk, doc_id, i, analysis, "simple_split"
                )
                documents.append(doc)
        
        return documents
    
    def _process_chapter_aware(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """章節感知處理"""
        config = self.config[analysis.text_type]
        
        print(f"   {config['description']} - 檢測到 {analysis.structure_info.get('chapter_count', 0)} 個章節")
        
        # 按章節分割
        chapters = self._split_by_chapters(text)
        
        if len(chapters) <= 1:
            # 沒有檢測到章節，回退到段落感知
            return self._process_paragraph_aware(text, doc_id, analysis)
        
        all_documents = []
        
        for chapter_idx, chapter_text in enumerate(chapters):
            chapter_doc_id = f"{doc_id}_ch{chapter_idx+1:02d}"
            
            # 每個章節內部使用段落感知分割
            chapter_analysis = analysis  # 使用相同的分析結果
            chapter_docs = self._process_paragraph_aware(chapter_text, chapter_doc_id, chapter_analysis)
            
            # 添加章節信息到元數據
            for doc in chapter_docs:
                doc.metadata.update({
                    'chapter_index': chapter_idx,
                    'chapter_total': len(chapters),
                    'parent_doc_id': doc_id,
                    'split_method': 'chapter_aware'
                })
            
            all_documents.extend(chapter_docs)
        
        return all_documents
    
    def _process_section_aware(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """小節感知處理"""
        config = self.config[analysis.text_type]
        
        print(f"   {config['description']} - 檢測到 {analysis.structure_info.get('section_count', 0)} 個小節")
        
        # 按小節分割
        sections = self._split_by_sections(text)
        
        if len(sections) <= 1:
            return self._process_simple_split(text, doc_id, analysis)
        
        all_documents = []
        
        for section_idx, section_text in enumerate(sections):
            section_doc_id = f"{doc_id}_sec{section_idx+1:02d}"
            
            # 小節內部簡單分割
            section_docs = self._process_simple_split(section_text, section_doc_id, analysis)
            
            # 添加小節信息
            for doc in section_docs:
                doc.metadata.update({
                    'section_index': section_idx,
                    'section_total': len(sections),
                    'parent_doc_id': doc_id,
                    'split_method': 'section_aware'
                })
            
            all_documents.extend(section_docs)
        
        return all_documents
    
    def _process_fine_split(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """精細分割處理"""
        config = self.config[analysis.text_type]
        
        print(f"   {config['description']}")
        
        return self._process_simple_split(text, doc_id, analysis)
    
    def _process_hierarchical_split(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """階層式分割處理"""
        config = self.config[analysis.text_type]
        
        print(f"   {config['description']}")
        
        # 第一層：章節分割
        chapters = self._split_by_chapters(text)
        
        if len(chapters) > 1:
            return self._process_chapter_aware(text, doc_id, analysis)
        
        # 第二層：段落分割
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if len(paragraphs) > 10:  # 段落較多，用段落感知
            return self._process_paragraph_aware(text, doc_id, analysis)
        
        # 第三層：精細分割
        return self._process_fine_split(text, doc_id, analysis)
    
    def _split_by_chapters(self, text: str) -> List[str]:
        """按章節分割文本"""
        chapter_patterns = [
            r'\n(第[一二三四五六七八九十\d]+章[^\n]*)',
            r'\n(Chapter\s+\d+[^\n]*)',
            r'\n(\d+\.\d+[^\n]*)',
            r'\n([A-Z][^\n]{10,50})\n(?=[A-Z]|$)'
        ]
        
        best_split = None
        best_count = 0
        
        for pattern in chapter_patterns:
            matches = list(re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE))
            
            if len(matches) >= 2:
                splits = []
                start = 0
                
                for match in matches:
                    if start < match.start():
                        splits.append(text[start:match.start()].strip())
                    start = match.start()
                
                splits.append(text[start:].strip())
                splits = [s for s in splits if s and len(s) > 50]
                
                if len(splits) > best_count:
                    best_split = splits
                    best_count = len(splits)
        
        return best_split if best_split else [text]
    
    def _split_by_sections(self, text: str) -> List[str]:
        """按小節分割文本"""
        section_patterns = [
            r'\n(第[一二三四五六七八九十\d]+節[^\n]*)',
            r'\n(\d+\.\d+[^\n]*)',
            r'\n([一二三四五六七八九十]、[^\n]*)',
            r'\n([A-Z]\.[^\n]*)'
        ]
        
        for pattern in section_patterns:
            matches = list(re.finditer(pattern, text, re.MULTILINE))
            
            if len(matches) >= 2:
                splits = []
                start = 0
                
                for match in matches:
                    if start < match.start():
                        splits.append(text[start:match.start()].strip())
                    start = match.start()
                
                splits.append(text[start:].strip())
                return [s for s in splits if s and len(s) > 30]
        
        return [text]
    
    def _add_overlap_to_documents(self, documents: List[Document], overlap_size: int) -> List[Document]:
        """為文檔添加重疊部分"""
        if len(documents) <= 1:
            return documents
        
        for i in range(1, len(documents)):
            # 從前一個文檔末尾取重疊內容
            prev_content = documents[i-1].page_content
            current_content = documents[i].page_content
            
            if len(prev_content) > overlap_size:
                overlap_text = prev_content[-overlap_size:]
                # 尋找完整的句子邊界
                sentence_end = overlap_text.rfind('。')
                if sentence_end == -1:
                    sentence_end = overlap_text.rfind('.')
                
                if sentence_end > overlap_size // 2:
                    overlap_text = overlap_text[sentence_end+1:]
                
                documents[i].page_content = overlap_text + "\n" + current_content
                documents[i].metadata['has_overlap'] = True
                documents[i].metadata['overlap_length'] = len(overlap_text)
        
        return documents
    
    def _create_chunk_document(self, content: str, doc_id: str, chunk_index: int, 
                      analysis: TextAnalysis, split_method: str) -> Document:
        """創建分塊文档 - 修復版：完整URL和標題提取"""
        token_count = self.token_estimator.estimate_tokens(content)
        
        # 完整的URL和標題提取#!/usr/bin/env python3
"""
向量操作核心層 - VectorOperationsCore
職責：底層數據操作、向量處理、文檔載入、集合管理
包含26個底層核心方法，從原 OptimizedVectorSystem 精確移動而來
"""

import time
import json
import hashlib
import re
import logging
import os
import gc
import threading
import tempfile
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Set, Any, Union
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, asdict

# LangChain核心
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    TextLoader, PyPDFLoader, Docx2txtLoader, 
    CSVLoader, JSONLoader, PandasExcelLoader
)

# --- Start of content from core_config.py ---

# 🆕 1. 添加 Railway 環境檢測函數（放在文件頂部，導入語句之後）
def detect_railway_environment():
    """檢測是否在 Railway 環境中運行"""
    railway_indicators = [
        os.getenv("RAILWAY_PROJECT_ID"),
        os.getenv("RAILWAY_SERVICE_ID"), 
        os.getenv("DATABASE_URL"),
        "railway.internal" in os.getenv("POSTGRES_HOST", "")
    ]
    
    is_railway = any(railway_indicators)
    if is_railway:
        print("🚂 檢測到 Railway 部署環境")
        # 顯示 Railway 特定信息
        project_id = os.getenv("RAILWAY_PROJECT_ID", "unknown")
        service_id = os.getenv("RAILWAY_SERVICE_ID", "unknown")  
        print(f"   項目ID: {project_id[:8]}***")
        print(f"   服務ID: {service_id[:8]}***")
    
    return is_railway

# 🔧 優化版系統配置
SYSTEM_CONFIG = {
    "persist_dir": "./chroma_langchain_db",
    "data_dir": "./data", 
    "device": "cpu",  
    "log_level": "INFO",
    "max_workers": 4,  # 並發處理數
    "cache_embeddings": True,  # 快取嵌入向量
    "backup_enabled": True  # 自動備份
}

# 🧠 智能文本分類配置
SMART_TEXT_CONFIG = {
    # 文本長度閾值（字符數）
    "micro_text_threshold": 100,      # 微文本：標題、摘要等
    "short_text_threshold": 500,      # 短文本：短文章、段落
    "medium_text_threshold": 2000,    # 中等文本：中篇文章
    "long_text_threshold": 8000,    # 長文本：長篇文章
    "ultra_long_threshold": 20000,    # 超長文本：書籍章節
    
    # 各類文本的處理策略
    "micro_text": {
        "min_length": 20,
        "process_whole": True,
        "merge_with_next": True,  # 嘗試與下一個片段合併
        "chunk_size": 0,
        "description": "微文本整體處理"
    },
    
    "short_text": {
        "min_length": 50,
        "process_whole": True,
        "preserve_structure": True,
        "chunk_size": 0,
        "allow_merge": True,  # 允許合併相鄰短文本
        "description": "短文本整體保存"
    },
    
    "medium_text": {
        "chunk_size": 800,
        "chunk_overlap": 120,
        "preserve_paragraphs": True,
        "smart_boundaries": True,  # 智能邊界檢測
        "min_chunk_ratio": 0.3,   # 最小片段比例
        "description": "中等文本段落感知分割"
    },
    
    "long_text": {
        "chunk_size": 500,
        "chunk_overlap": 80,
        "hierarchical_split": True,
        "section_aware": True,    # 章節感知
        "quality_check": True,    # 質量檢查
        "description": "長文本精細化分割"
    },
    
    "ultra_long": {
        "chunk_size": 600,
        "chunk_overlap": 100,
        "multi_level_split": True,  # 多層級分割
        "chapter_detection": True,  # 章節檢測
        "summary_chunks": True,     # 生成摘要片段
        "description": "超長文本階層式處理"
    },
    
    "mega_text": {
        "chunk_size": 500,
        "chunk_overlap": 80,
        "multi_level_split": True,  # 多層級分割
        "chapter_detection": True,  # 章節檢測
        "aggressive_split": True,   # 積極分割
        "quality_filter": True,     # 質量過濾
        "description": "超大文本積極分割處理"
    }
}

# 🔧 優化版 Token 限制配置
TOKEN_LIMITS = {
    "max_tokens_per_request": 150000,  # 更保守的限制
    "max_batch_size": 8,               # 減小批次大小
    "min_batch_size": 1,
    "batch_delay": 5.0,                # 增加批次間延遲
    "retry_delay": 10,
    "max_retries": 3,
    "token_safety_margin": 0.2,        # 20% 安全邊際
    "adaptive_batching": True           # 自適應批次大小
}

# 🔧 性能優化配置
PERFORMANCE_CONFIG = {
    "embedding_batch_size": 12,
    "parallel_processing": True,
    "memory_limit_mb": 1024,      # 記憶體限制
    "chunk_cache_size": 1000,     # 分塊快取大小
    "preload_models": True,       # 預載入模型
    "gc_frequency": 50            # 垃圾回收頻率
}

# 支持的文件格式擴展
SUPPORTED_EXTENSIONS = {
    '.txt', '.md', '.pdf', '.csv', '.json', '.py', '.js', 
    '.docx', '.doc', '.epub', '.rst', '.org', '.xlsx', '.xls'
}

from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from langchain_community.vectorstores import Chroma

try:
    from langchain_community.vectorstores import Chroma
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False
    # 創建一個虛擬的 Chroma 類型用於類型註解
    class Chroma:
        pass

# 🔧 檢查環境變數
def check_openai_api_key():
    """檢查 OpenAI API Key"""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("⚠️ 未設置 OPENAI_API_KEY 環境變數")
        print("   請在 .env 文件中設置: OPENAI_API_KEY=sk-your-api-key")
        return False
    
    if api_key.startswith("sk-") and len(api_key) > 20:
        print("✅ OpenAI API Key 格式正確")
        print(f"   API Key: {api_key[:8]}***{api_key[-4:]}")
        return True
    else:
        print("⚠️ OpenAI API Key 格式可能不正確")
        return False

# LangChain 核心組件導入
try:
    from langchain_postgres import PGVector
    PGVECTOR_AVAILABLE = True
    print("✅ PGVector 可用")
except ImportError:
    try:
        from langchain_community.vectorstores import PGVector
        PGVECTOR_AVAILABLE = True
        print("✅ PGVector (community) 可用")
    except ImportError:
        PGVECTOR_AVAILABLE = False
        print("❌ PGVector 不可用")
        # 可以回退到 Chroma
        from langchain_community.vectorstores import Chroma

# OpenAI embeddings 導入
try:
    from langchain_openai import OpenAIEmbeddings
    OPENAI_EMBEDDINGS_AVAILABLE = True
    print("✅ OpenAI Embeddings 可用")
    
    if check_openai_api_key():
        print("✅ OpenAI 環境檢查通過")
    else:
        print("⚠️ OpenAI 環境配置可能有問題")
        
except ImportError as e:
    OPENAI_EMBEDDINGS_AVAILABLE = False
    print(f"⚠️ OpenAI Embeddings 不可用: {e}")

# 🔧 PostgreSQL 依賴檢查
try:
    import psycopg2
    PSYCOPG2_AVAILABLE = True
except ImportError:
    PSYCOPG2_AVAILABLE = False
    print("⚠️ 警告: psycopg2 未安裝，請執行: pip install psycopg2-binary")

# 中文分詞和轉換
try:
    import jieba
    jieba.initialize()
    JIEBA_AVAILABLE = True
except ImportError:
    JIEBA_AVAILABLE = False
    print("⚠️ 警告: jieba 未安裝，中文分詞功能將被禁用")

try:
    import opencc
    OPENCC_AVAILABLE = True
except ImportError:
    OPENCC_AVAILABLE = False
    print("⚠️ 警告: opencc 未安裝，繁簡轉換功能將被禁用")

# 文檔處理支持
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
    DOCX_METHOD = "docx2txt"
    print("✅ DOCX 支持已啟用 (docx2txt)")
except ImportError:
    try:
        from docx import Document as DocxDocument
        DOCX2TXT_AVAILABLE = True
        DOCX_METHOD = "python-docx"
        print("✅ DOCX 支持已啟用 (python-docx)")
    except ImportError:
        DOCX2TXT_AVAILABLE = False
        DOCX_METHOD = None
        print("⚠️ 警告: docx2txt 或 python-docx 未安裝")

try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    EPUB_AVAILABLE = True
    print("✅ EPUB 支持已啟用")
except ImportError:
    EPUB_AVAILABLE = False
    print("⚠️ 警告: ebooklib 或 beautifulsoup4 未安裝")

@dataclass
class FileInfo:
    """文件資訊"""
    path: str
    size: int
    mtime: float
    hash: str
    encoding: str = "utf-8"
    file_type: str = ""

@dataclass
class TextAnalysis:
    """文本分析結果"""
    length: int
    text_type: str
    language: str
    encoding: str
    structure_info: Dict
    quality_score: float
    processing_strategy: str

@dataclass
class ChunkInfo:
    """分塊資訊"""
    chunk_id: str
    content: str
    metadata: Dict
    token_count: int
    quality_score: float
    relationships: List[str]  # 與其他分塊的關係

@dataclass
class SearchResult:
    content: str
    score: float
    metadata: Dict
    collection: str
    chunk_info: Optional[ChunkInfo] = None

# --- End of content from core_config.py ---


# --- Start of content from text_processing.py ---

logger = logging.getLogger(__name__)

class AdvancedTokenEstimator:
    """🔧 高級 Token 估算器"""
    
    def __init__(self):
        # 不同語言和內容類型的 Token 係數
        self.token_ratios = {
            "chinese": 2.5,      # 中文字符/token
            "english": 4.0,      # 英文字符/token
            "mixed": 3.0,        # 中英混合
            "code": 3.5,         # 程式碼
            "punctuation": 1.0,  # 標點符號
            "numbers": 2.0       # 數字
        }
        self.safety_margin = TOKEN_LIMITS.get("token_safety_margin", 0.15)
    
    def analyze_text_composition(self, text: str) -> Dict[str, int]:
        """分析文本組成"""
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_chars = len(re.findall(r'[a-zA-Z]', text))
        numbers = len(re.findall(r'\d', text))
        punctuation = len(re.findall(r'[^\w\s\u4e00-\u9fff]', text))
        other_chars = len(text) - chinese_chars - english_chars - numbers - punctuation
        
        return {
            "chinese": chinese_chars,
            "english": english_chars,
            "numbers": numbers,
            "punctuation": punctuation,
            "other": other_chars,
            "total": len(text)
        }
    
    def estimate_tokens(self, text: str, content_type: str = "mixed") -> int:
        """精確估算 Token 數量"""
        if not text:
            return 0
        
        composition = self.analyze_text_composition(text)
        
        # 根據文本組成計算 token
        estimated_tokens = (
            composition["chinese"] / self.token_ratios["chinese"] +
            composition["english"] / self.token_ratios["english"] +
            composition["numbers"] / self.token_ratios["numbers"] +
            composition["punctuation"] / self.token_ratios["punctuation"] +
            composition["other"] / self.token_ratios["mixed"]
        )
        
        # 根據內容類型調整
        if content_type == "code":
            estimated_tokens *= 1.2  # 程式碼通常 token 密度更高
        elif content_type == "academic":
            estimated_tokens *= 1.1  # 學術文本專業詞彙多
        
        # 加上安全邊際
        final_estimate = int(estimated_tokens * (1 + self.safety_margin))
        
        return max(final_estimate, 1)  # 至少1個token
    
    def estimate_embedding_cost(self, total_tokens: int, model: str = "text-embedding-3-small") -> float:
        """估算 Embedding 成本"""
        cost_per_1k_tokens = {
            "text-embedding-3-small": 0.00002,
            "text-embedding-3-large": 0.00013,
            "text-embedding-ada-002": 0.0001
        }
        
        rate = cost_per_1k_tokens.get(model, 0.00002)
        return (total_tokens / 1000) * rate

class ChineseTextNormalizer:
    """優化版中文文本標準化處理器"""
    
    def __init__(self):
        self.s2t_converter = None
        self.t2s_converter = None
        
        if OPENCC_AVAILABLE:
            try:
                self.s2t_converter = opencc.OpenCC('s2t')
                self.t2s_converter = opencc.OpenCC('t2s')
                print("🔤 繁簡轉換器初始化完成")
            except Exception as e:
                logger.warning(f"OpenCC 初始化失敗: {e}")
        
        # 中文文本正規化規則
        self.normalization_rules = [
            (r'[\u3000]+', ' '),
            (r'\r\n|\r', '\n'),                       # 統一換行
            (r'\n{3,}', '\n\n'),                      # 限制連續換行
            (r'[\u201C\u201D\u2018\u2019\u201E\u201A\u2033\u2032]', '"'),  # 統一引號
            (r'[——–∶]', '-'),                          # 統一破折號
            (r'[…⋯]', '...'),                         # 統一省略號
        ]
    
    def normalize_text(self, text: str) -> Tuple[str, Dict]:
        """標準化文本並返回處理資訊"""
        if not text:
            return "", {}
        
        original_length = len(text)
        processed_text = text
        
        # 應用正規化規則
        for pattern, replacement in self.normalization_rules:
            processed_text = re.sub(pattern, replacement, processed_text)
        
        # 移除首尾空白
        processed_text = processed_text.strip()
        
        # 檢測和轉換繁簡
        variant, confidence = self.detect_chinese_variant(processed_text)
        if variant == 'simplified' and self.s2t_converter:
            try:
                processed_text = self.s2t_converter.convert(processed_text)
                variant = 'traditional_converted'
            except Exception as e:
                logger.warning(f"繁簡轉換失敗: {e}")
        
        processing_info = {
            'original_length': original_length,
            'processed_length': len(processed_text),
            'variant': variant,
            'confidence': confidence,
            'normalized': original_length != len(processed_text)
        }
        
        return processed_text, processing_info
    
    def detect_chinese_variant(self, text: str) -> Tuple[str, float]:
        """檢測繁體或簡體中文"""
        if not text:
            return 'unknown', 0.0
        
        simplified_chars = set('国发会学习论问题业专长时间经济')
        traditional_chars = set('國發會學習論問題業專長時間經濟')
        
        simplified_count = sum(1 for char in text if char in simplified_chars)
        traditional_count = sum(1 for char in text if char in traditional_chars)
        total_chinese = len(re.findall(r'[\u4e00-\u9fff]', text))
        
        if total_chinese == 0:
            return 'unknown', 0.0
        
        simplified_ratio = simplified_count / total_chinese
        traditional_ratio = traditional_count / total_chinese
        
        if simplified_ratio > traditional_ratio:
            return 'simplified', simplified_ratio
        elif traditional_ratio > simplified_ratio:
            return 'traditional', traditional_ratio
        else:
            return 'mixed', max(simplified_ratio, traditional_ratio)
    
    def create_search_variants(self, query: str) -> List[str]:
        """創建繁簡搜索變體"""
        variants = [query]
        
        if not self.s2t_converter or not self.t2s_converter:
            return variants
        
        try:
            traditional = self.s2t_converter.convert(query)
            if traditional != query:
                variants.append(traditional)
            
            simplified = self.t2s_converter.convert(query)
            if simplified != query:
                variants.append(simplified)
        except Exception as e:
            logger.warning(f"創建搜索變體失敗: {e}")
        
        return list(set(variants))

class EpubProcessor:
    """📚 EPUB 文件處理器"""
    
    def __init__(self):
        self.normalizer = ChineseTextNormalizer()
        
    def extract_epub_content(self, file_path: Path) -> str:
        """提取 EPUB 內容"""
        try:
            if not EPUB_AVAILABLE:
                raise ImportError("EPUB 處理庫未安裝")
            
            print(f"📚 正在處理 EPUB 文件: {file_path.name}")
            
            # 讀取 EPUB 文件
            book = epub.read_epub(str(file_path))
            
            # 提取所有文本內容
            content_parts = []
            chapter_count = 0
            
            # 獲取書籍信息
            title = book.get_metadata('DC', 'title')
            author = book.get_metadata('DC', 'creator')
            
            book_info = []
            if title:
                book_info.append(f"書名: {title[0][0] if title else '未知'}")
            if author:
                book_info.append(f"作者: {author[0][0] if author else '未知'}")
            
            if book_info:
                content_parts.append("\n".join(book_info) + "\n\n")
            
            # 按順序處理章節
            spine_items = book.get_items_of_type(ebooklib.ITEM_DOCUMENT)
            
            for item in spine_items:
                try:
                    # 解析 HTML 內容
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    
                    # 移除腳本和樣式
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    # 提取文本
                    text = soup.get_text()
                    
                    if text and len(text.strip()) > 50:  # 過濾太短的內容
                        # 清理文本
                        cleaned_text = self._clean_epub_text(text)
                        
                        if cleaned_text.strip():
                            chapter_count += 1
                            
                            # 添加章節標記
                            chapter_title = self._extract_chapter_title(soup, cleaned_text)
                            if chapter_title:
                                content_parts.append(f"\n\n=== {chapter_title} ===\n")
                            else:
                                content_parts.append(f"\n\n=== 第 {chapter_count} 章 ===\n")
                            
                            content_parts.append(cleaned_text)
                            
                except Exception as e:
                    print(f"   ⚠️ 章節處理失敗: {e}")
                    continue
            
            if not content_parts:
                raise ValueError("EPUB 文件中沒有找到有效內容")
            
            full_content = "".join(content_parts)
            
            print(f"   ✅ EPUB 處理完成: {chapter_count} 個章節, {len(full_content):,} 字符")
            
            return full_content
            
        except Exception as e:
            print(f"   ❌ EPUB 處理失敗: {e}")
            raise
    
    def _clean_epub_text(self, text: str) -> str:
        """清理 EPUB 文本"""
        if not text:
            return ""
        
        # 標準化文本
        normalized_text, _ = self.normalizer.normalize_text(text)
        
        # EPUB 特定清理
        epub_cleaning_rules = [
            (r'\n{4,}', '\n\n'),                    # 限制連續換行
            (r'[ \t]{3,}', ' ' ),                    # 限制連續空格
            (r'^[ \t]+', '', re.MULTILINE),          # 移除行首空白
            (r'[ \t]+$', '', re.MULTILINE),          # 移除行尾空白
            (r'\n[ \t]*\n', '\n\n'),                # 清理空行
            (r'[\x00-\x08\x0B\x0C\x0E-\x1F]', ''),  # 移除控制字符
        ]
        
        cleaned_text = normalized_text
        for pattern, replacement, *flags in epub_cleaning_rules:
            flag = flags[0] if flags else 0
            cleaned_text = re.sub(pattern, replacement, cleaned_text, flags=flag)
        
        return cleaned_text.strip()
    
    def _extract_chapter_title(self, soup, text: str) -> Optional[str]:
        """嘗試提取章節標題"""
        try:
            # 從 HTML 標籤中查找標題
            for tag in ['h1', 'h2', 'h3', 'title']:
                title_element = soup.find(tag)
                if title_element and title_element.get_text().strip():
                    title = title_element.get_text().strip()
                    if 5 <= len(title) <= 100:  # 合理的標題長度
                        return title
            
            # 從文本開頭查找標題
            lines = text.split('\n')[:5]  # 檢查前5行
            for line in lines:
                line = line.strip()
                if line and 5 <= len(line) <= 100:
                    # 檢查是否包含章節關鍵詞
                    chapter_keywords = ['章', 'Chapter', '第', '卷', 'Part']
                    if any(keyword in line for keyword in chapter_keywords):
                        return line
                    # 或者是短行且看起來像標題
                    elif len(line) <= 50 and not line.endswith(('。', '！', '？', '.', '!', '?')):
                        return line
            
            return None
            
        except Exception:
            return None

class SmartTextAnalyzer:
    """🧠 智能文本分析器"""
    
    def __init__(self):
        self.normalizer = ChineseTextNormalizer()
        self.token_estimator = AdvancedTokenEstimator()
        
    def analyze_text(self, text: str, source_info: Dict = None) -> TextAnalysis:
        """綜合分析文本"""
        if not text:
            return TextAnalysis(
                length=0, text_type="empty", language="unknown",
                encoding="utf-8", structure_info={}, quality_score=0.0,
                processing_strategy="skip"
            )
        
        # 標準化文本
        normalized_text, norm_info = self.normalizer.normalize_text(text)
        length = len(normalized_text)
        
        # 分類文本長度
        text_type = self._classify_text_length(length)
        
        # 分析文本結構
        structure_info = self._analyze_structure(normalized_text)
        
        # 檢測語言
        language = self._detect_language(normalized_text)
        
        # 評估質量
        quality_score = self._evaluate_quality(normalized_text, structure_info)
        
        # 決定處理策略
        processing_strategy = self._determine_strategy(text_type, structure_info, quality_score)
        
        return TextAnalysis(
            length=length,
            text_type=text_type,
            language=language,
            encoding=norm_info.get('encoding', 'utf-8'),
            structure_info=structure_info,
            quality_score=quality_score,
            processing_strategy=processing_strategy
        )
    
    def _classify_text_length(self, length: int) -> str:
        """分類文本長度"""
        config = SMART_TEXT_CONFIG
        
        if length < config["micro_text_threshold"]:
            return "micro_text"
        elif length < config["short_text_threshold"]:
            return "short_text"
        elif length < config["medium_text_threshold"]:
            return "medium_text"
        elif length < config["long_text_threshold"]:
            return "long_text"
        elif length < config["ultra_long_threshold"]:
            return "ultra_long"
        else:
            return "mega_text"  # 超大文本
    
    def _analyze_structure(self, text: str) -> Dict:
        """分析文本結構"""
        structure_info = {
            'paragraphs': len(text.split('\n\n')),
            'lines': len(text.split('\n')),
            'sentences': len(re.findall(r'[。！？.!?]+', text)),
            'has_chapters': False,
            'has_sections': False,
            'has_lists': False,
            'has_tables': False,
            'chapter_count': 0,
            'section_count': 0
        }
        
        # 檢測章節結構
        chapter_patterns = [
            r'第[一二三四五六七八九十\d]+章',
            r'Chapter\s+\d+',
            r'\d+\.\s*[^\n]{1,50}\n'
        ]
        
        for pattern in chapter_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                structure_info['has_chapters'] = True
                structure_info['chapter_count'] = len(matches)
                break
        
        # 檢測小節結構
        section_patterns = [
            r'第[一二三四五六七八九十\d]+節',
            r'\d+\.\d+',
            r'[一二三四五六七八九十]、'
        ]
        
        for pattern in section_patterns:
            matches = re.findall(pattern, text)
            if matches:
                structure_info['has_sections'] = True
                structure_info['section_count'] = len(matches)
                break
        
        # 檢測列表
        if re.search(r'^\s*[•\-*]\s+', text, re.MULTILINE) or \
           re.search(r'^\s*\d+\.\s+', text, re.MULTILINE):
            structure_info['has_lists'] = True
        
        # 檢測表格
        if '|' in text and text.count('|') > 4:
            structure_info['has_tables'] = True
        
        return structure_info
    
    def _detect_language(self, text: str) -> str:
        """檢測文本主要語言"""
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_chars = len(re.findall(r'[a-zA-Z]', text))
        total_chars = len(text)
        
        if total_chars == 0:
            return "unknown"
        
        chinese_ratio = chinese_chars / total_chars
        english_ratio = english_chars / total_chars
        
        if chinese_ratio > 0.3:
            if english_ratio > 0.2:
                return "mixed_zh_en"
            else:
                return "chinese"
        elif english_ratio > 0.5:
            return "english"
        else:
            return "mixed"
    
    def _evaluate_quality(self, text: str, structure_info: Dict) -> float:
        """評估文本質量（0-1分數）"""
        score = 0.5  # 基礎分數
        
        # 長度合理性（太短或太長都扣分）
        length = len(text)
        if 50 <= length <= 10000:
            score += 0.2
        elif length < 20:
            score -= 0.3
        
        # 結構完整性
        if structure_info['paragraphs'] > 1:
            score += 0.1
        if structure_info['sentences'] > 2:
            score += 0.1
        
        # 內容豐富度
        unique_chars = len(set(text))
        if unique_chars > 20:
            score += 0.1
        
        # 避免重複內容
        lines = text.split('\n')
        unique_lines = len(set(lines))
        if len(lines) > 0:
            uniqueness_ratio = unique_lines / len(lines)
            score += uniqueness_ratio * 0.1
        
        return min(max(score, 0.0), 1.0)
    
    def _determine_strategy(self, text_type: str, structure_info: Dict, quality_score: float) -> str:
        """決定處理策略"""
        if quality_score < 0.3:
            return "low_quality_skip"
        
        if text_type in ["micro_text", "short_text"]:
            return "whole_document"
        elif text_type == "medium_text":
            if structure_info['has_chapters'] or structure_info['paragraphs'] > 5:
                return "paragraph_aware"
            else:
                return "simple_split"
        elif text_type == "long_text":
            if structure_info['has_chapters']:
                return "chapter_aware"
            elif structure_info['has_sections']:
                return "section_aware"
            else:
                return "fine_split"
        elif text_type in ["ultra_long", "mega_text"]:
            return "hierarchical_split"
        else:
            # 回退策略
            return "hierarchical_split"

class OptimizedTextSplitter:
    """🔧 優化版文本分割器"""
    
    def __init__(self):
        self.analyzer = SmartTextAnalyzer()
        self.token_estimator = AdvancedTokenEstimator()
        self.config = SMART_TEXT_CONFIG
        
        # 分割器實例池
        self._splitter_cache = {}
        
        print("🔧 優化版文本分割器初始化完成")
        print(f"   支持 {len(self.config)} 種文本類型")
        print(f"   智能策略選擇")
        print(f"   性能優化")
    
    def get_splitter(self, chunk_size: int, chunk_overlap: int) -> RecursiveCharacterTextSplitter:
        """獲取分割器實例（帶快取）"""
        cache_key = f"{chunk_size}_{chunk_overlap}"
        
        if cache_key not in self._splitter_cache:
            self._splitter_cache[cache_key] = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=[
                    "\n\n第", "\n第",      # 章節標題
                    "\n\n", "\n",          # 段落
                    "。", "！", "？", "；",  # 句子結束
                    "，", "、",            # 短語分隔
                    " ", ""
                ],
                keep_separator=True,
                length_function=len
            )
        
        return self._splitter_cache[cache_key]
    
    def smart_split_documents(self, text: str, doc_id: str, source_info: Dict = None) -> List[Document]:
        """🧠 智能文檔分割主入口"""
        if not text or not text.strip():
            return []
        
        # 分析文本
        analysis = self.analyzer.analyze_text(text, source_info)
        
        print(f"📊 文本分析: {analysis.text_type} ({analysis.length:,} 字符)")
        print(f"   處理策略: {analysis.processing_strategy}")
        print(f"   質量分數: {analysis.quality_score:.2f}")
        print(f"   主要語言: {analysis.language}")
        
        # 根據策略選擇處理方法
        documents = []
        if analysis.processing_strategy == "low_quality_skip":
            print("   跳過處理")
            return []
        elif analysis.processing_strategy == "whole_document":
            documents = self._process_whole_document(text, doc_id, analysis)
        elif analysis.processing_strategy == "paragraph_aware":
            documents = self._process_paragraph_aware(text, doc_id, analysis)
        elif analysis.processing_strategy == "simple_split":
            documents = self._process_simple_split(text, doc_id, analysis)
        elif analysis.processing_strategy == "chapter_aware":
            documents = self._process_chapter_aware(text, doc_id, analysis)
        elif analysis.processing_strategy == "section_aware":
            documents = self._process_section_aware(text, doc_id, analysis)
        elif analysis.processing_strategy == "fine_split":
            documents = self._process_fine_split(text, doc_id, analysis)
        elif analysis.processing_strategy == "hierarchical_split":
            documents = self._process_hierarchical_split(text, doc_id, analysis)
        else:
            # 回退到簡單分割
            documents = self._process_simple_split(text, doc_id, analysis)

        # 最終驗證和過濾，確保沒有空內容的文檔
        final_documents = [doc for doc in documents if doc.page_content and doc.page_content.strip()]
        if len(final_documents) != len(documents):
            logger.warning(f"過濾掉 {len(documents) - len(final_documents)} 個空內容的倫塊")
        
        return final_documents
    
    def _process_whole_document(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """處理整個文檔（不分割）"""
        config = self.config[analysis.text_type]
        
        if len(text.strip()) < config.get("min_length", 20):
            print("   文檔過短，跳過處理")
            return []
        
        print(f"   {config['description']}")
        
        # 計算 token 數
        token_count = self.token_estimator.estimate_tokens(text)
        
        metadata = {
            'doc_id': doc_id,
            'chunk_id': f"{doc_id}_whole",
            'chunk_index': 0,
            'total_chunks': 1,
            'text_type': analysis.text_type,
            'processing_strategy': analysis.processing_strategy,
            'is_complete_document': True,
            'token_count': token_count,
            'quality_score': analysis.quality_score,
            'language': analysis.language,
            'structure_info': analysis.structure_info
        }
        
        return [Document(page_content=text, metadata=metadata)]
    
    def _process_paragraph_aware(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """段落感知處理"""
        config = self.config[analysis.text_type]
        chunk_size = config["chunk_size"]
        chunk_overlap = config["chunk_overlap"]
        
        print(f"   {config['description']} (目標大小: {chunk_size})")
        
        # 按段落分割
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if not paragraphs:
            return self._process_simple_split(text, doc_id, analysis)
        
        documents = []
        current_chunk = ""
        chunk_index = 0
        
        for para_idx, paragraph in enumerate(paragraphs):
            # 檢查是否可以加入當前分塊
            potential_chunk = current_chunk + ("\n\n" if current_chunk else "") + paragraph
            
            if len(potential_chunk) <= chunk_size or not current_chunk:
                current_chunk = potential_chunk
            else:
                # 當前分塊已滿，創建文檔
                if current_chunk:
                    doc = self._create_chunk_document(
                        current_chunk, doc_id, chunk_index, analysis, "paragraph_aware"
                    )
                    documents.append(doc)
                    chunk_index += 1
                
                current_chunk = paragraph
        
        # 處理最後一個分塊
        if current_chunk:
            doc = self._create_chunk_document(
                current_chunk, doc_id, chunk_index, analysis, "paragraph_aware"
            )
            documents.append(doc)
        
        # 處理重疊
        if chunk_overlap > 0 and len(documents) > 1:
            documents = self._add_overlap_to_documents(documents, chunk_overlap)
        
        return documents
    
    def _process_simple_split(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """簡單分割處理"""
        config = self.config[analysis.text_type]
        chunk_size = config["chunk_size"]
        chunk_overlap = config["chunk_overlap"]
        
        print(f"   {config['description']} (大小: {chunk_size}, 重疊: {chunk_overlap})")
        
        splitter = self.get_splitter(chunk_size, chunk_overlap)
        chunks = splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            if len(chunk.strip()) >= config.get("min_length", 20):
                doc = self._create_chunk_document(
                    chunk, doc_id, i, analysis, "simple_split"
                )
                documents.append(doc)
        
        return documents
    
    def _process_chapter_aware(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """章節感知處理"""
        config = self.config[analysis.text_type]
        
        print(f"   {config['description']} - 檢測到 {analysis.structure_info.get('chapter_count', 0)} 個章節")
        
        # 按章節分割
        chapters = self._split_by_chapters(text)
        
        if len(chapters) <= 1:
            # 沒有檢測到章節，回退到段落感知
            return self._process_paragraph_aware(text, doc_id, analysis)
        
        all_documents = []
        
        for chapter_idx, chapter_text in enumerate(chapters):
            chapter_doc_id = f"{doc_id}_ch{chapter_idx+1:02d}"
            
            # 每個章節內部使用段落感知分割
            chapter_analysis = analysis  # 使用相同的分析結果
            chapter_docs = self._process_paragraph_aware(chapter_text, chapter_doc_id, chapter_analysis)
            
            # 添加章節信息到元數據
            for doc in chapter_docs:
                doc.metadata.update({
                    'chapter_index': chapter_idx,
                    'chapter_total': len(chapters),
                    'parent_doc_id': doc_id,
                    'split_method': 'chapter_aware'
                })
            
            all_documents.extend(chapter_docs)
        
        return all_documents
    
    def _process_section_aware(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """小節感知處理"""
        config = self.config[analysis.text_type]
        
        print(f"   {config['description']} - 檢測到 {analysis.structure_info.get('section_count', 0)} 個小節")
        
        # 按小節分割
        sections = self._split_by_sections(text)
        
        if len(sections) <= 1:
            return self._process_simple_split(text, doc_id, analysis)
        
        all_documents = []
        
        for section_idx, section_text in enumerate(sections):
            section_doc_id = f"{doc_id}_sec{section_idx+1:02d}"
            
            # 小節內部簡單分割
            section_docs = self._process_simple_split(section_text, section_doc_id, analysis)
            
            # 添加小節信息
            for doc in section_docs:
                doc.metadata.update({
                    'section_index': section_idx,
                    'section_total': len(sections),
                    'parent_doc_id': doc_id,
                    'split_method': 'section_aware'
                })
            
            all_documents.extend(section_docs)
        
        return all_documents
    
    def _process_fine_split(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """精細分割處理"""
        config = self.config[analysis.text_type]
        
        print(f"   {config['description']}")
        
        return self._process_simple_split(text, doc_id, analysis)
    
    def _process_hierarchical_split(self, text: str, doc_id: str, analysis: TextAnalysis) -> List[Document]:
        """階層式分割處理"""
        config = self.config[analysis.text_type]
        
        print(f"   {config['description']}")
        
        # 第一層：章節分割
        chapters = self._split_by_chapters(text)
        
        if len(chapters) > 1:
            return self._process_chapter_aware(text, doc_id, analysis)
        
        # 第二層：段落分割
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if len(paragraphs) > 10:  # 段落較多，用段落感知
            return self._process_paragraph_aware(text, doc_id, analysis)
        
        # 第三層：精細分割
        return self._process_fine_split(text, doc_id, analysis)
    
    def _split_by_chapters(self, text: str) -> List[str]:
        """按章節分割文本"""
        chapter_patterns = [
            r'\n(第[一二三四五六七八九十\d]+章[^\n]*)',
            r'\n(Chapter\s+\d+[^\n]*)',
            r'\n(\d+\.\d+[^\n]*)',
            r'\n([A-Z][^\n]{10,50})\n(?=[A-Z]|$)'
        ]
        
        best_split = None
        best_count = 0
        
        for pattern in chapter_patterns:
            matches = list(re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE))
            
            if len(matches) >= 2:
                splits = []
                start = 0
                
                for match in matches:
                    if start < match.start():
                        splits.append(text[start:match.start()].strip())
                    start = match.start()
                
                splits.append(text[start:].strip())
                splits = [s for s in splits if s and len(s) > 50]
                
                if len(splits) > best_count:
                    best_split = splits
                    best_count = len(splits)
        
        return best_split if best_split else [text]
    
    def _split_by_sections(self, text: str) -> List[str]:
        """按小節分割文本"""
        section_patterns = [
            r'\n(第[一二三四五六七八九十\d]+節[^\n]*)',
            r'\n(\d+\.\d+[^\n]*)',
            r'\n([一二三四五六七八九十]、[^\n]*)',
            r'\n([A-Z]\.[^\n]*)'
        ]
        
        for pattern in section_patterns:
            matches = list(re.finditer(pattern, text, re.MULTILINE))
            
            if len(matches) >= 2:
                splits = []
                start = 0
                
                for match in matches:
                    if start < match.start():
                        splits.append(text[start:match.start()].strip())
                    start = match.start()
                
                splits.append(text[start:].strip())
                return [s for s in splits if s and len(s) > 30]
        
        return [text]
    
    def _add_overlap_to_documents(self, documents: List[Document], overlap_size: int) -> List[Document]:
        """為文檔添加重疊部分"""
        if len(documents) <= 1:
            return documents
        
        for i in range(1, len(documents)):
            # 從前一個文檔末尾取重疊內容
            prev_content = documents[i-1].page_content
            current_content = documents[i].page_content
            
            if len(prev_content) > overlap_size:
                overlap_text = prev_content[-overlap_size:]
                # 尋找完整的句子邊界
                sentence_end = overlap_text.rfind('。')
                if sentence_end == -1:
                    sentence_end = overlap_text.rfind('.')
                
                if sentence_end > overlap_size // 2:
                    overlap_text = overlap_text[sentence_end+1:]
                
                documents[i].page_content = overlap_text + "\n" + current_content
                documents[i].metadata['has_overlap'] = True
                documents[i].metadata['overlap_length'] = len(overlap_text)
        
        return documents
    
    def _create_chunk_document(self, content: str, doc_id: str, chunk_index: int, 
                          analysis: TextAnalysis, split_method: str) -> Document:
        """創建分塊文檔 - 統一元數據格式"""
        token_count = self.token_estimator.estimate_tokens(content)
        
        # 搜尋URL
        url_regex = r'https?://[^\s\'"<>\[\]]+'  # 修正後的正則表達式，可以匹配更複雜的URL
        found_urls = re.findall(url_regex, content)
        if found_urls:
            logger.info(f"🔍 成功從文本塊中提取到URL: {found_urls}")

        # 基本元數據（確保都是簡單類型）
        metadata = {
            'doc_id': str(doc_id),
            'chunk_id': f"{doc_id}_{chunk_index+1:03d}",
            'chunk_index': int(chunk_index),
            'text_type': str(analysis.text_type),
            'processing_strategy': str(analysis.processing_strategy),
            'split_method': str(split_method),
            'chunk_length': int(len(content)),
            'token_count': int(token_count),
            'quality_score': float(analysis.quality_score),
            'language': str(analysis.language),
            'has_overlap': False
        }

        # 🎯 正確方案：URL列表 → 分隔符字串 (這是唯一可行的方案)
        if found_urls:
            metadata['contained_urls'] = '|'.join(found_urls)  # ✅ 字符串類型
            metadata['url_count'] = len(found_urls)           # ✅ 整數類型
            metadata['has_urls'] = True                       # ✅ 布林類型
        else:
            metadata['contained_urls'] = ''                   # ✅ 空字符串
            metadata['url_count'] = 0                        # ✅ 整數 0
            metadata['has_urls'] = False                     # ✅ 布林 False

        # 🎯 正確方案：複雜結構 → JSON字串
        if analysis.structure_info:
            metadata['structure_info'] = json.dumps(analysis.structure_info, ensure_ascii=False)  # ✅ JSON字符串
            # 同時保留關鍵信息作為簡單字段（便於查詢）
            metadata['has_chapters'] = bool(analysis.structure_info.get('has_chapters', False))    # ✅ 布林
            metadata['has_sections'] = bool(analysis.structure_info.get('has_sections', False))    # ✅ 布林  
            metadata['paragraph_count'] = int(analysis.structure_info.get('paragraphs', 0))       # ✅ 整數
            metadata['sentence_count'] = int(analysis.structure_info.get('sentences', 0))         # ✅ 整數
        else:
            metadata['structure_info'] = '{}'        # ✅ 空JSON字符串
            metadata['has_chapters'] = False         # ✅ 布林
            metadata['has_sections'] = False         # ✅ 布林
            metadata['paragraph_count'] = 0          # ✅ 整數
            metadata['sentence_count'] = 0           # ✅ 整數
        
        return Document(page_content=content, metadata=metadata)

# --- End of content from text_processing.py ---


# --- Start of content from vector_builder.py ---

class AdaptiveBatchProcessor:
    """🔧 自適應批次處理器"""
    
    def __init__(self):
        self.token_estimator = AdvancedTokenEstimator()
        self.max_tokens_per_request = TOKEN_LIMITS["max_tokens_per_request"]
        self.max_batch_size = TOKEN_LIMITS["max_batch_size"]
        self.adaptive_batching = TOKEN_LIMITS.get("adaptive_batching", True)
        
        # 性能統計
        self.batch_stats = {
            'total_batches': 0,
            'total_documents': 0,
            'total_tokens': 0,
            'avg_batch_time': 0,
            'success_rate': 0,
            'adaptive_adjustments': 0
        }
        
        # 自適應參數
        self.current_batch_size = self.max_batch_size
        self.success_streak = 0
        self.failure_streak = 0
    
    def create_smart_batches(self, documents: List[Document]) -> List[Tuple[List[Document], Dict]]:
        """創建智能批次"""
        if not documents:
            return []
        
        print("🔧 創建智能批次...")
        print(f"   文檔數: {len(documents)}")
        print(f"   Token 限制: {self.max_tokens_per_request:,}")
        print(f"   最大批次大小: {self.max_batch_size}")
        print(f"   自適應批次: {'✅' if self.adaptive_batching else '❌'}")
        
        batches = []
        current_batch = []
        current_tokens = 0
        batch_info = {'documents': 0, 'tokens': 0, 'types': defaultdict(int)}
        
        # 按 token 數排序文檔（大的在前面，更容易處理）
        sorted_docs = sorted(
            documents, 
            key=lambda doc: doc.metadata.get('token_count', 
                self.token_estimator.estimate_tokens(doc.page_content)),
            reverse=True
        )
        
        for doc_idx, doc in enumerate(sorted_docs):
            doc_tokens = doc.metadata.get('token_count') or \
                        self.token_estimator.estimate_tokens(doc.page_content)
            
            # 檢查單個文檔是否過大
            if doc_tokens > self.max_tokens_per_request:
                print(f"   文檔 {doc_idx+1} 過大 ({doc_tokens:,} tokens)，需要分割")
                
                # 完成當前批次
                if current_batch:
                    batches.append((current_batch, dict(batch_info)))
                    current_batch = []
                    current_tokens = 0
                    batch_info = {'documents': 0, 'tokens': 0, 'types': defaultdict(int)}
                
                # 分割大文檔
                split_docs = self._split_large_document(doc)
                for split_doc in split_docs:
                    split_tokens = self.token_estimator.estimate_tokens(split_doc.page_content)
                    split_doc.metadata['token_count'] = split_tokens
                    batches.append(([split_doc], {
                        'documents': 1, 
                        'tokens': split_tokens, 
                        'types': {split_doc.metadata.get('text_type', 'unknown'): 1},
                        'is_split': True
                    }))
                continue
            
            # 檢查是否可以加入當前批次
            would_exceed_tokens = current_tokens + doc_tokens > self.max_tokens_per_request
            would_exceed_size = len(current_batch) >= self._get_adaptive_batch_size()
            
            if would_exceed_tokens or would_exceed_size:
                # 完成當前批次
                if current_batch:
                    batches.append((current_batch, dict(batch_info)))
                
                # 開始新批次
                current_batch = [doc]
                current_tokens = doc_tokens
                batch_info = {
                    'documents': 1, 
                    'tokens': doc_tokens, 
                    'types': defaultdict(int)
                }
                batch_info['types'][doc.metadata.get('text_type', 'unknown')] += 1
            else:
                # 加入當前批次
                current_batch.append(doc)
                current_tokens += doc_tokens
                batch_info['documents'] += 1
                batch_info['tokens'] += doc_tokens
                batch_info['types'][doc.metadata.get('text_type', 'unknown')] += 1
        
        # 處理最後一個批次
        if current_batch:
            batches.append((current_batch, dict(batch_info)))
        
        # 統計信息
        total_docs = sum(info['documents'] for _, info in batches)
        total_tokens = sum(info['tokens'] for _, info in batches)
        avg_tokens_per_batch = total_tokens / len(batches) if batches else 0
        
        print(f"✅ 智能批次創建完成:")
        print(f"   總批次數: {len(batches)}")
        print(f"   總文檔數: {total_docs}")
        print(f"   總 tokens: {total_tokens:,}")
        print(f"   平均 tokens/批次: {avg_tokens_per_batch:.0f}")
        print(f"   估算成本: ${self.token_estimator.estimate_embedding_cost(total_tokens):.4f}")
        
        # 更新統計
        self.batch_stats['total_batches'] += len(batches)
        self.batch_stats['total_documents'] += total_docs
        self.batch_stats['total_tokens'] += total_tokens
        
        return batches
    
    def _get_adaptive_batch_size(self) -> int:
        """獲取自適應批次大小"""
        if not self.adaptive_batching:
            return self.max_batch_size
        
        # 根據成功率調整批次大小
        if self.success_streak >= 3:
            # 連續成功，可以增加批次大小
            self.current_batch_size = min(self.max_batch_size, self.current_batch_size + 1)
            self.batch_stats['adaptive_adjustments'] += 1
        elif self.failure_streak >= 2:
            # 連續失敗，減少批次大小
            self.current_batch_size = max(1, self.current_batch_size - 2)
            self.batch_stats['adaptive_adjustments'] += 1
        
        return self.current_batch_size
    
    def _split_large_document(self, doc: Document) -> List[Document]:
        """分割過大的文檔"""
        content = doc.page_content
        max_chars = int(self.max_tokens_per_request * 2.5)  # 估算字符數
        
        # 嘗試按段落分割
        paragraphs = content.split('\n\n')
        split_docs = []
        current_content = ""
        part_index = 0
        
        for para in paragraphs:
            if len(current_content + para) <= max_chars or not current_content:
                current_content += ("\n\n" if current_content else "") + para
            else:
                if current_content.strip():
                    split_doc = Document(
                        page_content=current_content,
                        metadata={
                            **doc.metadata,
                            'chunk_id': f"{doc.metadata.get('chunk_id', 'unknown')}_part{part_index+1}",
                            'is_split_part': True,
                            'part_index': part_index,
                            'split_reason': 'token_limit'
                        }
                    )
                    split_docs.append(split_doc)
                    part_index += 1
                current_content = para
        
        # 處理最後一部分
        if current_content.strip():
            split_doc = Document(
                page_content=current_content,
                metadata={
                    **doc.metadata,
                    'chunk_id': f"{doc.metadata.get('chunk_id', 'unknown')}_part{part_index+1}",
                    'is_split_part': True,
                    'part_index': part_index,
                    'split_reason': 'token_limit'
                }
            )
            split_docs.append(split_doc)
        
        return split_docs or [doc]  # 如果分割失敗，返回原文檔
    
    def record_batch_result(self, success: bool, processing_time: float = 0):
        """記錄批次處理結果"""
        if success:
            self.success_streak += 1
            self.failure_streak = 0
        else:
            self.failure_streak += 1
            self.success_streak = 0
        
        # 更新成功率
        total_attempts = self.batch_stats['total_batches']
        if total_attempts > 0:
            self.batch_stats['success_rate'] = (total_attempts - self.failure_streak) / total_attempts
        
        # 更新平均處理時間
        if processing_time > 0:
            current_avg = self.batch_stats['avg_batch_time']
            self.batch_stats['avg_batch_time'] = (current_avg + processing_time) / 2
    
    def get_performance_stats(self) -> Dict:
        """獲取性能統計"""
        return dict(self.batch_stats)

# --- End of content from vector_builder.py ---


# 向量存儲 (條件導入)
if PGVECTOR_AVAILABLE:
    try:
        from langchain_postgres import PGVector
    except ImportError:
        from langchain_community.vectorstores import PGVector
else:
    from langchain_community.vectorstores import Chroma

# OpenAI (條件導入)
if OPENAI_EMBEDDINGS_AVAILABLE:
    from langchain_openai import OpenAIEmbeddings

# PostgreSQL (條件導入)  
if PSYCOPG2_AVAILABLE:
    import psycopg2

logger = logging.getLogger(__name__)

class VectorOperationsCore:
    """向量操作核心類別 - 負責底層數據操作和向量處理"""
    
    def __init__(self, data_dir: str = None, model_type: str = None):
        """✅ 純PostgreSQL初始化 - 移除file_records依賴"""
        
        # 🔧 1. 基本變數設置
        self.data_dir = Path(data_dir or SYSTEM_CONFIG["data_dir"])
        self.model_type = model_type or "openai"
        self.persist_dir = Path(SYSTEM_CONFIG["persist_dir"])  # Chroma備用

        # ❌ 不再建立本地目錄 (純PostgreSQL方案)
        print("🚀 純PostgreSQL方案：不使用本地data目錄")
        
        # 🔧 2. 資料庫連接設置（但不測試）
        self.db_adapter = None
        self.connection_string = None
        self.use_postgres = False

        database_url = os.getenv("DATABASE_URL")
        if PGVECTOR_AVAILABLE and database_url:
            self.connection_string = database_url
            print("🔍 發現DATABASE_URL，準備測試PostgreSQL連接...")
        else:
            print("⚠️ DATABASE_URL未設置或PGVector不可用，將使用Chroma")

        if not PGVECTOR_AVAILABLE:
            print("⚠️ PGVector依賴未安裝，使用Chroma作為備用")
            self.persist_dir.mkdir(exist_ok=True)

        # ✅ 3. 先初始化Embedding模型（關鍵!)
        self._setup_embedding_model()
        print("✅ Embedding模型初始化完成")

        # ✅ 4. 現在可以測試PostgreSQL連接了（embeddings已存在）
        if PGVECTOR_AVAILABLE and database_url and hasattr(self, 'embeddings'):
            try:
                print("🔍 測試PostgreSQL + PGVector連接...")
                # 測試連接
                PGVector.from_existing_index(
                    collection_name="_test_connection",
                    embedding=self.embeddings,
                    connection=self.connection_string  # ✅ 修正參數名稱
                )
                self.use_postgres = True
                print("✅ PostgreSQL (pgvector) 連接成功")
            except Exception as e:
                print(f"⚠️ PostgreSQL (pgvector) 連接測試失敗: {e}")
                self.use_postgres = False
                print("📄 回退到Chroma本地存儲")
                self.persist_dir.mkdir(exist_ok=True)
        
        if not self.use_postgres:
            print("🔍 使用Chroma作為向量存儲")
            self.persist_dir.mkdir(exist_ok=True)
        
        # 🔧 5. 初始化文本處理組件
        self._setup_text_processing()
        
        # 🔧 6. 初始化處理器
        self.batch_processor = AdaptiveBatchProcessor()
        self.text_splitter = OptimizedTextSplitter()
        
        # 🔧 7. 初始化存儲（移除檔案記錄）
        self._vector_stores = {}
        
        # ✅ 添加file_records初始化
        self.file_records = {}
        
        # ✅ 改為純PostgreSQL初始化
        print("🚀 純PostgreSQL方案：所有檔案數據將直接存儲在PostgreSQL中")
        print("📄 不再維護本地檔案記錄 (file_records.json)")
        
        self.processing_lock = threading.Lock()
        
        print(f"🚀 向量操作核心初始化完成")
        print(f"   嵌入模型: {self.model_type}")
        print(f"   資料目錄: 不使用 (純PostgreSQL)")
        print(f"   向量庫: {'PostgreSQL + PGVector' if self.use_postgres else 'Chroma (本地)'}")
        print(f"   智能文本處理: ✅")
        print(f"   自適應批次: ✅")
        print(f"   純PostgreSQL方案: {'✅' if self.use_postgres else '❌'}")

    def _setup_embedding_model(self):
        """設定嵌入模型"""
        try:
            if self.model_type == "openai":
                if not OPENAI_EMBEDDINGS_AVAILABLE:
                    raise ImportError("OpenAI Embeddings不可用")
                
                print(f"🔧 初始化OpenAI Embeddings...")
                
                api_key = os.getenv("OPENAI_API_KEY")
                base_url = os.getenv("OPENAI_API_BASE")
                
                embedding_params = {
                    "model": "text-embedding-3-small",
                    "api_key": api_key,
                    "max_retries": 3,
                    "request_timeout": 60  # 增加超時時間到60秒
                }
                
                if base_url:
                    embedding_params["base_url"] = base_url
                    print(f"🔧 使用自定義API端點: {base_url}")
                
                self.embeddings = OpenAIEmbeddings(**embedding_params)
                print("✅ OpenAI Embeddings初始化成功")
                
            else:
                # HuggingFace模型
                print(f"🔧 初始化HuggingFace Embeddings...")
                
                self.embeddings = HuggingFaceEmbeddings(
                    model_name="BAAI/bge-small-zh-v1.5",
                    model_kwargs={"device": "cpu"},
                    encode_kwargs={"batch_size": 16, "normalize_embeddings": True}
                )
                print("✅ HuggingFace Embeddings初始化成功")
                
        except Exception as e:
            print(f"❌ 嵌入模型初始化失敗: {e}")
            
            # 回退機制
            if self.model_type == "openai":
                print("📄 嘗試HuggingFace備選...")
                try:
                    self.embeddings = HuggingFaceEmbeddings(
                        model_name="BAAI/bge-small-zh-v1.5",
                        model_kwargs={"device": "cpu"}
                    )
                    self.model_type = "huggingface"
                    print("✅ 已回退到HuggingFace")
                except Exception as e2:
                    raise RuntimeError(f"所有嵌入模型都初始化失敗: {e2}")
            else:
                raise

    def _setup_text_processing(self):
        """設定文本處理組件"""
        self.normalizer = ChineseTextNormalizer()
        self.analyzer = SmartTextAnalyzer()
        print("✅ 文本處理組件初始化完成")

    def get_or_create_vectorstore(self, collection_name: str):
        """獲取或創建向量存儲 - PostgreSQL優先"""
        if collection_name not in self._vector_stores:
            try:
                if self.use_postgres and PGVECTOR_AVAILABLE:
                    # 🔧 使用PGVector
                    try:
                        from langchain_postgres import PGVector
                        self._vector_stores[collection_name] = PGVector(
                            embeddings=self.embeddings,
                            collection_name=collection_name,
                            connection=self.connection_string,
                            use_jsonb=True,
                        )
                    except ImportError:
                        from langchain_community.vectorstores import PGVector
                        self._vector_stores[collection_name] = PGVector(
                            connection=self.connection_string,  # ✅ 修正參數
                            collection_name=collection_name,
                            embeddings=self.embeddings,  # ✅ 修正參數
                            distance_strategy="cosine",
                            pre_delete_collection=False,
                            logger=logger
                        )
                    print(f"✅ PGVector向量存儲就緒: {collection_name}")
                else:
                    # 🔧 備用Chroma - 確保導入
                    if CHROMA_AVAILABLE:
                        from langchain_community.vectorstores import Chroma
                        self._vector_stores[collection_name] = Chroma(
                            collection_name=collection_name,
                            embedding_function=self.embeddings,
                            persist_directory=str(self.persist_dir)
                        )
                        print(f"✅ Chroma向量存儲就緒: {collection_name}")
                    else:
                        raise ImportError("Chroma不可用且PostgreSQL也不可用")
                        
            except Exception as e:
                logger.error(f"向量存儲創建失敗: {e}")
                raise RuntimeError(f"無法創建向量存儲: {e}")
        
        return self._vector_stores[collection_name]

    def _generate_doc_id(self, file_path: Path) -> str:
        """生成文檔ID"""
        # ✅ 修正語法錯誤
        content_hash = hashlib.md5(str(file_path).encode()).hexdigest()[:8]
        return f"doc_{file_path.stem}_{content_hash}"

    def load_document(self, file_path: Path) -> List[Document]:
        """載入並智能處理文檔"""
        try:
            extension = file_path.suffix.lower()
            
            # 根據檔案類型載入
            if extension == '.pdf':
                loader = PyPDFLoader(str(file_path))
                documents = loader.load()
                full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
            elif extension in ['.docx', '.doc'] and DOCX2TXT_AVAILABLE:
                if DOCX_METHOD == "docx2txt":
                    import docx2txt
                    full_text = docx2txt.process(str(file_path))
                else:
                    loader = Docx2txtLoader(str(file_path))
                    documents = loader.load()
                    full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
            elif extension == '.epub' and EPUB_AVAILABLE:
                # 🆕 EPUB處理邏輯
                epub_processor = EpubProcessor()
                full_text = epub_processor.extract_epub_content(file_path)
            elif extension == '.csv':
                loader = CSVLoader(str(file_path), encoding='utf-8')
                documents = loader.load()
                full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
            elif extension in ['.xlsx', '.xls']:
                loader = PandasExcelLoader(str(file_path), sheet_name=None)
                documents = loader.load()
                full_text = "\n\n".join([doc.page_content for doc in documents if doc.page_content])
            elif extension == '.json':
                loader = JSONLoader(str(file_path), jq_schema='.', text_content=False)
                documents = loader.load()
                full_text = "\n".join([doc.page_content for doc in documents if doc.page_content])
            else:
                # 嘗試自動檢測編碼
                encodings = ['utf-8', 'gbk', 'gb2312', 'big5']
                full_text = None
                
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            full_text = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                
                if full_text is None:
                    raise ValueError(f"無法解碼文件: {file_path}")
            
            if not full_text or not full_text.strip():
                logger.warning(f"文件內容為空或無法提取: {file_path.name}")
                return []
            
            # 生成文檔ID
            doc_id = self._generate_doc_id(file_path)
            
            # 使用優化的文本分割器
            source_info = {
                'file_path': str(file_path),
                'file_type': extension,
                'file_size': file_path.stat().st_size if file_path.exists() else 0
            }
            
            documents = self.text_splitter.smart_split_documents(full_text, doc_id, source_info)
            
            # 添加統一元數據
            for doc in documents:
                doc.metadata.update({
                    'source': str(file_path),
                    'filename': file_path.name,
                    'extension': extension,
                    'file_size': source_info['file_size'],
                    'load_timestamp': time.time()
                })
            
            print(f"📄 文檔載入完成: {file_path.name} ({len(documents)} 個分塊)")
            
            return documents
            
        except Exception as e:
            logger.error(f"載入文檔失敗 {file_path}: {e}")
            print(f"❌ 文檔載入失敗: {file_path.name} - {e}")
            return []

    def get_file_info(self, file_path: Path) -> Optional[FileInfo]:
        """獲取檔案基本信息"""
        try:
            if not file_path.exists():
                return None
                
            stat = file_path.stat()
            
            # 計算檔案哈希（用於變更檢測）
            hash_md5 = hashlib.md5()
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            
            return FileInfo(
                path=str(file_path),
                size=stat.st_size,
                mtime=stat.st_mtime,
                hash=hash_md5.hexdigest(),
                encoding='utf-8',
                file_type=file_path.suffix.lower()
            )
            
        except Exception as e:
            logger.error(f"獲取檔案信息失敗 {file_path}: {e}")
            return None

    def _parallel_load_documents(self, file_paths: List[Path], collection_name: str) -> List[Document]:
        """並發載入文檔"""
        all_documents = []
        max_workers = min(SYSTEM_CONFIG.get("max_workers", 4), len(file_paths))
        
        print(f"   並發載入 (工作線程: {max_workers})")
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_file = {
                executor.submit(self.load_document, file_path): file_path 
                for file_path in file_paths
            }
            
            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    documents = future.result()
                    if documents:
                        for doc in documents:
                            doc.metadata['collection'] = collection_name
                        all_documents.extend(documents)
                        print(f"   {file_path.name}: {len(documents)} 分塊")
                    else:
                        print(f"   {file_path.name}: 無有效內容")
                except Exception as e:
                    print(f"   {file_path.name}: {e}")
                    logger.error(f"並發載入失敗 {file_path}: {e}")
        
        return all_documents

    def _sequential_load_documents(self, file_paths: List[Path], collection_name: str) -> List[Document]:
        """順序載入文檔"""
        all_documents = []
        
        print(f"   順序載入")
        
        for i, file_path in enumerate(file_paths, 1):
            try:
                print(f"   [{i}/{len(file_paths)}] 處理: {file_path.name}")
                documents = self.load_document(file_path)
                
                if documents:
                    for doc in documents:
                        doc.metadata['collection'] = collection_name
                    all_documents.extend(documents)
                    print(f"      載入 {len(documents)} 個分塊")
                else:
                    print(f"      無有效內容")
                    
            except Exception as e:
                print(f"      載入失敗: {e}")
                logger.error(f"文件載入失敗 {file_path}: {e}")
        
        return all_documents

    def _load_file_records(self) -> Dict[str, Dict[str, FileInfo]]:
        """載入檔案記錄 - 加強錯誤處理和恢復機制"""
        record_file = self.data_dir / "file_records.json"
        
        # 🔧 檢查檔案是否存在
        if not record_file.exists():
            print("🔍 檔案記錄不存在，將建立新的記錄")
            return {}
        
        try:
            # 🔧 讀取並檢查檔案內容
            with open(record_file, 'r', encoding='utf-8') as f:
                content = f.read().strip()
            
            # 🔧 檢查檔案是否為空
            if not content:
                print("⚠️ 檔案記錄為空，將建立新的記錄")
                return {}
            
            # 🔧 檢查是否以{開頭（基本JSON格式檢查）
            if not content.startswith('{'):
                print(f"⚠️ 檔案記錄格式錯誤，內容開頭: {repr(content[:50])}")
                return self._handle_corrupted_records(record_file, content)
            
            # 🔧 嘗試解析JSON
            try:
                data = json.loads(content)
            except json.JSONDecodeError as json_error:
                print(f"❌ JSON解析失敗: {json_error}")
                print(f"   錯誤位置: line {json_error.lineno}, column {json_error.colno}")
                print(f"   檔案前100字符: {repr(content[:100])}")
                return self._handle_corrupted_records(record_file, content)
            
            # 🔧 驗證資料格式
            if not isinstance(data, dict):
                print(f"⚠️ 檔案記錄格式錯誤，應為字典但得到: {type(data)}")
                return {}
            
            # 🔧 轉換為FileInfo物件
            records = {}
            for collection, files in data.items():
                records[collection] = {}
                for file_path, info in files.items():
                    try:
                        if isinstance(info, dict):
                            fileinfo_fields = {
                                'path': info.get('path', file_path),
                                'size': info.get('size', 0),
                                'mtime': info.get('mtime', time.time()),
                                'hash': info.get('hash', ''),
                                'encoding': info.get('encoding', 'utf-8'),
                                'file_type': info.get('file_type', '')
                            }
                            
                            file_info_obj = FileInfo(**fileinfo_fields)
                            
                            # 🔧 恢復額外屬性
                            if 'uploaded_by' in info:
                                file_info_obj.uploaded_by = info['uploaded_by']
                            if 'uploaded_at' in info:
                                file_info_obj.uploaded_at = info['uploaded_at']
                            if 'file_source' in info:
                                file_info_obj.file_source = info['file_source']
                            
                            records[collection][file_path] = file_info_obj
                        else:
                            records[collection][file_path] = info
                            
                    except Exception as e:
                        logger.warning(f"載入檔案記錄失敗 {file_path}: {e}")
                        # 🔧 建立預設記錄
                        try:
                            default_info = FileInfo(
                                path=file_path,
                                size=0,
                                mtime=time.time(),
                                hash='',
                                encoding='utf-8',
                                file_type=''
                            )
                            records[collection][file_path] = default_info
                        except Exception:
                            logger.error(f"無法建立預設FileInfo for {file_path}")
                            continue
            
            print(f"✅ 檔案記錄載入成功: {len(records)} 個集合")
            return records
            
        except Exception as e:
            logger.error(f"載入檔案記錄失敗: {e}")
            print(f"❌ 嚴重錯誤，載入檔案記錄失敗: {e}")
            return self._handle_corrupted_records(record_file, "")

    def _save_file_records(self):
        """保存檔案記錄"""
        record_file = self.data_dir / "file_records.json"
        
        try:
            # 確保目錄存在
            record_file.parent.mkdir(parents=True, exist_ok=True)
            
            # 轉換FileInfo物件為字典
            serializable_records = {}
            for collection, files in self.file_records.items():
                serializable_records[collection] = {}
                for file_path, file_info in files.items():
                    if isinstance(file_info, FileInfo):
                        serializable_records[collection][file_path] = asdict(file_info)
                    else:
                        serializable_records[collection][file_path] = file_info
            
            # 先寫入臨時檔案，然後原子性替換
            temp_file = record_file.with_suffix('.tmp')
            with open(temp_file, 'w', encoding='utf-8') as f:
                json.dump(serializable_records, f, indent=2, ensure_ascii=False)
            
            temp_file.replace(record_file)
            logger.info(f"檔案記錄保存成功: {len(serializable_records)} 個集合")
            
        except Exception as e:
            logger.error(f"保存檔案記錄失敗: {e}")

    def _handle_corrupted_records(self, record_file: Path, content: str) -> Dict:
        """處理損壞的檔案記錄"""
        print(f"🔧 處理損壞的檔案記錄...")
        
        # 創建備份
        try:
            backup_file = record_file.with_suffix(f'.backup_{int(time.time())}')
            if record_file.exists():
                shutil.copy2(record_file, backup_file)
                print(f"   備份: {backup_file.name}")
        except Exception as e:
            logger.warning(f"創建備份失敗: {e}")
        
        # 嘗試恢復
        try:
            return self._rebuild_file_records()
        except Exception as e:
            logger.error(f"重建檔案記錄失敗: {e}")
            return {}

    def _rebuild_file_records(self) -> Dict:
        """重建檔案記錄"""
        print("🔧 重建檔案記錄...")
        
        new_records = {}
        
        if not self.data_dir.exists():
            return new_records
        
        # 掃描所有子目錄
        for subdir in self.data_dir.iterdir():
            if subdir.is_dir() and not subdir.name.startswith('.'):
                collection_name = f"collection_{subdir.name}"
                new_records[collection_name] = {}
                
                # 掃描目錄中的檔案
                for file_path in subdir.rglob('*'):
                    if (
                        file_path.is_file() and 
                        file_path.suffix.lower() in SUPPORTED_EXTENSIONS
                    ):
                        
                        file_info = self.get_file_info(file_path)
                        if file_info:
                            new_records[collection_name][str(file_path)] = file_info
        
        print(f"✅ 重建完成: {len(new_records)} 個集合")
        return new_records

    def get_file_source_statistics(self) -> Dict[str, Dict[str, int]]:
        """獲取檔案來源統計"""
        stats = {}
        
        for collection_name, files in self.file_records.items():
            stats[collection_name] = {
                'total': len(files),
                'upload': 0,
                'sync': 0,
                'unknown': 0
            }
            
            for file_info in files.values():
                source = getattr(file_info, 'file_source', 'unknown')
                if source in stats[collection_name]:
                    stats[collection_name][source] += 1
                else:
                    stats[collection_name]['unknown'] += 1
        
        return stats

    def diagnose_file_records(self) -> Dict:
        """診斷檔案記錄"""
        diagnosis = {
            'total_collections': len(self.file_records),
            'total_files': sum(len(files) for files in self.file_records.values()),
            'missing_files': 0,
            'outdated_files': 0,
            'corrupted_records': 0
        }
        
        for collection_name, files in self.file_records.items():
            for file_path, file_info in files.items():
                try:
                    path_obj = Path(file_path)
                    
                    if not path_obj.exists():
                        diagnosis['missing_files'] += 1
                        continue
                    
                    current_mtime = path_obj.stat().st_mtime
                    if abs(current_mtime - file_info.mtime) > 1:
                        diagnosis['outdated_files'] += 1
                        
                except Exception:
                    diagnosis['corrupted_records'] += 1
        
        return diagnosis

    def cleanup_invalid_records(self) -> Dict:
        """清理無效記錄"""
        removed_count = 0
        
        for collection_name in list(self.file_records.keys()):
            files = self.file_records[collection_name]
            
            for file_path in list(files.keys()):
                try:
                    if not Path(file_path).exists():
                        del files[file_path]
                        removed_count += 1
                except Exception:
                    del files[file_path]
                    removed_count += 1
            
            # 移除空集合
            if not files:
                del self.file_records[collection_name]
        
        if removed_count > 0:
            self._save_file_records()
        
        return {'removed_records': removed_count}

    def _process_batches(self, vectorstore: Union["Chroma", Any], batches: List[Tuple[List[Document], Dict]]) -> int:
        """處理批次向量化 - 完整的錯誤處理和元數據修復"""

        success_count = 0
        total_docs = sum(len(batch_docs) for batch_docs, _ in batches)
        
        print(f"\n開始批次向量化...")
        print(f"   總批次數: {len(batches)}")
        print(f"   總文檔數: {total_docs}")
        
        for batch_num, (batch_docs, batch_info) in enumerate(batches, 1):
            print(f"\n   批次 {batch_num}/{len(batches)}")
            print(f"      文檔數: {batch_info['documents']}")
            print(f"      tokens: {batch_info['tokens']:,}")
            print(f"      使用率: {(batch_info['tokens']/TOKEN_LIMITS['max_tokens_per_request']*100):.1f}%")
            
            # 顯示文檔類型分布
            type_info = ", ".join([f"{k}:{v}" for k, v in batch_info['types'].items()])
            print(f"      類型: {type_info}")
            
            start_time = time.time()
            
            try:
                print(f"      開始處理批次 {batch_num}...")
                print(f"      正在調用OpenAI API... (這可能需要30-60秒)")
                
                # 🛠️ 修復：統一處理元數據，確保類型正確
                safe_docs = []
                for doc in batch_docs:
                    safe_metadata = self._ensure_simple_metadata(doc.metadata)
                    safe_doc = Document(page_content=doc.page_content, metadata=safe_metadata)
                    safe_docs.append(safe_doc)

                print(f"      已處理 {len(safe_docs)} 個文檔的元數據，確保類型兼容")
                
                vectorstore.add_documents(safe_docs)
                processing_time = time.time() - start_time
                
                success_count += len(batch_docs)
                self.batch_processor.record_batch_result(True, processing_time)
                
                print(f"      批次 {batch_num} 完成 ({processing_time:.1f}s)")
                print(f"      總進度: {success_count}/{total_docs} ({success_count/total_docs*100:.1f}%)")
                
                # 批次間延遲
                if batch_num < len(batches):
                    delay = TOKEN_LIMITS["batch_delay"]
                    print(f"      等待 {delay} 秒...")
                    time.sleep(delay)
                    
            except Exception as e:
                processing_time = time.time() - start_time
                self.batch_processor.record_batch_result(False, processing_time)
                
                error_msg = str(e)
                print(f"      批次 {batch_num} 失敗 ({processing_time:.1f}s)")
                print(f"         錯誤: {error_msg}")
                
                # 🔧 特別處理元數據錯誤
                if "metadata" in error_msg.lower():
                    print(f"         檢測到元數據錯誤，嘗試更嚴格的處理...")
                    try:
                        # 重新嘗試，使用最嚴格的元數據過濾
                        ultra_safe_docs = []
                        for doc in batch_docs:
                            # 只保留最基本的元數據字段
                            minimal_metadata = {
                                'doc_id': str(doc.metadata.get('doc_id', 'unknown')),
                                'chunk_id': str(doc.metadata.get('chunk_id', 'unknown')),
                                'chunk_index': int(doc.metadata.get('chunk_index', 0)),
                                'text_type': str(doc.metadata.get('text_type', 'unknown')),
                                'source': str(doc.metadata.get('source', 'unknown')),
                                'filename': str(doc.metadata.get('filename', 'unknown')),
                                'token_count': int(doc.metadata.get('token_count', 0)),
                                'chunk_length': int(doc.metadata.get('chunk_length', 0)),
                                # URL處理
                                'contained_urls': str(doc.metadata.get('contained_urls', '')),
                                'url_count': int(doc.metadata.get('url_count', 0)),
                                'has_urls': bool(doc.metadata.get('has_urls', False))
                            }
                            
                            ultra_safe_doc = Document(
                                page_content=doc.page_content,
                                metadata=minimal_metadata
                            )
                            ultra_safe_docs.append(ultra_safe_doc)
                        
                        vectorstore.add_documents(ultra_safe_docs)
                        success_count += len(batch_docs)
                        print(f"         使用最小化元數據重新處理成功")
                        continue
                        
                    except Exception as retry_e:
                        print(f"         重新處理也失敗: {retry_e}")
                # 其他錯誤處理
                if "timeout" in error_msg.lower():
                    print(f"         超時錯誤，延長等待時間...")
                    time.sleep(30)
                elif "rate_limit" in error_msg.lower() or "429" in error_msg:
                    print(f"         速率限制，延長等待...")
                    time.sleep(60)
                elif "token" in error_msg.lower() and batch_info['documents'] > 1:
                    print(f"         Token超限，嘗試單個處理...")
                    single_success = self._process_documents_individually(vectorstore, batch_docs)
                    success_count += single_success
                elif "connection" in error_msg.lower():
                    print(f"         連接錯誤，等待重試...")
                    time.sleep(20)
                    try:
                        print(f"         重試批次 {batch_num}...")
                        # 使用安全的元數據重試
                        safe_docs = []
                        for doc in batch_docs:
                            safe_metadata = self._ensure_simple_metadata(doc.metadata)
                            safe_doc = Document(page_content=doc.page_content, metadata=safe_metadata)
                            safe_docs.append(safe_doc)
                        
                        vectorstore.add_documents(safe_docs)
                        success_count += len(batch_docs)
                        print(f"         重試成功")
                    except Exception as retry_e:
                        print(f"         重試失敗: {retry_e}")
                else:
                    print(f"         跳過此批次")
                    
                # 每次錯誤後添加額外延遲
                print(f"         錯誤後暫停10秒...")
                time.sleep(10)
        
        return success_count

    def _ensure_simple_metadata(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """確保元數據只包含Chroma支持的簡單類型：string, int, float, bool"""
        safe_metadata = {}
        
        for key, value in metadata.items():
            if isinstance(value, (str, int, float, bool)) or value is None:
                # ✅ 已經是Chroma支持的簡單類型，直接保留
                safe_metadata[key] = value
            elif isinstance(value, list):
                # ❌ 列表不被支持 → 轉為分隔符字符串
                if key == 'contained_urls' or 'url' in key.lower():
                    safe_metadata[key] = '|'.join(str(v) for v in value) if value else ''
                    safe_metadata[f'{key}_count'] = len(value)
                else:
                    safe_metadata[key] = '|'.join(str(v) for v in value) if value else ''
                    safe_metadata[f'{key}_count'] = len(value)
            elif isinstance(value, dict):
                # ❌ 字典不被支持 → 轉為JSON字符串
                safe_metadata[key] = json.dumps(value, ensure_ascii=False)
            else:
                # ❌ 其他類型不被支持 → 轉為字符串
                safe_metadata[key] = str(value)
        
        return safe_metadata

    def _process_documents_individually(self, vectorstore, documents: List[Document]) -> int:
        """單個處理文檔"""
        success_count = 0
        
        for i, doc in enumerate(documents):
            try:
                doc_tokens = doc.metadata.get('token_count', 0)
                if doc_tokens > TOKEN_LIMITS['max_tokens_per_request']:
                    print(f"         文檔 {i+1} 仍然過大 ({doc_tokens:,} tokens)，跳過")
                    continue
                
                vectorstore.add_documents([doc])
                success_count += 1
                print(f"         單個文檔 {i+1}/{len(documents)} 完成")
                time.sleep(1)  # 單個處理時短暫延遲
                
            except Exception as e:
                print(f"         單個文檔 {i+1} 失敗: {e}")
        
        return success_count

    def incremental_update(self, collection_name: str, added_files: List[Path], 
                          modified_files: List[Path], deleted_files: List[str],
                          current_files: Dict[str, FileInfo]) -> bool:
        """🚀 優化版增量更新"""
        with self.processing_lock:
            try:
                vectorstore = self.get_or_create_vectorstore(collection_name)
                
                # 處理刪除和修改
                files_to_delete = deleted_files + [str(f) for f in modified_files]
                if files_to_delete:
                    for file_path in files_to_delete:
                        try:
                            vectorstore.delete(filter={"source": file_path})
                            print(f"已刪除: {Path(file_path).name}")
                        except Exception as e:
                            logger.warning(f"刪除文檔失敗 {file_path}: {e}")
                
                # 處理新增和修改的文件
                files_to_process = added_files + modified_files
                if not files_to_process:
                    print("   無新文件需要處理")
                    return True
                
                print(f"處理 {len(files_to_process)} 個文件...")
                print(f"   處理大文件可能需要較長時間，請耐心等待...")
                
                # 並發載入文檔
                all_documents = []
                if PERFORMANCE_CONFIG.get("parallel_processing", True):
                    all_documents = self._parallel_load_documents(files_to_process, collection_name)
                else:
                    all_documents = self._sequential_load_documents(files_to_process, collection_name)
                
                if not all_documents:
                    print("   沒有有效文檔需要向量化")
                    return True
                
                # 統計和成本估算
                total_tokens = sum(doc.metadata.get('token_count', 0) for doc in all_documents)
                estimated_cost = self.batch_processor.token_estimator.estimate_embedding_cost(total_tokens)
                
                print(f"\n向量化統計:")
                print(f"   總分塊數: {len(all_documents)}")
                print(f"   總tokens: {total_tokens:,}")
                print(f"   估算成本: ${estimated_cost:.4f}")
                
                # 創建智能批次並處理
                batches = self.batch_processor.create_smart_batches(all_documents)
                success_count = self._process_batches(vectorstore, batches)
                
                print(f"\n向量化完成！")
                print(f"   成功: {success_count}/{len(all_documents)} 個分塊")
                print(f"   成功率: {(success_count/len(all_documents)*100):.1f}%")
                
                # 更新文件記錄
                self.file_records[collection_name] = current_files
                self._save_file_records()
                
                # 記憶體清理
                if success_count % PERFORMANCE_CONFIG.get("gc_frequency", 50) == 0:
                    gc.collect()
                
                return success_count > 0
                
            except Exception as e:
                logger.error(f"增量更新失敗 {collection_name}: {e}")
                print(f"❌ 系統錯誤: {e}")
                return False

    def get_collection_name(self, dir_path: Path) -> str:
        """獲取集合名稱"""
        try:
            relative_path = dir_path.relative_to(self.data_dir)
            if relative_path.parts:
                return f"collection_{relative_path.parts[0]}"
        except ValueError:
            pass
        return "collection_other"

    def sync_collections(self) -> int:
        """純PostgreSQL方案：不再掃描本地目錄"""
        print("純PostgreSQL方案：跳過本地目錄掃描")
        print("所有數據都在PostgreSQL中，無需同步")
        return 0

    def scan_directory_changes(self, dir_path: Path, collection_name: str) -> Tuple[List[Path], List[Path], List[str], Dict[str, FileInfo]]:
        """掃描目錄變更 - 修正版：正確處理上傳檔案"""
        current_files = {}
        
        print(f"掃描目錄: {dir_path}")
        
        # 遞歸掃描目錄
        file_count = 0
        for file_path in dir_path.rglob('*'):
            if (
                file_path.is_file() and 
                file_path.suffix.lower() in SUPPORTED_EXTENSIONS and
                not file_path.name.startswith('.') and
                file_path.stat().st_size > 0
            ):  # 跳過空文件
                
                file_info = self.get_file_info(file_path)
                if file_info:
                    # 修正：使用標準化的絕對路徑作為鍵值
                    try:
                        # 使用absolute()避免符號連結問題
                        absolute_path = str(file_path.absolute())
                        current_files[absolute_path] = file_info
                        file_count += 1
                    except Exception as e:
                        logger.warning(f"路徑標準化失敗 {file_path}: {e}")
                        # 回退到原始路徑
                        current_files[str(file_path)] = file_info
                        file_count += 1
        
        print(f"找到 {file_count} 個有效檔案")
        
        old_files = self.file_records.get(collection_name, {})
        print(f"舊記錄中有 {len(old_files)} 個檔案")
        
        # 修正：正規化舊記錄的路徑鍵值
        normalized_old_files = {}
        normalization_errors = 0
        
        for old_path, old_info in old_files.items():
            try:
                old_path_obj = Path(old_path)
                
                if old_path_obj.is_absolute():
                    # 已經是絕對路徑
                    normalized_key = str(old_path_obj.absolute())
                else:
                    # 相對路徑轉絕對路徑
                    try:
                        abs_path = (dir_path / old_path).absolute()
                        normalized_key = str(abs_path)
                    except Exception:
                        # 如果無法轉換，保持原樣
                        normalized_key = old_path
                        
                normalized_old_files[normalized_key] = old_info
                
            except Exception as e:
                logger.warning(f"舊路徑正規化失敗 {old_path}: {e}")
                # 保持原路徑
                normalized_old_files[old_path] = old_info
                normalization_errors += 1
        
        if normalization_errors > 0:
            print(f"{normalization_errors} 個舊路徑正規化失敗")
        
        # 修正：智能變更檢測
        added_files = []
        modified_files = []
        
        print("檢測變更...")
        
        for file_path, file_info in current_files.items():
            current_file_name = Path(file_path).name
            current_hash = file_info.hash
            
            # 首先嘗試精確路徑匹配
            if file_path in normalized_old_files:
                old_info = normalized_old_files[file_path]
                if old_info.hash != current_hash:
                    modified_files.append(Path(file_path))
                    print(f"修改檔案: {current_file_name}")
            else:
                # 智能檔案匹配：檢查是否是同一檔案的不同路徑表示
                file_found = False
                
                for old_path, old_info in normalized_old_files.items():
                    old_file_name = Path(old_path).name
                    
                    # 檔案名相同且哈希相同 = 同一檔案
                    if (
                        current_file_name == old_file_name and 
                        current_hash == old_info.hash
                    ):
                        file_found = True
                        print(f"路徑變更但內容相同: {current_file_name}")
                        break
                        
                    # 檔案名相同但哈希不同 = 檔案被修改
                    elif (
                        current_file_name == old_file_name and 
                        current_hash != old_info.hash
                    ):
                        modified_files.append(Path(file_path))
                        file_found = True
                        print(f"修改檔案 (路徑變更): {current_file_name}")
                        break
                
                if not file_found:
                    added_files.append(Path(file_path))
                    print(f"新檔案: {current_file_name}")
        
        # 修正：智能刪除檢測
        deleted_files = []
        
        for old_path in normalized_old_files.keys():
            old_file_name = Path(old_path).name
            
            if old_path not in current_files:
                # 檢查檔案是否真的不存在（可能只是路徑表示不同）
                file_still_exists = False
                
                for current_path in current_files.keys():
                    current_file_name = Path(current_path).name
                    if current_file_name == old_file_name:
                        # 進一步檢查是否是同一檔案（通過內容哈希）
                        current_hash = current_files[current_path].hash
                        old_hash = normalized_old_files[old_path].hash
                        
                        if current_hash == old_hash:
                            file_still_exists = True
                            break
                
                if not file_still_exists:
                    deleted_files.append(old_path)
                    print(f"刪除檔案: {old_file_name}")
        
        print(f"變更統計:")
        print(f"   新增: {len(added_files)}")
        print(f"   修改: {len(modified_files)}")
        print(f"   刪除: {len(deleted_files)}")
        
        return added_files, modified_files, deleted_files, current_files


# ✅ 確保所有import都能正常工作
__all__ = ['VectorOperationsCore']